from __future__ import annotations

import io
import random
import tempfile
import time
import unittest
import zipfile
from datetime import date
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from fastapi import HTTPException

from grievance_api.db.db import Db
from grievance_api.db.migrate import migrate
from grievance_api.services import pay_portal
from grievance_api.services.pay_portal import (
    PayActor,
    approve_irs_rate_candidate,
    calculate_commission_compensation,
    calculate_pay_profile_snapshot,
    calculate_president_differential,
    create_mileage_attachment,
    create_revision,
    handle_pay_docuseal_completion,
    list_irs_rate_candidates,
    lock_period_and_send_packet,
    parse_irs_mileage_rate_candidates,
    pay_profile_by_email,
    remove_mileage_attachment,
    delete_pay_entry,
    pay_settings,
    pay_packet_signer_order,
    store_attachment,
    store_compensation_stub,
    sync_irs_mileage_rate_candidates,
    upsert_entry,
    upsert_pay_profile,
    upsert_wage_scale,
    validate_receipt_file,
)
from grievance_api.services.sharepoint_graph import DirectoryUserRef, GraphUploader
from grievance_api.web.routes_pay import (
    InternalRoleAssignmentRequest,
    PayEntryCorrectionRequest,
    PayEntryReviewRequest,
    PayInternalUserImportRequest,
    PayDemoArtifactRequest,
    PayDemoFeedbackStatusRequest,
    PayDemoFeedbackRequest,
    PayDemoSettingsUpdateRequest,
    PayProfileUpsertRequest,
    approve_pay_irs_rate,
    create_pay_entry_correction_route,
    create_pay_profile,
    delete_pay_attachment,
    delete_pay_entry_route,
    remove_pay_profile,
    download_pay_attachment,
    review_pay_entry_route,
    generate_pay_demo_output_files,
    import_pay_internal_users,
    pay_context,
    pay_demo_context,
    pay_directory_users,
    pay_page,
    pay_start_page,
    pay_view_page,
    remove_pay_internal_role,
    save_pay_internal_role,
    submit_pay_demo_feedback,
    update_pay_demo_feedback,
    update_pay_demo_settings,
)


class _FakeGraph:
    def __init__(self) -> None:
        self.calls: list[dict[str, object]] = []

    def upload_local_file_to_folder_path(self, **kwargs):  # noqa: ANN003
        self.calls.append(dict(kwargs))
        return SimpleNamespace(
            web_url=f"https://sharepoint.local/{kwargs['folder_path']}/{kwargs['filename']}",
            path=f"{kwargs['folder_path']}/{kwargs['filename']}",
        )


class _FakeDocuSeal:
    def __init__(self) -> None:
        self.create_calls: list[dict[str, object]] = []
        self.download_calls: list[str] = []

    def create_submission(self, **kwargs):  # noqa: ANN003
        self.create_calls.append(dict(kwargs))
        return SimpleNamespace(
            submission_id="pay-ds-1",
            signing_link="https://docuseal.local/sign/pay-ds-1",
            raw={},
        )

    def extract_signing_links_by_email(self, raw):  # noqa: ANN001
        _ = raw
        if not self.create_calls:
            return {}
        return {
            str(email).lower(): f"https://docuseal.local/sign/{index}"
            for index, email in enumerate(self.create_calls[-1].get("signers", []), start=1)
        }

    def fetch_signing_links_by_email(self, submission_id: str):
        _ = submission_id
        return {}

    def download_completed_artifacts(self, *, submission_id: str):
        self.download_calls.append(submission_id)
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("signed.pdf", b"%PDF-1.4\nsigned packet\n")
        return {"completed_zip_bytes": buf.getvalue(), "submission": {}}


class _FakeMailer:
    def __init__(self) -> None:
        self.calls: list[dict[str, object]] = []

    def send_mail(self, **kwargs):  # noqa: ANN003
        self.calls.append(dict(kwargs))
        return SimpleNamespace(graph_message_id="mail-1", internet_message_id=None)


class _FakeHttpResponse:
    def __init__(self, text: str, status_code: int = 200) -> None:
        self.text = text
        self.status_code = status_code


class _InlineExecutor:
    def submit(self, fn, **kwargs):  # noqa: ANN001
        fn(**kwargs)
        return SimpleNamespace(done=lambda: True)


class _DirectoryGraphStub:
    def __init__(self, rows: list[DirectoryUserRef], licensed_rows: list[DirectoryUserRef] | None = None) -> None:
        self.rows = rows
        self.licensed_rows = licensed_rows if licensed_rows is not None else rows
        self.last_search: str | None = None
        self.last_limit: int | None = None
        self.last_import_limit: int | None = None

    def search_directory_users(self, search_text: str, *, limit: int = 10) -> list[DirectoryUserRef]:
        self.last_search = search_text
        self.last_limit = limit
        return list(self.rows)

    def list_licensed_directory_users(self, *, limit: int = 999) -> list[DirectoryUserRef]:
        self.last_import_limit = limit
        return list(self.licensed_rows)[:limit]


class _LicensedGraphUploader(GraphUploader):
    def __init__(self) -> None:
        super().__init__(
            tenant_id="tenant",
            client_id="client",
            cert_thumbprint="thumb",
            cert_pem_path="/tmp/missing.pem",
        )
        self.calls: list[tuple[str, str, dict[str, object] | None]] = []

    def _request(self, method: str, endpoint: str, **kwargs):  # noqa: ANN003
        self.calls.append((method, endpoint, kwargs.get("params")))
        if endpoint == "/subscribedSkus":
            return {
                "value": [
                    {
                        "skuId": "paid-sku",
                        "skuPartNumber": "O365_BUSINESS_PREMIUM",
                        "capabilityStatus": "Enabled",
                        "prepaidUnits": {"enabled": 12},
                    },
                    {
                        "skuId": "free-sku",
                        "skuPartNumber": "POWER_BI_STANDARD",
                        "capabilityStatus": "Enabled",
                        "prepaidUnits": {"enabled": 200},
                    },
                    {
                        "skuId": "trial-sku",
                        "skuPartNumber": "TEAMS_EXPLORATORY",
                        "capabilityStatus": "Enabled",
                        "prepaidUnits": {"enabled": 50},
                    },
                    {
                        "skuId": "suspended-paid-sku",
                        "skuPartNumber": "ENTERPRISEPACK",
                        "capabilityStatus": "Suspended",
                        "prepaidUnits": {"enabled": 8},
                    },
                ]
            }
        if endpoint == "/users":
            return {
                "value": [
                    {
                        "id": "user-paid",
                        "displayName": "Paid User",
                        "mail": "paid@example.org",
                        "userPrincipalName": "paid@example.org",
                        "accountEnabled": True,
                        "assignedLicenses": [{"skuId": "paid-sku"}],
                    },
                    {
                        "id": "user-free",
                        "displayName": "Free User",
                        "mail": "free@example.org",
                        "userPrincipalName": "free@example.org",
                        "accountEnabled": True,
                        "assignedLicenses": [{"skuId": "free-sku"}],
                    },
                    {
                        "id": "user-trial",
                        "displayName": "Trial User",
                        "mail": "trial@example.org",
                        "userPrincipalName": "trial@example.org",
                        "accountEnabled": True,
                        "assignedLicenses": [{"skuId": "trial-sku"}],
                    },
                    {
                        "id": "user-disabled",
                        "displayName": "Disabled User",
                        "mail": "disabled@example.org",
                        "userPrincipalName": "disabled@example.org",
                        "accountEnabled": False,
                        "assignedLicenses": [{"skuId": "paid-sku"}],
                    },
                    {
                        "id": "user-both",
                        "displayName": "Paid Plus Free",
                        "mail": "both@example.org",
                        "userPrincipalName": "both@example.org",
                        "accountEnabled": True,
                        "assignedLicenses": [{"skuId": "free-sku"}, {"skuId": "paid-sku"}],
                    },
                ]
            }
        raise AssertionError(f"unexpected Graph endpoint {endpoint}")


class _Request:
    def __init__(self, *, state, session: dict | None = None) -> None:  # noqa: ANN001
        self.app = SimpleNamespace(state=state)
        self.session = session if session is not None else {}
        self.headers = {"host": "127.0.0.1"}
        self.client = SimpleNamespace(host="127.0.0.1")


def _fake_docx_to_pdf(docx_path: str, out_dir: str, timeout_seconds: int, **kwargs) -> str:  # noqa: ANN001, ANN003
    _ = timeout_seconds
    _ = kwargs
    out = Path(out_dir) / f"{Path(docx_path).stem}.pdf"
    out.write_bytes(f"PDF:{Path(docx_path).name}\n".encode("utf-8"))
    return str(out)


def _fake_fill_docx(**kwargs) -> None:  # noqa: ANN003
    Path(kwargs["output_path"]).parent.mkdir(parents=True, exist_ok=True)
    Path(kwargs["output_path"]).write_text(
        "anchor" if kwargs.get("include_signature_placeholders") else "voucher",
        encoding="utf-8",
    )


def _fake_merge(input_paths: list[str], output_path: str) -> None:
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    payload = "|".join(Path(path).name for path in input_paths)
    Path(output_path).write_bytes(payload.encode("utf-8"))


def _fake_watermark(input_path: str, output_path: str | None = None, *, text: str = "DEMO") -> None:
    target = Path(output_path or input_path)
    if output_path:
        target.write_bytes(Path(input_path).read_bytes())
    target.write_text(target.read_text(encoding="utf-8") + f"|WATERMARK:{text}", encoding="utf-8")


class PayPortalTests(unittest.IsolatedAsyncioTestCase):
    async def asyncSetUp(self) -> None:
        self.tmpdir = tempfile.TemporaryDirectory()
        self.addAsyncCleanup(self._cleanup_tmpdir)
        self.db_path = str(Path(self.tmpdir.name) / "test.sqlite3")
        self.data_root = str(Path(self.tmpdir.name) / "data")
        migrate(self.db_path)
        self.db = Db(self.db_path)

    async def _cleanup_tmpdir(self) -> None:
        self.tmpdir.cleanup()

    def _cfg(self):
        return SimpleNamespace(
            data_root=self.data_root,
            libreoffice_timeout_seconds=45,
            docx_pdf_engine="libreoffice",
            docx_pdf_graph_temp_folder="_docx_pdf_convert",
            pay_portal=SimpleNamespace(
                enabled=True,
                voucher_template_path=str(Path(self.tmpdir.name) / "voucher.docx"),
                sharepoint_root_folder="Pay Portal",
                receipt_max_file_bytes=1_000_000,
                receipt_max_entry_bytes=5_000_000,
                clamav_host="clamav",
                clamav_port=3310,
                clamav_timeout_seconds=30,
                president_email="",
                treasurer_emails=(),
                president_target_scale="36",
                president_target_multiplier=1.20,
                google_maps_api_key="",
                irs_rates={"2025": "0.70", "2026": "0.725"},
                irs_rate_sync_enabled=True,
                irs_rate_source_urls=("https://www.irs.gov/example",),
                common_places=(),
                common_places_sharepoint_folder="Local3106/Mileage/Config",
                pay_users=(),
            ),
            officer_auth=SimpleNamespace(enabled=False),
            external_steward_auth=SimpleNamespace(enabled=False),
            graph=SimpleNamespace(
                site_hostname="contoso.sharepoint.com",
                site_path="/sites/local",
                document_library="Documents",
            ),
            docuseal=SimpleNamespace(default_template_id=1, template_ids={}, strict_template_ids=False),
            email=SimpleNamespace(
                enabled=True,
                internal_recipients=("treasurer@example.org",),
                max_attachment_bytes=2_000_000,
            ),
        )

    @staticmethod
    def _staff_session(role: str, *, email: str = "admin@example.org") -> dict[str, object]:
        return {
            "officer_user": {
                "email": email,
                "display_name": email,
                "role": role,
                "contract_scopes": [],
                "exp": int(time.time()) + 3600,
            }
        }

    async def _save_profile(
        self,
        *,
        email: str,
        display_name: str | None = None,
        pay_basis: str = "hourly",
        base_wage_input_type: str = "hourly",
        base_wage_amount: float = 45,
        weekly_basis_hours: float = 40,
        commission_month_1_amount: float = 0,
        commission_month_2_amount: float = 0,
        commission_month_3_amount: float = 0,
    ) -> dict[str, object]:
        return await upsert_pay_profile(
            self.db,
            principal_id=None,
            principal_email=email,
            principal_display_name=display_name or email,
            pay_basis=pay_basis,
            base_wage_input_type=base_wage_input_type,
            base_wage_amount=base_wage_amount,
            weekly_basis_hours=weekly_basis_hours,
            commission_month_1_amount=commission_month_1_amount,
            commission_month_2_amount=commission_month_2_amount,
            commission_month_3_amount=commission_month_3_amount,
            status="active",
            notes=None,
            updated_by="test",
        )

    async def _open_period_and_entry(self, *, actor: PayActor | None = None) -> tuple[str, str]:
        actor = actor or PayActor(
            email="officer@example.org",
            display_name="Officer Example",
            role="officer",
            can_view_all=True,
            can_edit_all=False,
            can_lock=False,
        )
        await self._save_profile(email=actor.email, display_name=actor.display_name)
        period = await pay_portal.ensure_pay_period(self.db)
        entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={
                "entry_date": str(period["period_start"]),
                "display_name": actor.display_name,
                "hours": 8,
                "local_number": "3106",
                "notes": "meeting",
            },
            pay_cfg=self._cfg().pay_portal,
        )
        return str(period["id"]), str(entry["id"])

    async def test_pay_settings_falls_back_to_configured_mileage_defaults(self) -> None:
        cfg = SimpleNamespace(
            president_email="president@example.org",
            treasurer_emails=("treasurer@example.org",),
            irs_rates={"2024": "0.67", "2025": "0.70", "2026": "0.725"},
            common_places=(
                {"label": "Union Hall", "address": "4076 Union Hall Pl, Jacksonville, FL 32205, USA"},
            ),
        )
        settings = await pay_portal.pay_settings(self.db, pay_cfg=cfg)
        self.assertEqual(settings["president_email"], "president@example.org")
        self.assertEqual(settings["treasurer_emails"], ["treasurer@example.org"])
        self.assertEqual(settings["irs_rates"], {"2024": "0.67", "2025": "0.70", "2026": "0.725"})
        demo_settings = await pay_portal.pay_demo_settings(self.db)
        self.assertTrue(demo_settings["demo_mode_enabled"])
        self.assertEqual(demo_settings["demo_cycle_title"], "Training Demo Cycle")
        self.assertIn(
            {"label": "Union Hall", "address": "4076 Union Hall Pl, Jacksonville, FL 32205, USA"},
            settings["common_places"],
        )
        self.assertIn(
            {"label": "State Road 16", "address": "4875 State Road 16, St. Augustine, FL 32092, USA"},
            settings["common_places"],
        )

    async def test_legacy_common_places_pick_up_seeded_defaults(self) -> None:
        cfg = SimpleNamespace(
            president_email="",
            treasurer_emails=(),
            irs_rates={},
            common_places=(),
        )
        await self.db.upsert_app_setting(
            setting_key="pay_portal",
            setting={"common_places": [{"label": "Legacy", "address": "1 Existing St"}]},
            updated_by="test",
        )
        settings = await pay_portal.pay_settings(self.db, pay_cfg=cfg)
        self.assertIn({"label": "Legacy", "address": "1 Existing St"}, settings["common_places"])
        self.assertIn(
            {"label": "State Road 16", "address": "4875 State Road 16, St. Augustine, FL 32092, USA"},
            settings["common_places"],
        )

    async def test_pay_settings_common_places_can_be_overridden_for_removal(self) -> None:
        cfg = SimpleNamespace(
            president_email="",
            treasurer_emails=(),
            irs_rates={},
            common_places=(),
        )
        saved = await pay_portal.save_pay_settings(
            self.db,
            setting={"common_places": [{"label": "Only", "address": "1 Test St"}]},
            updated_by="test",
            pay_cfg=cfg,
        )
        settings = await pay_portal.pay_settings(self.db, pay_cfg=cfg)
        self.assertEqual(saved["common_places"], [{"label": "Only", "address": "1 Test St"}])
        self.assertEqual(settings["common_places"], [{"label": "Only", "address": "1 Test St"}])

    async def test_pay_demo_feedback_is_saved_for_training_review(self) -> None:
        actor = PayActor(
            email="officer@example.org",
            display_name="Officer Example",
            role="officer",
            can_view_all=True,
            can_edit_all=False,
            can_lock=False,
        )
        saved = await pay_portal.create_pay_demo_feedback(
            self.db,
            actor=actor,
            screen="mileage",
            category="missing",
            demo_step=2,
            demo_cycle_title="Officer Training",
            comment="Show the generated mileage PDF before attaching it.",
        )
        rows = await pay_portal.list_pay_demo_feedback(self.db)

        self.assertEqual(saved["screen"], "mileage")
        self.assertEqual(saved["category"], "missing")
        self.assertEqual(saved["demo_step"], 2)
        self.assertEqual(saved["demo_cycle_title"], "Officer Training")
        self.assertEqual(rows[0]["actor_email"], "officer@example.org")
        self.assertEqual(rows[0]["comment"], "Show the generated mileage PDF before attaching it.")

        closed = await pay_portal.update_pay_demo_feedback_status(self.db, feedback_id=saved["id"], status="closed")
        self.assertEqual(closed["status"], "closed")

    async def test_pay_demo_feedback_api_respects_demo_mode_setting(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer"),
        )
        saved = await submit_pay_demo_feedback(
            PayDemoFeedbackRequest(
                screen="demo",
                category="suggestion",
                demo_step=4,
                demo_cycle_title="May training",
                comment="Add a printable checklist.",
            ),
            request,
        )
        self.assertEqual(saved["category"], "suggestion")
        self.assertEqual(saved["demo_step"], 4)

        admin_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("admin"),
        )
        updated = await update_pay_demo_settings(
            PayDemoSettingsUpdateRequest(demo_mode_enabled=False, demo_cycle_title="Paused Demo"),
            admin_request,
        )
        self.assertFalse(updated["demo_mode_enabled"])
        with self.assertRaises(HTTPException) as denied:
            await submit_pay_demo_feedback(
                PayDemoFeedbackRequest(screen="demo", category="suggestion", comment="Try while disabled."),
                request,
            )
        self.assertEqual(denied.exception.status_code, 400)

        reopened = await update_pay_demo_feedback(
            int(saved["id"]),
            PayDemoFeedbackStatusRequest(status="closed"),
            admin_request,
        )
        self.assertEqual(reopened["status"], "closed")

    async def test_mileage_attachment_updates_entry_totals_and_filename(self) -> None:
        period_id, entry_id = await self._open_period_and_entry()
        actor = PayActor(
            email="officer@example.org",
            display_name="Officer Example",
            role="officer",
            can_view_all=True,
            can_edit_all=False,
            can_lock=False,
        )

        with patch(
            "grievance_api.services.pay_portal.build_mileage_pdf",
            return_value=(b"%PDF-1.4\nmileage\n", pay_portal.Decimal("10.88"), pay_portal.Decimal("15.00")),
        ) as build_pdf:
            result = await create_mileage_attachment(
                db=self.db,
                cfg=self._cfg(),
                period_id=period_id,
                entry_id=entry_id,
                actor=actor,
                name="test test",
                local_number="3106",
                date_str="2026-05-09",
                description="Union business",
                locations=["Union Hall", "Worksite"],
                rate_text="0.725",
            )

        self.assertEqual(result["filename"], "20260509 test test.pdf")
        self.assertEqual(result["attachment_type"], "mileage_pdf")
        self.assertEqual(result["mileage_miles"], 15.0)
        self.assertEqual(result["reimbursement"], 10.88)
        build_pdf.assert_called_once()
        row = await self.db.fetchone(
            "SELECT mileage_miles, mileage_rate, mileage_amount FROM pay_entries WHERE id=?",
            (entry_id,),
        )
        self.assertEqual(row[0], 15.0)
        self.assertEqual(row[1], 0.725)
        self.assertEqual(row[2], 10.88)
        attachments = await pay_portal.list_attachments(self.db, period_id=period_id, actor=actor)
        mileage = [row for row in attachments if row["attachment_type"] == "mileage_pdf"][0]
        self.assertEqual(mileage["mileage_miles"], 15.0)
        self.assertEqual(mileage["mileage_rate"], 0.725)
        self.assertEqual(mileage["mileage_amount"], 10.88)
        self.assertTrue(mileage["can_remove"])

    async def test_mileage_report_download_and_remove_are_permissioned(self) -> None:
        period_id, entry_id = await self._open_period_and_entry()
        actor = PayActor(
            email="officer@example.org",
            display_name="Officer Example",
            role="officer",
            can_view_all=True,
            can_edit_all=False,
            can_lock=False,
        )
        with patch(
            "grievance_api.services.pay_portal.build_mileage_pdf",
            return_value=(b"%PDF-1.4\nmileage\n", pay_portal.Decimal("10.88"), pay_portal.Decimal("15.00")),
        ):
            result = await create_mileage_attachment(
                db=self.db,
                cfg=self._cfg(),
                period_id=period_id,
                entry_id=entry_id,
                actor=actor,
                name="test test",
                local_number="3106",
                date_str="2026-05-09",
                description="Union business",
                locations=["Union Hall", "Worksite"],
                rate_text="0.725",
            )

        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        owner_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="officer@example.org"),
        )
        download = await download_pay_attachment(str(result["id"]), owner_request)
        self.assertEqual(download.filename, "20260509 test test.pdf")

        other_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="other@example.org"),
        )
        with self.assertRaises(HTTPException) as denied:
            await download_pay_attachment(str(result["id"]), other_request)
        self.assertEqual(denied.exception.status_code, 403)

        removed = await delete_pay_attachment(str(result["id"]), owner_request)
        self.assertEqual(removed["mileage_amount"], 10.88)
        row = await self.db.fetchone(
            "SELECT mileage_miles, mileage_rate, mileage_amount FROM pay_entries WHERE id=?",
            (entry_id,),
        )
        self.assertEqual(row[0], 0.0)
        self.assertEqual(row[1], 0.0)
        self.assertEqual(row[2], 0.0)
        visible = await pay_portal.list_attachments(self.db, period_id=period_id, actor=actor)
        self.assertEqual([row for row in visible if row["attachment_type"] == "mileage_pdf"], [])

    async def test_mileage_report_remove_denies_locked_and_unsafe_legacy_reports(self) -> None:
        period_id, entry_id = await self._open_period_and_entry()
        actor = PayActor("officer@example.org", "Officer Example", "officer", True, False, False)
        legacy = await pay_portal.store_attachment(
            self.db,
            cfg=self._cfg(),
            period_id=period_id,
            entry_id=entry_id,
            actor=actor,
            attachment_type="mileage_pdf",
            filename="legacy-mileage.pdf",
            content_type="application/pdf",
            content=b"%PDF-1.4\nlegacy\n%%EOF",
            scan=False,
        )
        await self.db.exec(
            "UPDATE pay_entries SET mileage_miles=?, mileage_rate=?, mileage_amount=? WHERE id=?",
            (12.0, 0.725, 8.70, entry_id),
        )
        removed = await remove_mileage_attachment(self.db, attachment_id=str(legacy["id"]), actor=actor)
        self.assertEqual(removed["mileage_amount"], 8.7)

        period_id, entry_id = await self._open_period_and_entry()
        first = await pay_portal.store_attachment(
            self.db,
            cfg=self._cfg(),
            period_id=period_id,
            entry_id=entry_id,
            actor=actor,
            attachment_type="mileage_pdf",
            filename="legacy-mileage-1.pdf",
            content_type="application/pdf",
            content=b"%PDF-1.4\nlegacy1\n%%EOF",
            scan=False,
        )
        second = await pay_portal.store_attachment(
            self.db,
            cfg=self._cfg(),
            period_id=period_id,
            entry_id=entry_id,
            actor=actor,
            attachment_type="mileage_pdf",
            filename="legacy-mileage-2.pdf",
            content_type="application/pdf",
            content=b"%PDF-1.4\nlegacy2\n%%EOF",
            scan=False,
        )
        await self.db.exec(
            "UPDATE pay_entries SET mileage_miles=?, mileage_rate=?, mileage_amount=? WHERE id=?",
            (12.0, 0.725, 8.70, entry_id),
        )
        forms = await pay_portal.list_attachments(self.db, period_id=period_id, actor=actor)
        first_form = next(row for row in forms if row["id"] == first["id"])
        self.assertTrue(first_form["legacy_summary"])
        self.assertTrue(first_form["can_remove"])
        self.assertEqual(first_form["mileage_miles"], 12.0)
        self.assertEqual(first_form["mileage_rate"], 0.725)
        self.assertEqual(first_form["mileage_amount"], 8.7)

        duplicate_removed = await remove_mileage_attachment(self.db, attachment_id=str(first["id"]), actor=actor)
        self.assertFalse(duplicate_removed["subtracted_entry_totals"])
        row = await self.db.fetchone(
            "SELECT mileage_miles, mileage_rate, mileage_amount FROM pay_entries WHERE id=?",
            (entry_id,),
        )
        self.assertEqual(row[0], 12.0)
        self.assertEqual(row[1], 0.725)
        self.assertEqual(row[2], 8.7)

        await self.db.exec("UPDATE pay_entries SET locked_at_utc=? WHERE id=?", (pay_portal.utcnow(), entry_id))
        locked_actor = PayActor("officer@example.org", "Officer Example", "officer", True, False, False)
        with self.assertRaisesRegex(ValueError, "entry is locked"):
            await remove_mileage_attachment(self.db, attachment_id=str(second["id"]), actor=locked_actor)

    async def test_public_pay_start_page_has_login_links_without_portal_data(self) -> None:
        response = await pay_start_page(_Request(state=SimpleNamespace(cfg=self._cfg(), db=self.db)))
        html = response.body.decode("utf-8")

        self.assertIn("Lost Wage Portal", html)
        self.assertIn("CWA Local 3106", html)
        self.assertIn("Private access required", html)
        self.assertIn("Officers and admins", html)
        self.assertIn("Approved pay users", html)
        self.assertIn('href="/auth/login?next=/pay"', html)
        self.assertIn('href="/auth/steward/login?next=/pay"', html)
        self.assertIn('href="/pay/start" aria-current="page"', html)
        self.assertNotIn("entryForm", html)
        self.assertNotIn("/pay/api/context", html)

    async def test_pay_page_redirects_unauthenticated_users_to_public_start(self) -> None:
        response = await pay_page(_Request(state=SimpleNamespace(cfg=self._cfg(), db=self.db)))

        self.assertEqual(response.status_code, 303)
        self.assertEqual(response.headers["location"], "/pay/start")

    async def test_authenticated_pay_home_redirects_by_role(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True

        officer_response = await pay_page(
            _Request(
                state=SimpleNamespace(cfg=cfg, db=self.db),
                session={
                    "officer_user": {
                        "email": "officer@example.org",
                        "display_name": "Officer",
                        "role": "officer",
                        "contract_scopes": [],
                        "exp": int(time.time()) + 3600,
                    }
                },
            )
        )
        self.assertEqual(officer_response.headers["location"], "/pay/entry")

        admin_response = await pay_page(
            _Request(
                state=SimpleNamespace(cfg=cfg, db=self.db),
                session={
                    "officer_user": {
                        "email": "admin@example.org",
                        "display_name": "Admin",
                        "role": "admin",
                        "contract_scopes": [],
                        "exp": int(time.time()) + 3600,
                    }
                },
            )
        )
        self.assertEqual(admin_response.headers["location"], "/pay/treasurer")

    async def test_mileage_address_check_resolves_route_addresses(self) -> None:
        def fake_leg(*, api_key: str, origin: str, destination: str):  # noqa: ANN001
            self.assertEqual(api_key, "maps-key")
            return {
                "origin": f"Resolved {origin}",
                "destination": f"Resolved {destination}",
                "distance_text": "12.3 mi",
                "distance_miles": pay_portal.Decimal("12.30"),
            }

        with patch("grievance_api.services.pay_portal._google_leg", side_effect=fake_leg):
            result = pay_portal.validate_mileage_locations(
                google_maps_api_key="maps-key",
                locations=["union hall", "member house", "worksite"],
            )

        self.assertEqual(result["locations"], ["Resolved union hall", "Resolved member house", "Resolved worksite"])
        self.assertEqual(result["total_miles"], 24.6)
        self.assertEqual(len(result["legs"]), 2)

    async def test_pay_demo_context_is_isolated_from_live_period_records(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="officer@example.org"),
        )
        context = await pay_demo_context(request)
        pay_periods = await self.db.fetchall("SELECT id FROM pay_periods")

        self.assertEqual(context["period"]["id"], "demo")
        self.assertEqual(context["period"]["status"], "demo")
        self.assertEqual(context["entries"], [])
        self.assertEqual(context["demo_artifacts"], [])
        self.assertEqual(pay_periods, [])

    async def test_pay_demo_artifact_generation_writes_demo_only_files(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        Path(cfg.pay_portal.voucher_template_path).write_text("template", encoding="utf-8")
        request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="officer@example.org"),
        )
        with patch(
            "grievance_api.services.pay_portal.build_mileage_pdf",
            return_value=(b"%PDF-1.4\nmileage\n", pay_portal.Decimal("24.65"), pay_portal.Decimal("34.00")),
        ) as build_mileage, patch(
            "grievance_api.services.pay_portal.fill_pay_voucher_docx",
            side_effect=_fake_fill_docx,
        ) as fill_docx, patch(
            "grievance_api.services.pay_portal.watermark_pdf",
            side_effect=_fake_watermark,
        ) as watermark, patch(
            "grievance_api.services.pay_portal.merge_pdfs",
            side_effect=_fake_merge,
        ) as merge, patch(
            "grievance_api.web.routes_pay.docx_to_pdf",
            side_effect=_fake_docx_to_pdf,
        ), patch(
            "grievance_api.web.routes_pay._DEMO_JOB_EXECUTOR",
            _InlineExecutor(),
        ), patch(
            "grievance_api.services.pay_portal._prefetch_google_legs",
            return_value=None,
        ):
            result = await generate_pay_demo_output_files(
                PayDemoArtifactRequest(demo_step=4, demo_cycle_title="Officer Training"),
                request,
            )
        pay_periods = await self.db.fetchall("SELECT id FROM pay_periods")
        filenames = {row["filename"] for row in result["rows"]}
        artifact_dir = Path(self.data_root) / "pay" / "demo" / "officerexample.org"
        packet_text = (artifact_dir / "demo-payroll-packet.pdf").read_text(encoding="utf-8")

        self.assertEqual(filenames, {"demo-payroll-packet.pdf"})
        self.assertEqual(result["status"], "completed")
        self.assertIn("job_id", result)
        self.assertEqual(pay_periods, [])
        self.assertIn("demo-voucher-01", packet_text)
        self.assertIn("demo-voucher-10", packet_text)
        self.assertIn("demo-mileage-01", packet_text)
        self.assertIn("demo-mileage-10", packet_text)
        self.assertIn("WATERMARK:DEMO", packet_text)
        self.assertEqual(build_mileage.call_count, 10)
        self.assertEqual(fill_docx.call_count, 10)
        merge.assert_called_once()
        watermark.assert_called_once()
        filled_entries = [entry for call in fill_docx.call_args_list for entry in call.kwargs["entries"]]
        self.assertEqual(len({row["display_name"] for row in filled_entries}), 10)
        nick_entries = [row for row in filled_entries if row["display_name"] == "Nick Craig"]
        self.assertTrue(nick_entries)
        self.assertTrue(all(row["lost_wage_hourly_rate"] == 250.0 for row in nick_entries))
        self.assertGreater(len({row["lost_wage_hourly_rate"] for row in filled_entries}), 1)
        self.assertGreater(len({row["entry_date"] for row in filled_entries}), 3)
        president_diff_entries = [row for row in filled_entries if row.get("president_diff_hours")]
        self.assertTrue(president_diff_entries)
        self.assertTrue(all(float(row["president_diff_amount"]) > 0 for row in president_diff_entries))
        self.assertTrue(all(float(row["lost_wage_hourly_rate"]) == 62.0 for row in president_diff_entries))
        self.assertTrue(all(float(row["president_diff_rate"]) == 17.0 for row in president_diff_entries))
        self.assertEqual(len(filled_entries), 38)
        self.assertEqual(sorted(len(call.kwargs["entries"]) for call in fill_docx.call_args_list)[:7], [2] * 7)
        self.assertEqual(sorted(len(call.kwargs["entries"]) for call in fill_docx.call_args_list)[7:], [8] * 3)
        self.assertTrue(all(float(row["mileage_amount"]) > 0 for row in filled_entries))
        self.assertTrue(all(str(row["notes"]).startswith("DEMO TRAINING - ") for row in filled_entries))
        self.assertTrue(all(row.get("submitter_certified_at_utc") for row in filled_entries))
        self.assertTrue(all(row.get("submitter_certified_by") for row in filled_entries))
        self.assertTrue(all("signed off" in str(row.get("submitter_certification_text", "")).lower() for row in filled_entries))

    async def test_pay_demo_packet_entries_have_demo_narratives_and_varied_days(self) -> None:
        rng = random.Random("training-test")
        entries = pay_portal._demo_entries_for_packet(  # noqa: SLF001
            period_start="2026-05-03",
            demo_step=4,
            person={
                "display_name": "Demo Steward",
                "email": "demo.steward@cwa3106.local",
                "address": "11700 Phillips Hwy, Jacksonville, FL 32256, USA",
                "role": "steward",
            },
            day_offsets=[1, 5, 9],
            hourly_rate=pay_portal.Decimal("42.25"),
            mileage_miles=pay_portal.Decimal("34.00"),
            mileage_rate=pay_portal.Decimal("0.725"),
            mileage_amount=pay_portal.Decimal("24.65"),
            rng=rng,
        )

        self.assertEqual(len(entries), 3)
        self.assertEqual({row["entry_date"] for row in entries}, {"2026-05-04", "2026-05-08", "2026-05-12"})
        self.assertTrue(all(str(row["notes"]).startswith("DEMO TRAINING - ") for row in entries))
        self.assertTrue(all(row["lost_wage_hourly_rate"] == 42.25 for row in entries))
        self.assertTrue(all(float(row["mileage_amount"]) > 0 for row in entries))
        self.assertTrue(all(float(row["president_diff_hours"]) == 0 for row in entries))

        president_entries = pay_portal._demo_entries_for_packet(  # noqa: SLF001
            period_start="2026-05-03",
            demo_step=4,
            person={
                "display_name": "Demo President",
                "email": "demo.president@cwa3106.local",
                "address": "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
                "role": "president",
            },
            day_offsets=[1],
            hourly_rate=pay_portal.Decimal("54.75"),
            mileage_miles=pay_portal.Decimal("34.00"),
            mileage_rate=pay_portal.Decimal("0.725"),
            mileage_amount=pay_portal.Decimal("24.65"),
            rng=random.Random("president-demo"),
        )
        self.assertGreater(float(president_entries[0]["president_diff_hours"]), 0)
        self.assertEqual(float(president_entries[0]["lost_wage_hourly_rate"]), 62.0)
        self.assertEqual(float(president_entries[0]["president_diff_rate"]), 17.0)
        self.assertGreater(float(president_entries[0]["president_diff_amount"]), 0)

    async def test_pay_context_can_load_prior_period_entries(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        actor = PayActor("officer@example.org", "Officer Example", "officer", True, False, False)
        await self._save_profile(email=actor.email, display_name=actor.display_name)
        previous = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-03"))
        await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-17"))
        await upsert_entry(
            self.db,
            period_id=str(previous["id"]),
            actor=actor,
            data={"entry_date": "2026-05-16", "hours": 8, "notes": "prior period work"},
            pay_cfg=cfg.pay_portal,
        )
        request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="officer@example.org"),
        )
        request.query_params = {"period_id": str(previous["id"])}

        context = await pay_context(request)

        self.assertEqual(context["period"]["id"], previous["id"])
        self.assertEqual([row["entry_date"] for row in context["entries"]], ["2026-05-16"])
        self.assertGreaterEqual(len(context["periods"]), 2)
        period_rows = {row["id"]: row for row in context["periods"]}
        self.assertEqual(period_rows[str(previous["id"])] ["entry_count"], 1)

    async def test_submitter_can_delete_own_unlocked_entry(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        actor = PayActor("officer@example.org", "Officer Example", "officer", True, False, False)
        await self._save_profile(email=actor.email, display_name=actor.display_name)
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-17"))
        entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={"entry_date": str(period["period_start"]), "hours": 1, "notes": "wrong day"},
            pay_cfg=cfg.pay_portal,
        )
        await store_attachment(
            self.db,
            cfg=cfg,
            period_id=str(period["id"]),
            entry_id=str(entry["id"]),
            actor=actor,
            attachment_type="receipt",
            filename="receipt.pdf",
            content_type="application/pdf",
            content=b"%PDF-1.4\nreceipt\n",
            scan=False,
        )

        deleted = await delete_pay_entry(self.db, entry_id=str(entry["id"]), actor=actor)

        self.assertTrue(deleted["ok"])
        self.assertEqual(deleted["attachment_count"], 1)
        self.assertEqual(await self.db.fetchone("SELECT id FROM pay_entries WHERE id=?", (entry["id"],)), None)
        self.assertEqual(await self.db.fetchone("SELECT id FROM pay_attachments WHERE entry_id=?", (entry["id"],)), None)
        event = await self.db.fetchone("SELECT event_type, details_json FROM pay_events WHERE event_type='entry_deleted'")
        self.assertEqual(event[0], "entry_deleted")
        self.assertIn("wrong day", event[1])

    async def test_delete_entry_permission_and_lock_rules(self) -> None:
        cfg = self._cfg()
        actor = PayActor("officer@example.org", "Officer Example", "officer", True, False, False)
        other = PayActor("other@example.org", "Other", "officer", True, False, False)
        await self._save_profile(email=actor.email, display_name=actor.display_name)
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-17"))
        entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={"entry_date": str(period["period_start"]), "hours": 1},
            pay_cfg=cfg.pay_portal,
        )

        with self.assertRaises(PermissionError):
            await delete_pay_entry(self.db, entry_id=str(entry["id"]), actor=other)

        await self.db.exec("UPDATE pay_entries SET locked_at_utc=? WHERE id=?", (pay_portal.utcnow(), entry["id"]))
        with self.assertRaisesRegex(ValueError, "locked"):
            await delete_pay_entry(self.db, entry_id=str(entry["id"]), actor=actor)

    async def test_delete_entry_route_returns_for_submitter(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        actor = PayActor("officer@example.org", "Officer Example", "officer", True, False, False)
        await self._save_profile(email=actor.email, display_name=actor.display_name)
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-17"))
        entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={"entry_date": str(period["period_start"]), "hours": 1},
            pay_cfg=cfg.pay_portal,
        )
        request = _Request(state=SimpleNamespace(cfg=cfg, db=self.db), session=self._staff_session("officer", email=actor.email))

        deleted = await delete_pay_entry_route(str(entry["id"]), request)

        self.assertTrue(deleted["ok"])
        self.assertEqual(await self.db.fetchone("SELECT id FROM pay_entries WHERE id=?", (entry["id"],)), None)

    async def test_pay_profile_default_address_seeds_entries_and_context(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="officer@example.org"),
        )

        profile = await create_pay_profile(
            PayProfileUpsertRequest(
                principal_email="officer@example.org",
                principal_display_name="Officer Example",
                pay_basis="hourly",
                base_wage_input_type="hourly",
                base_wage_amount=42,
                default_address="123 Union Way, Jacksonville, FL",
            ),
            request,
        )
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-17"))
        entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=PayActor("officer@example.org", "Officer Example", "officer", True, False, False),
            data={"entry_date": str(period["period_start"]), "hours": 2, "address": ""},
            pay_cfg=cfg.pay_portal,
        )

        self.assertEqual(profile["default_address"], "123 Union Way, Jacksonville, FL")
        self.assertEqual(entry["address"], "123 Union Way, Jacksonville, FL")
        context = await pay_context(request)
        self.assertEqual(context["pay_profile"]["default_address"], "123 Union Way, Jacksonville, FL")

    async def test_first_entry_address_seeds_empty_profile_default(self) -> None:
        cfg = self._cfg()
        actor = PayActor("officer@example.org", "Officer Example", "officer", True, False, False)
        await self._save_profile(email=actor.email, display_name=actor.display_name)
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-17"))

        await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={"entry_date": str(period["period_start"]), "hours": 1, "address": "456 New Address"},
            pay_cfg=cfg.pay_portal,
        )

        profile = await pay_profile_by_email(self.db, email=actor.email)
        self.assertEqual(profile["default_address"], "456 New Address")

    async def test_treasurer_review_and_submitter_edit_resets_status(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        period_id, entry_id = await self._open_period_and_entry()
        treasurer_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("admin"),
        )
        officer_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="officer@example.org"),
        )

        reviewed = await review_pay_entry_route(
            entry_id,
            PayEntryReviewRequest(review_status="needs_fix", review_note="Add missing receipt"),
            treasurer_request,
        )
        self.assertEqual(reviewed["review_status"], "needs_fix")
        self.assertEqual(reviewed["review_note"], "Add missing receipt")
        with self.assertRaises(HTTPException) as denied:
            await review_pay_entry_route(entry_id, PayEntryReviewRequest(review_status="approved"), officer_request)
        self.assertEqual(denied.exception.status_code, 403)

        await upsert_entry(
            self.db,
            period_id=period_id,
            actor=PayActor("officer@example.org", "Officer Example", "officer", True, False, False),
            data={"entry_date": "2026-05-17", "hours": 3, "notes": "fixed"},
            pay_cfg=cfg.pay_portal,
        )
        row = await self.db.fetchone("SELECT review_status, review_note FROM pay_entries WHERE id=?", (entry_id,))
        self.assertEqual(row[0], "pending")
        self.assertEqual(row[1], "Add missing receipt")

    async def test_treasurer_positive_corrections_add_without_reducing(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        period_id, entry_id = await self._open_period_and_entry()
        request = _Request(state=SimpleNamespace(cfg=cfg, db=self.db), session=self._staff_session("admin"))

        result = await create_pay_entry_correction_route(
            PayEntryCorrectionRequest(
                period_id=period_id,
                user_email="officer@example.org",
                display_name="Officer Example",
                entry_date="2026-05-17",
                hours=1.5,
                mileage_miles=10,
                mileage_rate=0.725,
                mileage_amount=7.25,
                meals_amount=5,
                notes="Forgot afternoon visit",
            ),
            request,
        )

        entry = result["entry"]
        self.assertEqual(entry["hours"], 9.5)
        self.assertEqual(entry["mileage_miles"], 10.0)
        self.assertEqual(entry["mileage_rate"], 0.725)
        self.assertEqual(entry["mileage_amount"], 7.25)
        self.assertEqual(entry["meals_amount"], 5.0)
        self.assertEqual(entry["correction_count"], 1)
        with self.assertRaises(HTTPException) as denied:
            await create_pay_entry_correction_route(
                PayEntryCorrectionRequest(period_id=period_id, user_email="officer@example.org", entry_date="2026-05-17", hours=-1),
                request,
            )
        self.assertEqual(denied.exception.status_code, 400)

    async def test_treasurer_correction_creates_missing_day_entry(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        await self._save_profile(email="officer@example.org", display_name="Officer Example")
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-17"))
        request = _Request(state=SimpleNamespace(cfg=cfg, db=self.db), session=self._staff_session("admin"))

        result = await create_pay_entry_correction_route(
            PayEntryCorrectionRequest(period_id=str(period["id"]), user_email="officer@example.org", entry_date="2026-05-18", hours=2, notes="Missing day"),
            request,
        )

        self.assertEqual(result["entry"]["entry_date"], "2026-05-18")
        self.assertEqual(result["entry"]["hours"], 2.0)
        self.assertEqual(result["entry"]["correction_count"], 1)

    async def test_packet_excludes_rejected_and_needs_fix_entries(self) -> None:
        cfg = self._cfg()
        cfg.pay_portal.president_email = "president@example.org"
        Path(cfg.pay_portal.voucher_template_path).write_text("template", encoding="utf-8")
        actor = PayActor("treasurer@example.org", "Treasurer", "treasurer", True, True, True)
        officer = PayActor("officer@example.org", "Officer Example", "officer", True, False, False)
        await self._save_profile(email=officer.email, display_name=officer.display_name)
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-17"))
        included = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=officer,
            data={"entry_date": "2026-05-17", "hours": 1, "notes": "include"},
            pay_cfg=cfg.pay_portal,
        )
        rejected = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=officer,
            data={"entry_date": "2026-05-18", "hours": 1, "notes": "reject"},
            pay_cfg=cfg.pay_portal,
        )
        await pay_portal.review_pay_entry(
            self.db,
            entry_id=str(rejected["id"]),
            actor=actor,
            review_status="rejected",
            review_note="Duplicate",
        )
        filled_dates: list[str] = []

        def _capture_fill(**kwargs):  # noqa: ANN003
            filled_dates.extend(str(entry["entry_date"]) for entry in kwargs["entries"])
            _fake_fill_docx(**kwargs)

        with patch("grievance_api.services.pay_portal.fill_pay_voucher_docx", side_effect=_capture_fill), patch(
            "grievance_api.services.pay_portal.merge_pdfs", side_effect=_fake_merge
        ), patch("grievance_api.web.routes_pay.docx_to_pdf", side_effect=_fake_docx_to_pdf):
            result = await lock_period_and_send_packet(
                db=self.db,
                cfg=cfg,
                graph=_FakeGraph(),
                docuseal=_FakeDocuSeal(),
                period_id=str(period["id"]),
                actor=actor,
                president_signer_email="president@example.org",
                docx_to_pdf_func=_fake_docx_to_pdf,
            )

        self.assertIn("2026-05-17", filled_dates)
        self.assertNotIn("2026-05-18", filled_dates)
        self.assertEqual(result["excluded_entries"][0]["entry_date"], "2026-05-18")

    async def test_pay_workspace_pages_are_role_focused(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        officer_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session={
                "officer_user": {
                    "email": "officer@example.org",
                    "display_name": "Officer",
                    "role": "officer",
                    "contract_scopes": [],
                    "exp": int(time.time()) + 3600,
                }
            },
        )

        entry_html = (await pay_view_page("entry", officer_request)).body.decode("utf-8")
        self.assertIn("My Pay Profile", entry_html)
        self.assertIn('id="myPayProfileForm"', entry_html)
        self.assertIn('id="entryForm"', entry_html)
        self.assertIn('id="periodSelect"', entry_html)
        self.assertIn('renderPeriodSelect', entry_html)
        self.assertIn('data-review="false"', entry_html)
        self.assertNotIn('<th>Review</th>', entry_html)
        self.assertIn("canEditEntryForm ? openEntryForEdit(row) : openCorrectionForEntry(row)", entry_html)
        self.assertIn("deleteEntry(row)", entry_html)
        self.assertIn("/pay/api/entries/${encodeURIComponent(row.id)}", entry_html)
        self.assertIn("Sign Off and Save Entry", entry_html)
        self.assertIn('name="submitter_certified"', entry_html)
        self.assertIn("Sign off and save to update this voucher", entry_html)
        self.assertNotIn('name="display_name"', entry_html)
        self.assertNotIn("Lost Wage Type", entry_html)
        self.assertNotIn("Lost Wage Amount", entry_html)
        self.assertNotIn("Mileage Miles<input", entry_html)
        self.assertNotIn("Mileage Amount<input", entry_html)
        self.assertNotIn("IRS Rate<input", entry_html)
        self.assertLess(entry_html.index("Daily Input"), entry_html.index("My Pay Profile"))
        self.assertLess(entry_html.index("My Pay Profile"), entry_html.index("Commission Pay Proof"))
        self.assertIn('id="commissionProofPanel"', entry_html)
        self.assertIn("Last Month Payroll", entry_html)
        self.assertIn('name="payroll_month"', entry_html)
        self.assertIn("syncPayProfileVisibility", entry_html)
        self.assertNotIn("Pay Stub Proof", entry_html)
        self.assertIn("President Differential", entry_html)
        self.assertIn('id="presidentDifferentialPanel"', entry_html)
        self.assertIn("disabled-panel", entry_html)
        self.assertIn("Daily Tally", entry_html)
        self.assertIn('id="dailyTallyBody"', entry_html)
        self.assertIn("<th>President Diff</th><th>Total</th>", entry_html)
        self.assertIn("tally.president_diff += Number(row.president_diff_amount || 0)", entry_html)
        self.assertIn("tally.lost + tally.mileage + tally.expenses + tally.president_diff", entry_html)
        self.assertIn('href="/pay/mileage"', entry_html)
        self.assertIn('href="/pay/demo"', entry_html)
        self.assertNotIn('href="/pay/president"', entry_html)
        self.assertIn('href="/pay/treasurer"', entry_html)
        self.assertNotIn("Pay User Allowlist", entry_html)

        president_redirect = await pay_view_page("president", officer_request)
        self.assertEqual(president_redirect.status_code, 303)
        self.assertEqual(president_redirect.headers["location"], "/pay/entry")

        president_cfg = self._cfg()
        president_cfg.officer_auth.enabled = True
        president_cfg.pay_portal.president_email = "president@example.org"
        president_request = _Request(
            state=SimpleNamespace(cfg=president_cfg, db=self.db),
            session={
                "officer_user": {
                    "email": "president@example.org",
                    "display_name": "President",
                    "role": "officer",
                    "contract_scopes": [],
                    "exp": int(time.time()) + 3600,
                }
            },
        )
        president_html = (await pay_view_page("entry", president_request)).body.decode("utf-8")
        president_panel = president_html.split('id="presidentDifferentialPanel"', 1)[1].split("</section>", 1)[0]
        self.assertNotIn("disabled", president_panel)

        mileage_html = (await pay_view_page("mileage", officer_request)).body.decode("utf-8")
        self.assertIn("Mileage Tracker", mileage_html)
        self.assertIn('id="mileageForm"', mileage_html)
        self.assertNotIn('label for="name"', mileage_html)
        self.assertIn('type="hidden" id="name" name="name"', mileage_html)
        self.assertIn('id="commonPlaceSelect"', mileage_html)
        self.assertIn('id="irsRateDisplay"', mileage_html)
        self.assertNotIn('id="irs_rate"', mileage_html)
        self.assertIn('id="locations"', mileage_html)
        self.assertIn("Add Location", mileage_html)
        self.assertIn("Check Addresses", mileage_html)
        self.assertIn("/pay/api/mileage/check-addresses", mileage_html)
        self.assertIn("Addresses updated from Google route results", mileage_html)
        self.assertIn("Generate Mileage PDF", mileage_html)
        self.assertIn("Mileage report breakdown", mileage_html)
        self.assertIn("IRS Standard Mileage Rate", mileage_html)
        self.assertIn("Total Distance", mileage_html)
        self.assertIn("Total Reimbursement", mileage_html)
        self.assertIn("Download PDF", mileage_html)
        self.assertIn("Remove Report", mileage_html)
        self.assertIn("removeMileageReport", mileage_html)
        self.assertIn('/pay/api/attachments/${encodeURIComponent(row.id)}/download', mileage_html)
        self.assertIn("rate: null", mileage_html)
        self.assertIn("/pay/api/entries/${entryId}/mileage", mileage_html)
        self.assertNotIn("Pay User Allowlist", mileage_html)

        demo_html = (await pay_view_page("demo", officer_request)).body.decode("utf-8")
        self.assertIn("Training Demo Cycle", demo_html)
        self.assertIn('id="demoCyclePanel"', demo_html)
        self.assertIn('id="demoChecklist"', demo_html)
        self.assertIn('id="demoActivityLog"', demo_html)
        self.assertIn('id="demoFilesBody"', demo_html)
        self.assertIn('id="generateDemoFilesBtn"', demo_html)
        self.assertIn('id="demoFeedbackForm"', demo_html)
        self.assertIn("/pay/api/demo/context", demo_html)
        self.assertIn("/pay/api/demo/artifacts", demo_html)
        self.assertIn("/pay/api/demo/feedback", demo_html)
        self.assertIn("Signed off in Pay Portal", demo_html)
        self.assertIn("president-only packet signing", demo_html)
        self.assertIn("Confirm only the president signs the packet", demo_html)
        self.assertNotIn("Lock And Send For Signature", demo_html)

        treasurer_html = (await pay_view_page("treasurer", officer_request)).body.decode("utf-8")
        self.assertIn("Review, Lock, and Send", treasurer_html)
        self.assertIn("Voucher Packet", treasurer_html)
        self.assertIn("Read-only review access", treasurer_html)
        self.assertNotIn("Lock And Send For Signature", treasurer_html)
        self.assertNotIn('id="correctionForm"', treasurer_html)
        self.assertNotIn("First-Time Wage Setup", treasurer_html)
        self.assertNotIn("Pay User Allowlist", treasurer_html)

        with self.assertRaises(HTTPException) as denied:
            await pay_view_page("admin", officer_request)
        self.assertEqual(denied.exception.status_code, 403)

        admin_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session={
                "officer_user": {
                    "email": "admin@example.org",
                    "display_name": "Admin",
                    "role": "admin",
                    "contract_scopes": [],
                    "exp": int(time.time()) + 3600,
                }
            },
        )
        treasurer_admin_html = (await pay_view_page("treasurer", admin_request)).body.decode("utf-8")
        self.assertIn('data-review="true"', treasurer_admin_html)
        self.assertIn('<th>Review</th>', treasurer_admin_html)
        self.assertIn('id="correctionPanel"', treasurer_admin_html)
        self.assertIn('class="panel hidden" id="correctionPanel"', treasurer_admin_html)
        self.assertIn("Choose Edit Voucher on a row below", treasurer_admin_html)
        self.assertIn("Edit Voucher", treasurer_admin_html)
        self.assertIn("openCorrectionForEntry(row)", treasurer_admin_html)

        admin_html = (await pay_view_page("admin", admin_request)).body.decode("utf-8")
        self.assertIn("Pay User Allowlist", admin_html)
        self.assertIn("Microsoft People", admin_html)
        self.assertIn("Import Microsoft Paid Users", admin_html)
        self.assertIn("/pay/api/internal-users/import", admin_html)
        self.assertIn("People and Pay Profiles", admin_html)
        self.assertIn("Save Pay Profile", admin_html)
        self.assertIn("Edit Profile", admin_html)
        self.assertIn("Remove Person", admin_html)
        self.assertIn("removePayProfile", admin_html)
        self.assertIn("mergedPayProfileRows", admin_html)
        self.assertIn("Demo Mode", admin_html)
        self.assertIn('id="demoSettingsForm"', admin_html)
        self.assertIn("/pay/api/demo/settings", admin_html)
        self.assertIn("function latestSavedWageScale()", admin_html)
        self.assertIn("form.target_weekly_amount.value", admin_html)
        self.assertIn("wageScale.effective_date", admin_html)
        self.assertIn("IRS Rate Updates", admin_html)
        self.assertIn('id="commonPlaceForm"', admin_html)
        self.assertIn('id="commonPlacesBody"', admin_html)
        self.assertNotIn('id="entryForm"', admin_html)

    async def test_pay_admin_can_search_internal_microsoft_users(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        graph = _DirectoryGraphStub(
            [
                DirectoryUserRef(
                    id="oid-president-1",
                    display_name="Pat President",
                    email="president@cwa3106.com",
                    user_principal_name="president@cwa3106.com",
                )
            ]
        )
        request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db, graph=graph),
            session=self._staff_session("admin"),
        )

        response = await pay_directory_users(request, search="pres", limit=8)

        self.assertEqual(response.count, 1)
        self.assertEqual(response.rows[0].principal_id, "oid-president-1")
        self.assertEqual(response.rows[0].email, "president@cwa3106.com")
        self.assertEqual(graph.last_search, "pres")
        self.assertEqual(graph.last_limit, 8)

    def test_graph_licensed_directory_users_filters_free_and_trial_skus(self) -> None:
        graph = _LicensedGraphUploader()

        rows = graph.list_licensed_directory_users(limit=50)

        self.assertEqual([row.email for row in rows], ["paid@example.org", "both@example.org"])
        subscribed_call = [call for call in graph.calls if call[1] == "/subscribedSkus"][0]
        self.assertNotIn("$top", subscribed_call[2] or {})
        self.assertIn(("GET", "/users"), [(call[0], call[1]) for call in graph.calls])

    async def test_pay_admin_can_import_licensed_internal_users_as_profiles(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        graph = _DirectoryGraphStub(
            [],
            licensed_rows=[
                DirectoryUserRef(
                    id="oid-paid-1",
                    display_name="Paid User",
                    email="paid.user@cwa3106.com",
                    user_principal_name="paid.user@cwa3106.com",
                ),
                DirectoryUserRef(
                    id="oid-existing-1",
                    display_name="Existing User",
                    email="existing@cwa3106.com",
                    user_principal_name="existing@cwa3106.com",
                ),
            ],
        )
        request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db, graph=graph),
            session=self._staff_session("admin"),
        )
        await upsert_pay_profile(
            self.db,
            principal_id="oid-existing-1",
            principal_email="existing@cwa3106.com",
            principal_display_name="Existing User",
            pay_basis="hourly",
            base_wage_input_type="hourly",
            base_wage_amount=44,
            weekly_basis_hours=40,
            commission_month_1_amount=0,
            commission_month_2_amount=0,
            commission_month_3_amount=0,
            status="active",
            notes=None,
            updated_by="admin@example.org",
        )

        result = await import_pay_internal_users(PayInternalUserImportRequest(limit=50), request)

        self.assertEqual(graph.last_import_limit, 50)
        self.assertEqual(result["imported_count"], 1)
        self.assertEqual(result["skipped_count"], 1)
        profile = await pay_profile_by_email(self.db, email="paid.user@cwa3106.com")
        self.assertIsNotNone(profile)
        self.assertEqual(profile["pay_basis"], "expense_only")
        self.assertEqual(profile["status"], "active")
        self.assertIn("Auto-imported", str(profile["notes"]))
        existing = await pay_profile_by_email(self.db, email="existing@cwa3106.com")
        self.assertEqual(existing["pay_basis"], "hourly")
        self.assertEqual(existing["base_wage_amount"], 44.0)

    async def test_pay_admin_can_assign_president_and_treasurer_from_internal_accounts(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        admin_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("admin"),
        )

        president = await save_pay_internal_role(
            InternalRoleAssignmentRequest(
                principal_id="oid-president-1",
                principal_email="president@cwa3106.com",
                principal_display_name="Pat President",
                role="president",
            ),
            admin_request,
        )
        treasurer = await save_pay_internal_role(
            InternalRoleAssignmentRequest(
                principal_id="oid-treasurer-1",
                principal_email="treasurer@cwa3106.com",
                principal_display_name="Terry Treasurer",
                role="treasurer",
            ),
            admin_request,
        )

        settings = await pay_settings(self.db, pay_cfg=cfg.pay_portal)
        self.assertEqual(settings["president_email"], "president@cwa3106.com")
        context = await pay_context(admin_request)
        self.assertEqual(
            {(row["principal_email"], row["role"]) for row in context["internal_roles"]},
            {("president@cwa3106.com", "president"), ("treasurer@cwa3106.com", "treasurer")},
        )
        profile = await create_pay_profile(
            PayProfileUpsertRequest(
                principal_id="oid-president-1",
                principal_email="president@cwa3106.com",
                principal_display_name="Pat President",
                pay_basis="president",
                base_wage_input_type="hourly",
                base_wage_amount=45,
                commission_month_1_amount=999,
                commission_month_2_amount=999,
                commission_month_3_amount=999,
            ),
            admin_request,
        )
        self.assertEqual(profile["pay_basis"], "president")
        self.assertEqual(profile["commission_hourly_rate"], 0.0)
        context_with_profile = await pay_context(admin_request)
        self.assertIn("pay_profiles", context_with_profile)
        self.assertEqual(context_with_profile["pay_profiles"][0]["principal_email"], "president@cwa3106.com")

        treasurer_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="treasurer@cwa3106.com"),
        )
        redirect = await pay_page(treasurer_request)
        self.assertEqual(redirect.headers["location"], "/pay/treasurer")

        deleted = await remove_pay_internal_role(int(president["assignment_id"]), admin_request)
        self.assertEqual(deleted["role"], "president")
        settings_after_delete = await pay_settings(self.db, pay_cfg=cfg.pay_portal)
        self.assertEqual(settings_after_delete["president_email"], "")
        self.assertEqual(treasurer["role"], "treasurer")

    async def test_treasurer_can_remove_pay_profile_from_people_roster(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        admin_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("admin"),
        )
        officer_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="officer@example.org"),
        )
        await upsert_pay_profile(
            self.db,
            principal_id="oid-remove-1",
            principal_email="remove.me@example.org",
            principal_display_name="Remove Me",
            pay_basis="expense_only",
            base_wage_input_type="hourly",
            base_wage_amount=0,
            weekly_basis_hours=40,
            commission_month_1_amount=0,
            commission_month_2_amount=0,
            commission_month_3_amount=0,
            status="active",
            notes=None,
            updated_by="admin@example.org",
        )

        removed = await remove_pay_profile("remove.me@example.org", admin_request)

        self.assertEqual(removed["removed"]["principal_email"], "remove.me@example.org")
        self.assertIsNone(await pay_profile_by_email(self.db, email="remove.me@example.org"))
        with self.assertRaises(HTTPException) as denied:
            await remove_pay_profile("remove.me@example.org", officer_request)
        self.assertEqual(denied.exception.status_code, 403)

    async def test_pay_profile_update_preserves_base_wage_input_type(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        admin_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("admin", email="admin@example.org"),
        )

        hourly = await create_pay_profile(
            PayProfileUpsertRequest(
                principal_email="member@example.org",
                principal_display_name="Member Example",
                pay_basis="hourly",
                base_wage_input_type="weekly",
                base_wage_amount=2000,
                weekly_basis_hours=40,
            ),
            admin_request,
        )
        weekly = await create_pay_profile(
            PayProfileUpsertRequest(
                principal_email="member@example.org",
                principal_display_name="Member Example",
                pay_basis="weekly",
                base_wage_input_type="hourly",
                base_wage_amount=55,
                weekly_basis_hours=40,
            ),
            admin_request,
        )
        hourly_again = await create_pay_profile(
            PayProfileUpsertRequest(
                principal_email="member@example.org",
                principal_display_name="Member Example",
                pay_basis="hourly",
                base_wage_input_type="weekly",
                base_wage_amount=55,
                weekly_basis_hours=40,
            ),
            admin_request,
        )

        self.assertEqual(hourly["pay_basis"], "hourly")
        self.assertEqual(hourly["base_wage_input_type"], "weekly")
        self.assertEqual(hourly["base_wage_amount"], 2000.0)
        self.assertEqual(hourly["calculated_hourly_rate"], 50.0)
        self.assertEqual(weekly["pay_basis"], "weekly")
        self.assertEqual(weekly["base_wage_input_type"], "hourly")
        self.assertEqual(weekly["base_wage_amount"], 55.0)
        self.assertEqual(weekly["calculated_hourly_rate"], 55.0)
        self.assertEqual(hourly_again["pay_basis"], "hourly")
        self.assertEqual(hourly_again["base_wage_input_type"], "weekly")
        self.assertEqual(hourly_again["base_wage_amount"], 55.0)
        self.assertEqual(hourly_again["calculated_hourly_rate"], 1.38)

    async def test_pay_profile_self_service_permissions(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        officer_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="officer@example.org"),
        )

        own_profile = await create_pay_profile(
            PayProfileUpsertRequest(
                principal_email="officer@example.org",
                principal_display_name="Officer Example",
                pay_basis="hourly",
                base_wage_input_type="hourly",
                base_wage_amount=42,
                weekly_basis_hours=40,
                notes="self entered",
            ),
            officer_request,
        )
        self.assertEqual(own_profile["principal_email"], "officer@example.org")
        self.assertEqual(own_profile["pay_basis"], "hourly")
        self.assertEqual(own_profile["status"], "active")
        self.assertEqual(own_profile["updated_by"], "officer@example.org")

        cfg.external_steward_auth.enabled = True
        now = pay_portal.utcnow()
        external_id = await self.db.insert(
            """
            INSERT INTO external_steward_users(
              email, display_name, status, auth_source, auth_issuer, auth_subject, invited_by,
              created_at_utc, updated_at_utc
            ) VALUES(?,?,?,?,?,?,?,?,?)
            """,
            (
                "guest@example.org",
                "Guest Example",
                "active",
                "external_microsoft",
                "issuer",
                "subject-guest",
                "admin@example.org",
                now,
                now,
            ),
        )
        await pay_portal.upsert_pay_user(
            self.db,
            email="guest@example.org",
            display_name="Guest Example",
            role="guest",
            status="active",
            actor="admin@example.org",
        )
        external_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session={
                "external_steward_user": {
                    "external_user_id": external_id,
                    "email": "guest@example.org",
                    "display_name": "Guest Example",
                    "role": "external_steward",
                    "issuer": "issuer",
                    "provider_subject": "subject-guest",
                    "verified_email": True,
                    "exp": int(time.time()) + 3600,
                }
            },
        )
        external_profile = await create_pay_profile(
            PayProfileUpsertRequest(
                principal_email="guest@example.org",
                principal_display_name="Guest Example",
                pay_basis="commission",
                base_wage_input_type="hourly",
                base_wage_amount=30,
                weekly_basis_hours=40,
                commission_month_1_amount=100,
                commission_month_2_amount=200,
                commission_month_3_amount=300,
            ),
            external_request,
        )
        self.assertEqual(external_profile["principal_email"], "guest@example.org")
        self.assertEqual(external_profile["pay_basis"], "commission")
        self.assertEqual(external_profile["updated_by"], "guest@example.org")

        with self.assertRaises(HTTPException) as other_denied:
            await create_pay_profile(
                PayProfileUpsertRequest(
                    principal_email="other@example.org",
                    principal_display_name="Other Person",
                    pay_basis="hourly",
                    base_wage_input_type="hourly",
                    base_wage_amount=44,
                ),
                officer_request,
            )
        self.assertEqual(other_denied.exception.status_code, 403)

        with self.assertRaises(HTTPException) as president_denied:
            await create_pay_profile(
                PayProfileUpsertRequest(
                    principal_email="officer@example.org",
                    principal_display_name="Officer Example",
                    pay_basis="president",
                    base_wage_input_type="hourly",
                    base_wage_amount=45,
                ),
                officer_request,
            )
        self.assertEqual(president_denied.exception.status_code, 403)

        admin_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("admin", email="admin@example.org"),
        )
        president_profile = await create_pay_profile(
            PayProfileUpsertRequest(
                principal_email="president@example.org",
                principal_display_name="President",
                pay_basis="president",
                base_wage_input_type="hourly",
                base_wage_amount=45,
                weekly_basis_hours=40,
            ),
            admin_request,
        )
        self.assertEqual(president_profile["pay_basis"], "president")
        self.assertEqual(president_profile["updated_by"], "admin@example.org")

        cfg.pay_portal.president_email = "president@example.org"
        president_request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session=self._staff_session("officer", email="president@example.org"),
        )
        updated_president_profile = await create_pay_profile(
            PayProfileUpsertRequest(
                principal_email="president@example.org",
                principal_display_name="President",
                pay_basis="hourly",
                base_wage_input_type="hourly",
                base_wage_amount=47,
                weekly_basis_hours=40,
            ),
            president_request,
        )
        self.assertEqual(updated_president_profile["pay_basis"], "president")
        self.assertEqual(updated_president_profile["base_wage_amount"], 47.0)
        self.assertEqual(updated_president_profile["updated_by"], "president@example.org")

    async def test_irs_rate_parser_detects_official_business_rate(self) -> None:
        candidates = parse_irs_mileage_rate_candidates(
            content="""
              <title>IRS sets 2027 business standard mileage rate at 75.5 cents per mile</title>
              <p>Beginning Jan. 1, 2027, the standard mileage rates are:
              75.5 cents per mile driven for business use.</p>
            """,
            source_url="https://www.irs.gov/newsroom/example",
        )

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].rate_year, "2027")
        self.assertEqual(candidates[0].effective_date, "2027-01-01")
        self.assertEqual(candidates[0].rate_per_mile, pay_portal.Decimal("0.755"))

    def test_common_places_config_parses_sharepoint_file_shapes(self) -> None:
        json_rows = pay_portal.parse_common_places_config(
            filename="places.json",
            content='{"places":[{"name":"Union Hall","address":"4076 Union Hall Pl, Jacksonville, FL"}]}',
        )
        csv_rows = pay_portal.parse_common_places_config(
            filename="addresses.csv",
            content='label,address\nPhillips Hwy,"11700 Phillips Hwy, Jacksonville, FL"\n',
        )
        text_rows = pay_portal.parse_common_places_config(
            filename="legacy.txt",
            content="Penman Rd=98 Penman Rd S, Jacksonville Beach, FL\n",
        )

        merged = pay_portal.merge_common_places(json_rows, csv_rows, text_rows)
        self.assertEqual(
            merged,
            [
                {"label": "Union Hall", "address": "4076 Union Hall Pl, Jacksonville, FL"},
                {"label": "Phillips Hwy", "address": "11700 Phillips Hwy, Jacksonville, FL"},
                {"label": "Penman Rd", "address": "98 Penman Rd S, Jacksonville Beach, FL"},
            ],
        )

    def test_common_places_cache_writes_app_side_file(self) -> None:
        path = pay_portal.write_common_places_cache(
            data_root=self.data_root,
            places=[
                {"label": "Union Hall", "address": "4076 Union Hall Pl, Jacksonville, FL"},
                {"label": "Union Hall", "address": "4076 Union Hall Pl, Jacksonville, FL"},
            ],
        )

        self.assertEqual(path, Path(self.data_root) / "pay" / "common_places.json")
        self.assertEqual(
            pay_portal.load_common_places_cache(data_root=self.data_root),
            [{"label": "Union Hall", "address": "4076 Union Hall Pl, Jacksonville, FL"}],
        )

    def test_sharepoint_common_places_loader_reads_config_folder(self) -> None:
        class _SharePointConfigGraph:
            def __init__(self) -> None:
                self.list_call: dict[str, object] | None = None

            def list_files_in_folder_path(self, **kwargs):  # noqa: ANN003
                self.list_call = dict(kwargs)
                return [
                    SimpleNamespace(drive_id="drive", item_id="places", name="places.csv"),
                    SimpleNamespace(drive_id="drive", item_id="ignored", name="notes.pdf"),
                ]

            def download_item_bytes(self, *, drive_id: str, item_id: str) -> bytes:
                if (drive_id, item_id) == ("drive", "places"):
                    return b'label,address\nUnion Hall,"4076 Union Hall Pl, Jacksonville, FL"\n'
                return b""

        graph = _SharePointConfigGraph()
        rows = pay_portal.load_sharepoint_common_places(
            graph=graph,
            graph_cfg=SimpleNamespace(
                site_hostname="contoso.sharepoint.com",
                site_path="/sites/cwa3106",
                document_library="Officer Eboard committee - Documents",
            ),
            pay_cfg=SimpleNamespace(common_places_sharepoint_folder=r"Local3106\Mileage\Config"),
        )

        self.assertEqual(graph.list_call["folder_path"], "Local3106/Mileage/Config")
        self.assertEqual(graph.list_call["library"], "Officer Eboard committee - Documents")
        self.assertEqual(rows, [{"label": "Union Hall", "address": "4076 Union Hall Pl, Jacksonville, FL"}])

    def test_sharepoint_common_places_loader_accepts_full_synced_path(self) -> None:
        class _SharePointConfigGraph:
            def __init__(self) -> None:
                self.calls: list[dict[str, object]] = []

            def list_files_in_folder_path(self, **kwargs):  # noqa: ANN003
                self.calls.append(dict(kwargs))
                if kwargs["library"] == "Documents":
                    raise RuntimeError("wrong library")
                return [SimpleNamespace(drive_id="drive", item_id="nested", name="nested/places.csv")]

            def download_item_bytes(self, *, drive_id: str, item_id: str) -> bytes:
                if (drive_id, item_id) == ("drive", "nested"):
                    return b'label,address\nNested,"123 Nested St, Jacksonville, FL"\n'
                return b""

        graph = _SharePointConfigGraph()
        rows = pay_portal.load_sharepoint_common_places(
            graph=graph,
            graph_cfg=SimpleNamespace(
                site_hostname="cwa3106.sharepoint.com",
                site_path="/sites/GrievancesLibrary",
                document_library="Documents",
            ),
            pay_cfg=SimpleNamespace(
                common_places_sharepoint_library="",
                common_places_sharepoint_folder=(
                    r"CWA 3106\Officer Eboard committee - Documents\Local3106\Mileage\Config"
                ),
            ),
        )

        self.assertEqual(graph.calls[0]["library"], "Officer Eboard committee - Documents")
        self.assertEqual(graph.calls[0]["folder_path"], "Local3106/Mileage/Config")
        self.assertTrue(graph.calls[0]["recursive"])
        self.assertEqual(rows, [{"label": "Nested", "address": "123 Nested St, Jacksonville, FL"}])

    def test_daily_pay_narrative_groups_notes_by_date(self) -> None:
        lines = pay_portal.pay_entry_daily_narrative_lines(
            [
                {"entry_date": "2026-05-11", "notes": "Prepared grievance packet"},
                {"entry_date": "2026-05-10", "notes": "Met with member\nand steward"},
                {"entry_date": "2026-05-10", "notes": "Reviewed payroll records"},
                {"entry_date": "2026-05-12", "notes": ""},
            ]
        )

        self.assertEqual(
            lines,
            [
                "Sun 05/10/2026: Met with member and steward; Reviewed payroll records",
                "Mon 05/11/2026: Prepared grievance packet",
            ],
        )

    def test_voucher_entry_amounts_use_profile_lost_wage_rate(self) -> None:
        amounts = pay_portal._entry_amounts(  # noqa: SLF001
            {
                "hours": 4,
                "hourly_rate": 1,
                "lost_wage_hourly_rate": 38.5,
                "mileage_amount": 12.34,
            }
        )

        self.assertEqual(amounts["hours"], pay_portal.Decimal("154.00"))
        self.assertEqual(amounts["mileage"], pay_portal.Decimal("12.34"))

    def test_voucher_weekly_quantities_use_raw_activity_values(self) -> None:
        quantities = pay_portal._entry_weekly_quantities(  # noqa: SLF001
            {
                "hours": 2.5,
                "hourly_rate": 52.25,
                "lost_wage_hourly_rate": 52.25,
                "mileage_miles": 12.25,
                "mileage_amount": 8.88,
                "president_diff_hours": 1.0,
                "president_diff_amount": 17.00,
            }
        )

        self.assertEqual(quantities["hours"], pay_portal.Decimal("2.50"))
        self.assertEqual(quantities["mileage"], pay_portal.Decimal("12.25"))
        self.assertEqual(quantities["president_diff"], pay_portal.Decimal("1.00"))
        self.assertEqual(pay_portal._voucher_quantity_text(quantities["hours"]), "2.50")  # noqa: SLF001

    def test_voucher_week_tables_show_quantities_and_bottom_table_shows_cash(self) -> None:
        template_path = Path(self.tmpdir.name) / "voucher-template.docx"
        output_path = Path(self.tmpdir.name) / "voucher-filled.docx"
        doc = Document()
        for text in ("Local #", "Name", "Address", "Hourly Rate", "Signature"):
            doc.add_paragraph(text)

        def add_week_table() -> None:
            table = doc.add_table(rows=10, cols=10)
            labels = [
                "",
                "Hours",
                "Mileage",
                "Rentals",
                "Meals",
                "Hotel Room",
                "Miscellaneous",
                "President Diff",
                "Total",
                "",
            ]
            for index, label in enumerate(labels):
                table.rows[index].cells[0].text = label

        add_week_table()
        add_week_table()
        totals_table = doc.add_table(rows=2, cols=8)
        totals_table.rows[0].cells[0].text = "For Use of Secy- Treas."
        totals_table.rows[1].cells[0].text = "TOTALS"
        doc.save(template_path)

        pay_portal.fill_pay_voucher_docx(
            template_path=str(template_path),
            output_path=str(output_path),
            period_start="2026-05-03",
            period_end="2026-05-16",
            entries=[
                {
                    "entry_date": "2026-05-04",
                    "display_name": "Nick Craig",
                    "local_number": "3106",
                    "address": "4076 Union Hall Pl",
                    "hours": 2.5,
                    "lost_wage_hourly_rate": 52.25,
                    "mileage_miles": 12.25,
                    "mileage_amount": 8.88,
                    "meals_amount": 22.50,
                    "president_diff_hours": 1.0,
                    "president_diff_amount": 17.00,
                },
                {
                    "entry_date": "2026-05-12",
                    "display_name": "Nick Craig",
                    "local_number": "3106",
                    "hours": 4.0,
                    "lost_wage_hourly_rate": 40.00,
                    "mileage_miles": 3.5,
                    "mileage_amount": 2.54,
                    "miscellaneous_amount": 10.00,
                    "president_diff_hours": 0.5,
                    "president_diff_amount": 8.50,
                },
            ],
        )

        filled = Document(output_path)
        week1 = filled.tables[0]
        week2 = filled.tables[1]
        grey_totals = filled.tables[2].rows[1].cells

        self.assertEqual(week1.rows[1].cells[2].text, "2.50")
        self.assertEqual(week1.rows[2].cells[2].text, "12.25")
        self.assertEqual(week1.rows[7].cells[2].text, "1.00")
        self.assertEqual(week1.rows[8].cells[2].text, "3.50")
        self.assertEqual(week1.rows[1].cells[8].text, "2.50")
        self.assertEqual(week1.rows[2].cells[8].text, "12.25")
        self.assertEqual(week1.rows[7].cells[8].text, "1.00")
        self.assertEqual(week1.rows[8].cells[8].text, "3.50")

        self.assertEqual(week2.rows[1].cells[3].text, "4.00")
        self.assertEqual(week2.rows[2].cells[3].text, "3.50")
        self.assertEqual(week2.rows[7].cells[3].text, "0.50")
        self.assertEqual(week2.rows[8].cells[3].text, "4.50")

        self.assertEqual(grey_totals[1].text, "$ 290.63")
        self.assertEqual(grey_totals[2].text, "$ 11.42")
        self.assertEqual(grey_totals[4].text, "$ 22.50")
        self.assertEqual(grey_totals[6].text, "$ 10.00")
        self.assertEqual(grey_totals[7].text, "$ 360.05")

    def test_voucher_narrative_line_splits_date_label_for_back_page(self) -> None:
        label, narrative = pay_portal._split_narrative_line(  # noqa: SLF001
            "Mon 05/04/2026: DEMO TRAINING - reviewed mileage."
        )

        self.assertEqual(label, "Mon 05/04/2026")
        self.assertEqual(narrative, "DEMO TRAINING - reviewed mileage.")

    def test_voucher_front_receipt_narrative_rows_are_removed(self) -> None:
        class _TableXml:
            def __init__(self) -> None:
                self.removed = []

            def remove(self, row_tr):  # noqa: ANN001
                self.removed.append(row_tr)

        receipt_row = SimpleNamespace(
            cells=[SimpleNamespace(text="Attach necessary receipts - Explain reason for expense:")],
            _tr="receipt-row",
        )
        keep_row = SimpleNamespace(cells=[SimpleNamespace(text="Hours")], _tr="keep-row")
        table_xml = _TableXml()
        doc = SimpleNamespace(
            tables=[
                SimpleNamespace(rows=[keep_row], _tbl=_TableXml()),
                SimpleNamespace(rows=[receipt_row], _tbl=table_xml),
                SimpleNamespace(rows=[receipt_row], _tbl=_TableXml()),
            ]
        )

        pay_portal._remove_front_receipt_explanation_rows(doc)  # noqa: SLF001

        self.assertEqual(table_xml.removed, ["receipt-row"])

    def test_voucher_totals_row_found_by_label(self) -> None:
        totals_row = SimpleNamespace(
            cells=[SimpleNamespace(text="TOTALS"), SimpleNamespace(text="$"), SimpleNamespace(text="$")]
        )
        doc = SimpleNamespace(
            tables=[
                SimpleNamespace(rows=[SimpleNamespace(cells=[SimpleNamespace(text="Hours")])]),
                SimpleNamespace(rows=[totals_row]),
            ]
        )

        self.assertIs(pay_portal._find_totals_row(doc), totals_row)  # noqa: SLF001

    async def test_irs_sync_stages_new_rate_and_ignores_active_rate(self) -> None:
        cfg = self._cfg().pay_portal

        active = await sync_irs_mileage_rate_candidates(
            self.db,
            pay_cfg=cfg,
            http_get=lambda _url, timeout=15: _FakeHttpResponse(
                "<p>Beginning Jan. 1, 2026, 72.5 cents per mile driven for business use.</p>"
            ),
            source_urls=("https://www.irs.gov/active",),
        )
        self.assertEqual(active["detected"], [])
        self.assertEqual(active["skipped_existing"], 1)

        detected = await sync_irs_mileage_rate_candidates(
            self.db,
            pay_cfg=cfg,
            http_get=lambda _url, timeout=15: _FakeHttpResponse(
                "<p>Beginning Jan. 1, 2027, 75.5 cents per mile driven for business use.</p>"
            ),
            source_urls=("https://www.irs.gov/new",),
        )
        self.assertEqual(len(detected["detected"]), 1)
        self.assertEqual(detected["detected"][0]["rate_year"], "2027")
        self.assertEqual(detected["detected"][0]["rate_per_mile"], "0.755")

        rows = await list_irs_rate_candidates(self.db, status="pending")
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["status"], "pending")

    async def test_irs_sync_failure_leaves_active_rates_unchanged(self) -> None:
        cfg = self._cfg().pay_portal

        result = await sync_irs_mileage_rate_candidates(
            self.db,
            pay_cfg=cfg,
            http_get=lambda _url, timeout=15: _FakeHttpResponse("<p>No mileage data here.</p>"),
            source_urls=("https://www.irs.gov/malformed",),
        )
        settings = await pay_settings(self.db, pay_cfg=cfg)

        self.assertEqual(result["detected"], [])
        self.assertEqual(len(result["failures"]), 1)
        self.assertEqual(settings["irs_rates"]["2026"], "0.725")

    async def test_irs_rate_approval_updates_active_settings(self) -> None:
        cfg = self._cfg().pay_portal
        result = await sync_irs_mileage_rate_candidates(
            self.db,
            pay_cfg=cfg,
            http_get=lambda _url, timeout=15: _FakeHttpResponse(
                "<p>Beginning Jan. 1, 2027, 75.5 cents per mile driven for business use.</p>"
            ),
            source_urls=("https://www.irs.gov/new",),
        )
        candidate_id = int(result["detected"][0]["id"])

        approved = await approve_irs_rate_candidate(
            self.db,
            candidate_id=candidate_id,
            actor="treasurer@example.org",
            pay_cfg=cfg,
        )
        settings = await pay_settings(self.db, pay_cfg=cfg)

        self.assertEqual(approved["status"], "approved")
        self.assertEqual(approved["approved_by"], "treasurer@example.org")
        self.assertEqual(settings["irs_rates"]["2027"], "0.755")

    async def test_non_treasurer_route_cannot_approve_irs_rate_candidate(self) -> None:
        cfg = self._cfg()
        cfg.officer_auth.enabled = True
        request = _Request(
            state=SimpleNamespace(cfg=cfg, db=self.db),
            session={
                "officer_user": {
                    "email": "officer@example.org",
                    "display_name": "Officer",
                    "role": "officer",
                    "contract_scopes": [],
                    "exp": int(time.time()) + 3600,
                }
            },
        )

        with self.assertRaises(HTTPException) as exc:
            await approve_pay_irs_rate(1, request)

        self.assertEqual(exc.exception.status_code, 403)

    async def test_president_differential_uses_effective_wage_scale_and_basis(self) -> None:
        await upsert_wage_scale(
            self.db,
            effective_date="2026-01-01",
            weekly_basis_hours=40,
            target_weekly_amount=2065.50,
            actual_weekly_amount=None,
            target_multiplier=1.20,
            target_scale="36",
            actual_scale="base",
            notes=None,
            updated_by="test",
        )
        result = await calculate_president_differential(
            self.db,
            entry_date="2026-05-10",
            weekly_basis_hours=40,
            president_diff_hours=10,
            target_scale="36",
            target_multiplier=1.20,
            lost_wage_input_type="hourly",
            lost_wage_amount=45,
        )
        self.assertEqual(result.presidential_hourly_rate, pay_portal.Decimal("62.00"))
        self.assertEqual(result.diff_rate, pay_portal.Decimal("17.00"))
        self.assertEqual(result.diff_amount, pay_portal.Decimal("170.00"))
        self.assertEqual(result.lost_wage_hourly_rate, pay_portal.Decimal("45.00"))

        await upsert_wage_scale(
            self.db,
            effective_date="2026-01-01",
            weekly_basis_hours=37.5,
            target_weekly_amount=2065.50,
            actual_weekly_amount=None,
            target_multiplier=1.20,
            target_scale="36",
            actual_scale="base",
            notes=None,
            updated_by="test",
        )
        result_37 = await calculate_president_differential(
            self.db,
            entry_date="2026-05-10",
            weekly_basis_hours=37.5,
            president_diff_hours=2,
            target_scale="36",
            target_multiplier=1.20,
            lost_wage_input_type="weekly",
            lost_wage_amount=1900.00,
        )
        self.assertEqual(result_37.presidential_hourly_rate, pay_portal.Decimal("67.00"))
        self.assertEqual(result_37.diff_amount, pay_portal.Decimal("32.66"))
        self.assertEqual(result_37.lost_wage_hourly_rate, pay_portal.Decimal("50.67"))

    async def test_president_differential_zero_floors_when_actual_exceeds_target(self) -> None:
        await upsert_wage_scale(
            self.db,
            effective_date="2026-01-01",
            weekly_basis_hours=40,
            target_weekly_amount=1000,
            actual_weekly_amount=None,
            target_multiplier=1.20,
            target_scale="36",
            actual_scale="base",
            notes=None,
            updated_by="test",
        )
        result = await calculate_president_differential(
            self.db,
            entry_date="2026-02-01",
            weekly_basis_hours=40,
            president_diff_hours=8,
            target_scale="36",
            target_multiplier=1.20,
            lost_wage_input_type="hourly",
            lost_wage_amount=50,
        )
        self.assertEqual(result.diff_rate, pay_portal.Decimal("0.00"))
        self.assertEqual(result.diff_amount, pay_portal.Decimal("0.00"))

    def test_pay_profile_snapshot_calculates_supported_pay_bases(self) -> None:
        hourly = calculate_pay_profile_snapshot(
            pay_basis="hourly",
            base_wage_input_type="weekly",
            base_wage_amount=1800,
            weekly_basis_hours=40,
        )
        weekly = calculate_pay_profile_snapshot(
            pay_basis="weekly",
            base_wage_input_type="hourly",
            base_wage_amount=45,
            weekly_basis_hours=40,
        )
        commission = calculate_pay_profile_snapshot(
            pay_basis="commission",
            base_wage_input_type="hourly",
            base_wage_amount=45,
            weekly_basis_hours=40,
            commission_month_1_amount=1000,
            commission_month_2_amount=1400,
            commission_month_3_amount=1800,
        )
        president = calculate_pay_profile_snapshot(
            pay_basis="president",
            base_wage_input_type="hourly",
            base_wage_amount=45,
            weekly_basis_hours=40,
            commission_month_1_amount=1000,
            commission_month_2_amount=1400,
            commission_month_3_amount=1800,
        )
        expense_only = calculate_pay_profile_snapshot(
            pay_basis="expense_only",
            base_wage_input_type="weekly",
            base_wage_amount=1800,
            weekly_basis_hours=40,
        )

        self.assertEqual(hourly["base_wage_input_type"], "weekly")
        self.assertEqual(hourly["calculated_hourly_rate"], pay_portal.Decimal("45.00"))
        self.assertEqual(weekly["base_wage_input_type"], "hourly")
        self.assertEqual(weekly["calculated_hourly_rate"], pay_portal.Decimal("45.00"))
        self.assertEqual(commission["commission_hourly_rate"], pay_portal.Decimal("8.75"))
        self.assertEqual(commission["calculated_hourly_rate"], pay_portal.Decimal("53.75"))
        self.assertEqual(president["commission_hourly_rate"], pay_portal.Decimal("0.00"))
        self.assertEqual(president["calculated_hourly_rate"], pay_portal.Decimal("45.00"))
        self.assertEqual(expense_only["calculated_hourly_rate"], pay_portal.Decimal("0.00"))

    async def test_president_entry_uses_stored_base_wage_profile_without_commission(self) -> None:
        await upsert_wage_scale(
            self.db,
            effective_date="2026-01-01",
            weekly_basis_hours=40,
            target_weekly_amount=2065.50,
            actual_weekly_amount=None,
            target_multiplier=1.20,
            target_scale="36",
            actual_scale="base",
            notes=None,
            updated_by="test",
        )
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-10"))
        actor = PayActor("president@example.org", "President", "officer", True, False, False)
        profile = await self._save_profile(
            email="president@example.org",
            display_name="President",
            pay_basis="president",
            base_wage_input_type="weekly",
            base_wage_amount=1800,
            commission_month_1_amount=1000,
            commission_month_2_amount=1400,
            commission_month_3_amount=1800,
        )
        entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={
                "entry_date": "2026-05-10",
                "display_name": "President",
                "hours": 1,
                "president_diff_hours": 3,
            },
            pay_cfg=self._cfg().pay_portal,
        )
        self.assertEqual(profile["commission_hourly_rate"], 0.0)
        self.assertEqual(entry["lost_wage_input_type"], "president")
        self.assertEqual(entry["lost_wage_amount"], 1800.0)
        self.assertEqual(entry["lost_wage_hourly_rate"], 62.0)
        self.assertEqual(entry["hourly_rate"], 62.0)
        self.assertEqual(entry["president_diff_hours"], 3.0)
        self.assertEqual(entry["president_diff_amount"], 51.0)
        treasurer_entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=PayActor("treasurer@example.org", "Treasurer", "treasurer", True, True, True),
            data={
                "user_email": "president@example.org",
                "entry_date": "2026-05-11",
                "display_name": "President",
                "hours": 1,
                "president_diff_hours": 2,
            },
            pay_cfg=self._cfg().pay_portal,
        )
        self.assertEqual(treasurer_entry["lost_wage_hourly_rate"], 62.0)
        self.assertEqual(treasurer_entry["president_diff_hours"], 2.0)
        self.assertEqual(treasurer_entry["president_diff_amount"], 34.0)
        president_actor_entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=PayActor(
                "president@example.org",
                "President",
                "president",
                True,
                False,
                False,
                is_president=True,
            ),
            data={
                "entry_date": "2026-05-12",
                "display_name": "President",
                "hours": 1,
                "president_diff_hours": 2.5,
            },
            pay_cfg=self._cfg().pay_portal,
        )
        self.assertEqual(president_actor_entry["lost_wage_hourly_rate"], 62.0)
        self.assertEqual(president_actor_entry["president_diff_hours"], 2.5)
        self.assertEqual(president_actor_entry["president_diff_amount"], 42.5)

    async def test_president_scheduled_hours_cap_includes_union_and_differential_hours(self) -> None:
        await upsert_wage_scale(
            self.db,
            effective_date="2026-01-01",
            weekly_basis_hours=40,
            target_weekly_amount=2065.50,
            actual_weekly_amount=None,
            target_multiplier=1.20,
            target_scale="36",
            actual_scale="base",
            notes=None,
            updated_by="test",
        )
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-10"))
        actor = PayActor("president@example.org", "President", "president", True, False, False, is_president=True)
        await self._save_profile(
            email="president@example.org",
            display_name="President",
            pay_basis="president",
            base_wage_input_type="hourly",
            base_wage_amount=45,
        )
        entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={"entry_date": "2026-05-10", "hours": 4, "president_diff_hours": 4},
            pay_cfg=self._cfg().pay_portal,
        )
        self.assertEqual(entry["hours"], 4.0)
        self.assertEqual(entry["president_diff_hours"], 4.0)

        fills_week = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={"entry_date": "2026-05-11", "hours": 30, "president_diff_hours": 2},
            pay_cfg=self._cfg().pay_portal,
        )
        self.assertEqual(fills_week["hours"], 30.0)
        self.assertEqual(fills_week["president_diff_hours"], 2.0)

        with self.assertRaisesRegex(ValueError, "cannot exceed 40"):
            await upsert_entry(
                self.db,
                period_id=str(period["id"]),
                actor=actor,
                data={"entry_date": "2026-05-12", "hours": 0, "president_diff_hours": 0.25},
                pay_cfg=self._cfg().pay_portal,
            )

    async def test_commission_compensation_uses_three_month_average_over_160(self) -> None:
        result = calculate_commission_compensation(
            base_wage_input_type="hourly",
            base_wage_amount=45,
            weekly_basis_hours=40,
            commission_month_1_amount=1000,
            commission_month_2_amount=1400,
            commission_month_3_amount=1800,
        )
        self.assertEqual(result.commission_average_monthly, pay_portal.Decimal("1400.00"))
        self.assertEqual(result.commission_hourly_rate, pay_portal.Decimal("8.75"))
        self.assertEqual(result.calculated_hourly_rate, pay_portal.Decimal("53.75"))

    async def test_commission_entry_requires_last_month_payroll_proof(self) -> None:
        await upsert_wage_scale(
            self.db,
            effective_date="2026-01-01",
            weekly_basis_hours=40,
            target_weekly_amount=2065.50,
            actual_weekly_amount=None,
            target_multiplier=1.20,
            target_scale="36",
            actual_scale="base",
            notes=None,
            updated_by="test",
        )
        actor = PayActor("commission@example.org", "Commission", "officer", True, False, False)
        profile = await self._save_profile(
            email="commission@example.org",
            display_name="Commission",
            pay_basis="commission",
            base_wage_input_type="hourly",
            base_wage_amount=45,
            weekly_basis_hours=40,
            commission_month_1_amount=1000,
            commission_month_2_amount=1400,
            commission_month_3_amount=1800,
        )
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-10"))
        expense_entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={
                "entry_date": "2026-05-09",
                "mileage_miles": 10,
                "mileage_rate": 0.725,
                "mileage_amount": 7.25,
            },
            pay_cfg=self._cfg().pay_portal,
        )
        self.assertIsNone(expense_entry["compensation_stub_id"])
        self.assertEqual(expense_entry["mileage_rate"], 0.725)
        self.assertEqual(expense_entry["mileage_amount"], 7.25)

        with self.assertRaisesRegex(ValueError, "commission payroll proof for 2026-04"):
            await upsert_entry(
                self.db,
                period_id=str(period["id"]),
                actor=actor,
                data={
                    "entry_date": "2026-05-10",
                    "display_name": "Typed Overwrite",
                    "hours": 2,
                },
                pay_cfg=self._cfg().pay_portal,
            )

        wrong_month_stub = await store_compensation_stub(
            self.db,
            cfg=self._cfg(),
            actor=actor,
            user_email=None,
            base_wage_input_type="hourly",
            base_wage_amount=45,
            weekly_basis_hours=40,
            commission_month_1_amount=1000,
            commission_month_2_amount=1400,
            commission_month_3_amount=1800,
            payroll_month="2026-03",
            filename="march-payroll.pdf",
            content_type="application/pdf",
            content=b"%PDF-1.4\nmarch stub\n%%EOF",
            scan=False,
        )
        self.assertEqual(wrong_month_stub["payroll_month"], "2026-03")
        with self.assertRaisesRegex(ValueError, "commission payroll proof for 2026-04"):
            await upsert_entry(
                self.db,
                period_id=str(period["id"]),
                actor=actor,
                data={
                    "entry_date": "2026-05-10",
                    "display_name": "Typed Overwrite",
                    "hours": 2,
                },
                pay_cfg=self._cfg().pay_portal,
            )

        matching_stub = await store_compensation_stub(
            self.db,
            cfg=self._cfg(),
            actor=actor,
            user_email=None,
            base_wage_input_type="hourly",
            base_wage_amount=45,
            weekly_basis_hours=40,
            commission_month_1_amount=1000,
            commission_month_2_amount=1400,
            commission_month_3_amount=1800,
            payroll_month="2026-04",
            filename="april-payroll.pdf",
            content_type="application/pdf",
            content=b"%PDF-1.4\napril stub\n%%EOF",
            scan=False,
        )
        entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={
                "entry_date": "2026-05-10",
                "display_name": "Typed Overwrite",
                "hours": 2,
                "hourly_rate": 1,
                "lost_wage_amount": 1,
            },
            pay_cfg=self._cfg().pay_portal,
        )
        self.assertEqual(profile["calculated_hourly_rate"], 53.75)
        self.assertEqual(entry["display_name"], "Commission")
        self.assertEqual(entry["lost_wage_input_type"], "commission")
        self.assertEqual(entry["compensation_stub_id"], matching_stub["id"])
        self.assertEqual(entry["lost_wage_hourly_rate"], 53.75)
        self.assertEqual(entry["hourly_rate"], 53.75)
        self.assertEqual(entry["president_diff_amount"], 0.0)

    async def test_missing_profile_allows_expenses_only_and_rejects_lost_wage_hours(self) -> None:
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-10"))
        actor = PayActor("no-profile@example.org", "No Profile", "officer", True, False, False)

        expense_entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={
                "entry_date": "2026-05-10",
                "mileage_miles": 10,
                "mileage_rate": 0.725,
                "mileage_amount": 7.25,
            },
            pay_cfg=self._cfg().pay_portal,
        )

        self.assertEqual(expense_entry["lost_wage_hourly_rate"], 0.0)
        self.assertEqual(expense_entry["mileage_amount"], 7.25)

        with self.assertRaisesRegex(ValueError, "pay profile required"):
            await upsert_entry(
                self.db,
                period_id=str(period["id"]),
                actor=actor,
                data={
                    "entry_date": "2026-05-11",
                    "hours": 1,
                },
                pay_cfg=self._cfg().pay_portal,
            )

    async def test_treasurer_entry_for_another_user_uses_target_users_profile(self) -> None:
        await self._save_profile(
            email="member@example.org",
            display_name="Member Example",
            pay_basis="weekly",
            base_wage_input_type="weekly",
            base_wage_amount=2000,
            weekly_basis_hours=40,
        )
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-10"))
        treasurer = PayActor("treasurer@example.org", "Treasurer", "treasurer", True, True, True)

        entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=treasurer,
            data={
                "user_email": "member@example.org",
                "entry_date": "2026-05-10",
                "hours": 2,
                "hourly_rate": 1,
                "lost_wage_amount": 1,
            },
            pay_cfg=self._cfg().pay_portal,
        )

        self.assertEqual(entry["user_email"], "member@example.org")
        self.assertEqual(entry["display_name"], "Member Example")
        self.assertEqual(entry["lost_wage_input_type"], "weekly")
        self.assertEqual(entry["lost_wage_amount"], 2000.0)
        self.assertEqual(entry["lost_wage_hourly_rate"], 50.0)

    async def test_guest_cannot_upload_compensation_stub_for_another_member(self) -> None:
        actor = PayActor("guest@example.org", "Guest", "guest", False, False, False, is_guest=True)
        with self.assertRaises(PermissionError):
            await store_compensation_stub(
                self.db,
                cfg=self._cfg(),
                actor=actor,
                user_email="other@example.org",
                base_wage_input_type="hourly",
                base_wage_amount=45,
                weekly_basis_hours=40,
                commission_month_1_amount=0,
                commission_month_2_amount=0,
                commission_month_3_amount=0,
                filename="pay-stub.pdf",
                content_type="application/pdf",
                content=b"%PDF-1.4\nstub\n%%EOF",
                scan=False,
            )

    async def test_entry_records_submitter_daily_certification(self) -> None:
        actor = PayActor("member@example.org", "Member Example", "officer", False, False, False)
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-10"))

        entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=actor,
            data={
                "entry_date": "2026-05-10",
                "mileage_miles": 4,
                "mileage_rate": 0.725,
                "mileage_amount": 2.90,
                "submitter_certified": True,
                "submitter_certification_text": "I certify this entry.",
            },
            pay_cfg=self._cfg().pay_portal,
        )

        self.assertIsNotNone(entry["submitter_certified_at_utc"])
        self.assertEqual(entry["submitter_certified_by"], "member@example.org")
        self.assertEqual(entry["submitter_certification_text"], "I certify this entry.")

    async def test_pay_packet_signer_order_routes_only_to_president(self) -> None:
        signers, president_index = pay_packet_signer_order(
            grouped_entry_emails=["member@example.org", "President@Example.Org", "other@example.org"],
            president_signer_email="president@example.org",
        )
        self.assertEqual(signers, ["president@example.org"])
        self.assertEqual(president_index, 1)

    async def test_receipt_validation_blocks_bad_types_and_infected_scan(self) -> None:
        with self.assertRaisesRegex(ValueError, "only PDF"):
            validate_receipt_file(
                filename="receipt.txt",
                content_type="text/plain",
                content=b"hello",
                max_file_bytes=1000,
            )

        period_id, entry_id = await self._open_period_and_entry()
        actor = PayActor("officer@example.org", "Officer Example", "officer", True, False, False)
        with patch("grievance_api.services.pay_portal.scan_with_clamav", side_effect=ValueError("FOUND")):
            with self.assertRaisesRegex(ValueError, "FOUND"):
                await store_attachment(
                    self.db,
                    cfg=self._cfg(),
                    period_id=period_id,
                    entry_id=entry_id,
                    actor=actor,
                    attachment_type="receipt",
                    filename="receipt.pdf",
                    content_type="application/pdf",
                    content=b"%PDF-1.4\n%%EOF",
                )

        with patch("grievance_api.services.pay_portal.scan_with_clamav", return_value="stream: OK"):
            saved = await store_attachment(
                self.db,
                cfg=self._cfg(),
                period_id=period_id,
                entry_id=entry_id,
                actor=actor,
                attachment_type="receipt",
                filename="receipt.pdf",
                content_type="application/pdf",
                content=b"%PDF-1.4\n%%EOF",
            )
        self.assertEqual(saved["scan_status"], "clean")

    async def test_guest_cannot_attach_to_another_users_entry(self) -> None:
        owner = PayActor("owner@example.org", "Owner", "guest", False, False, False, is_guest=True)
        period_id, entry_id = await self._open_period_and_entry(actor=owner)
        guest = PayActor("guest@example.org", "Guest", "guest", False, False, False, is_guest=True)
        with patch("grievance_api.services.pay_portal.scan_with_clamav", return_value="stream: OK"):
            with self.assertRaises(PermissionError):
                await store_attachment(
                    self.db,
                    cfg=self._cfg(),
                    period_id=period_id,
                    entry_id=entry_id,
                    actor=guest,
                    attachment_type="receipt",
                    filename="receipt.pdf",
                    content_type="application/pdf",
                    content=b"%PDF-1.4\n%%EOF",
                )

    async def test_lock_period_groups_each_person_voucher_with_their_supporting_docs(self) -> None:
        cfg = self._cfg()
        await pay_portal.save_pay_settings(
            self.db,
            setting={"president_email": "president@example.org", "treasurer_emails": ["treasurer@example.org"]},
            updated_by="test",
        )
        treasurer = PayActor("treasurer@example.org", "Treasurer", "treasurer", True, True, True)
        nick = PayActor("nick@example.org", "Nick Craig", "officer", False, False, False)
        vicky = PayActor("vicky@example.org", "Vicky Goll", "officer", False, False, False)
        await self._save_profile(email=nick.email, display_name=nick.display_name)
        await self._save_profile(email=vicky.email, display_name=vicky.display_name)
        period = await pay_portal.ensure_pay_period(self.db, for_date=date.fromisoformat("2026-05-17"))
        nick_entry_1 = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=nick,
            data={"entry_date": "2026-05-17", "hours": 2, "notes": "Nick day 1"},
            pay_cfg=cfg.pay_portal,
        )
        await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=nick,
            data={"entry_date": "2026-05-18", "hours": 3, "notes": "Nick day 2"},
            pay_cfg=cfg.pay_portal,
        )
        vicky_entry = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=vicky,
            data={"entry_date": "2026-05-17", "hours": 4, "notes": "Vicky day 1"},
            pay_cfg=cfg.pay_portal,
        )
        support_root = Path(self.tmpdir.name) / "packet-support"
        support_root.mkdir(parents=True, exist_ok=True)
        nick_support = support_root / "nick-mileage.pdf"
        vicky_support = support_root / "vicky-receipt.pdf"
        excluded_support = support_root / "nick-rejected.pdf"
        for file_path in (nick_support, vicky_support, excluded_support):
            file_path.write_bytes(b"%PDF-1.4\n%%EOF")

        async def _insert_support(attachment_id: str, entry_id: str, user_email: str, file_path: Path) -> None:
            await self.db.exec(
                """INSERT INTO pay_attachments(
                     id, period_id, entry_id, uploaded_by, attachment_type, original_filename, stored_filename,
                     local_path, content_type, size_bytes, sha256, scan_status, scan_result, created_at_utc
                   ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (
                    attachment_id,
                    str(period["id"]),
                    entry_id,
                    user_email,
                    "receipt",
                    file_path.name,
                    file_path.name,
                    str(file_path),
                    "application/pdf",
                    file_path.stat().st_size,
                    attachment_id,
                    "clean",
                    "test clean",
                    pay_portal.utcnow(),
                ),
            )

        await _insert_support("nick-support", str(nick_entry_1["id"]), nick.email, nick_support)
        await _insert_support("vicky-support", str(vicky_entry["id"]), vicky.email, vicky_support)
        rejected = await upsert_entry(
            self.db,
            period_id=str(period["id"]),
            actor=nick,
            data={"entry_date": "2026-05-19", "hours": 1, "notes": "Rejected Nick day"},
            pay_cfg=cfg.pay_portal,
        )
        await _insert_support("excluded-support", str(rejected["id"]), nick.email, excluded_support)
        await pay_portal.review_pay_entry(
            self.db,
            entry_id=str(rejected["id"]),
            actor=treasurer,
            review_status="rejected",
            review_note="wrong period",
        )

        voucher_fill_calls: list[tuple[str, list[str]]] = []

        def _capture_fill(**kwargs):  # noqa: ANN003
            if not kwargs.get("include_signature_placeholders"):
                voucher_fill_calls.append(
                    (
                        str(kwargs["entries"][0]["user_email"]),
                        [str(entry["entry_date"]) for entry in kwargs["entries"]],
                    )
                )
            _fake_fill_docx(**kwargs)

        with patch("grievance_api.services.pay_portal.fill_pay_voucher_docx", side_effect=_capture_fill), patch(
            "grievance_api.services.pay_portal.merge_pdfs", side_effect=_fake_merge
        ):
            await lock_period_and_send_packet(
                db=self.db,
                cfg=cfg,
                graph=_FakeGraph(),
                docuseal=_FakeDocuSeal(),
                period_id=str(period["id"]),
                actor=treasurer,
                president_signer_email=None,
                docx_to_pdf_func=_fake_docx_to_pdf,
            )

        self.assertEqual(
            voucher_fill_calls,
            [
                ("nick@example.org", ["2026-05-17", "2026-05-18"]),
                ("vicky@example.org", ["2026-05-17"]),
            ],
        )
        packet_row = await self.db.fetchone(
            "SELECT unsigned_packet_path FROM pay_packets WHERE period_id=?",
            (str(period["id"]),),
        )
        packet_text = Path(str(packet_row[0])).read_text(encoding="utf-8")
        self.assertLess(packet_text.index("01-Nick Craig.pdf"), packet_text.index("nick-mileage.pdf"))
        self.assertLess(packet_text.index("nick-mileage.pdf"), packet_text.index("02-Vicky Goll.pdf"))
        self.assertLess(packet_text.index("02-Vicky Goll.pdf"), packet_text.index("vicky-receipt.pdf"))
        self.assertNotIn("nick-rejected.pdf", packet_text)

    async def test_lock_period_builds_voucher_first_packet_and_sends_docuseal(self) -> None:
        cfg = self._cfg()
        await pay_portal.save_pay_settings(
            self.db,
            setting={"president_email": "president@example.org", "treasurer_emails": ["treasurer@example.org"]},
            updated_by="test",
        )
        actor = PayActor("treasurer@example.org", "Treasurer", "treasurer", True, True, True)
        period_id, _ = await self._open_period_and_entry(actor=actor)
        graph = _FakeGraph()
        docuseal = _FakeDocuSeal()
        with patch("grievance_api.services.pay_portal.fill_pay_voucher_docx", _fake_fill_docx), patch(
            "grievance_api.services.pay_portal.merge_pdfs",
            _fake_merge,
        ):
            result = await lock_period_and_send_packet(
                db=self.db,
                cfg=cfg,
                graph=graph,
                docuseal=docuseal,
                period_id=period_id,
                actor=actor,
                president_signer_email=None,
                docx_to_pdf_func=_fake_docx_to_pdf,
        )

        self.assertEqual(result["status"], "awaiting_signature")
        self.assertEqual(docuseal.create_calls[0]["signers"], ["president@example.org"])
        self.assertEqual(result["signer_order"], ["president@example.org"])
        self.assertEqual(result["signing_link"], "https://docuseal.local/sign/1")
        self.assertEqual(docuseal.create_calls[0]["metadata"]["president_signer_email"], "president@example.org")
        self.assertIn(b".pdf", docuseal.create_calls[0]["pdf_bytes"])
        packet_row = await self.db.fetchone("SELECT status, docuseal_submission_id FROM pay_packets WHERE period_id=?", (period_id,))
        self.assertEqual(packet_row[0], "awaiting_signature")
        self.assertEqual(packet_row[1], "pay-ds-1")

    async def test_pay_docuseal_completion_files_signed_packet_and_notifies_treasurer(self) -> None:
        cfg = self._cfg()
        period_id, _ = await self._open_period_and_entry()
        packet_dir = Path(self.data_root) / "pay" / period_id / "packet" / "packet-1"
        packet_dir.mkdir(parents=True)
        unsigned = packet_dir / "unsigned.pdf"
        unsigned.write_bytes(b"%PDF-1.4\nunsigned\n")
        await self.db.exec(
            "UPDATE pay_periods SET status='awaiting_signature', sharepoint_folder_path='Pay Portal/2026/test' WHERE id=?",
            (period_id,),
        )
        await self.db.exec(
            """INSERT INTO pay_packets(
                 id, period_id, revision, status, voucher_paths_json, voucher_pdf_paths_json,
                 unsigned_packet_path, unsigned_packet_sha256, docuseal_submission_id,
                 created_at_utc, updated_at_utc
               ) VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
            (
                "packet-1",
                period_id,
                1,
                "awaiting_signature",
                "[]",
                "[]",
                str(unsigned),
                "sha",
                "pay-ds-1",
                "now",
                "now",
            ),
        )
        graph = _FakeGraph()
        mailer = _FakeMailer()
        result = await handle_pay_docuseal_completion(
            db=self.db,
            cfg=cfg,
            graph=graph,
            mailer=mailer,
            docuseal=_FakeDocuSeal(),
            submission_id="pay-ds-1",
            payload={"event_type": "completed"},
        )
        self.assertEqual(result["handled"], True)
        period_row = await self.db.fetchone("SELECT status FROM pay_periods WHERE id=?", (period_id,))
        self.assertEqual(period_row[0], "completed")
        packet_row = await self.db.fetchone("SELECT status, signed_packet_path FROM pay_packets WHERE id='packet-1'")
        self.assertEqual(packet_row[0], "completed")
        self.assertTrue(Path(packet_row[1]).exists())
        self.assertEqual(mailer.calls[0]["to_recipients"], ["treasurer@example.org"])

    async def test_create_revision_opens_new_period_revision(self) -> None:
        period_id, _ = await self._open_period_and_entry()
        actor = PayActor("treasurer@example.org", "Treasurer", "treasurer", True, True, True)
        revised = await create_revision(self.db, period_id=period_id, actor=actor)
        self.assertEqual(revised["revision"], 2)
        self.assertEqual(revised["status"], "open")
