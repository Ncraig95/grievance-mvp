from __future__ import annotations

import asyncio
import base64
import csv
import hashlib
import html
import io
import json
import logging
import os
import random
import re
import shutil
import socket
import subprocess
import threading
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from decimal import Decimal, ROUND_CEILING, ROUND_HALF_UP
from pathlib import Path
from typing import Any
from uuid import uuid4

import requests
from docx import Document
from docx.shared import Pt
from PIL import Image

from ..db.db import Db, utcnow
from .graph_mail import MailAttachment
from .signature_workflow import resolve_docuseal_template_id


_SAFE_FILENAME_RE = re.compile(r"[^A-Za-z0-9._ -]+")
_PAYROLL_MONTH_RE = re.compile(r"^\d{4}-(0[1-9]|1[0-2])$")
_ANCHOR_PERIOD_START = date(2025, 9, 7)
_CURRENCY = Decimal("0.01")
_MILES = Decimal("0.01")
_METERS_PER_MILE = Decimal("1609.344")
_COMMISSION_HOUR_DIVISOR = Decimal("160")
_PAY_FORM_KEY = "pay_portal_packet"
_PAY_PORTAL_LOGGER = logging.getLogger("grievance_api.pay_portal")
_IRS_RATE_QUANT = Decimal("0.001")
_GOOGLE_LEG_CACHE_LOCK = threading.Lock()
_GOOGLE_LEG_CACHE: dict[tuple[str, str], dict[str, object]] = {}
_GOOGLE_LEG_PREFETCH_WORKERS = 3
_DEFAULT_IRS_RATE_SOURCE_URLS = (
    "https://www.irs.gov/tax-professionals/standard-mileage-rates",
    "https://www.irs.gov/newsroom/irs-sets-2026-business-standard-mileage-rate-at-725-cents-per-mile-up-25-cents",
)
_COMMON_PLACE_CONFIG_EXTENSIONS = {".json", ".yaml", ".yml", ".csv", ".txt"}
_RECEIPT_CONTENT_TYPES = {
    "application/pdf": ".pdf",
    "image/jpeg": ".jpg",
    "image/png": ".png",
}
_PAY_PROFILE_BASIS_VALUES = {"hourly", "weekly", "commission", "president", "expense_only"}
_PAY_PROFILE_STATUS_VALUES = {"active", "disabled"}
_PAY_FUND_TYPES = {"sif", "growth"}
_PAY_FUND_STATUS_VALUES = {"active", "closed"}
_PAY_FUND_LEDGER_TYPES = {"advance", "reimbursement_submitted", "reimbursement_received", "adjustment"}
_FUND_FICA_RATE = Decimal("0.0765")


def _pay_lock_log(stage: str, **fields: object) -> None:
    safe_fields: dict[str, object] = {"stage": stage}
    parts = [f"stage={stage}"]
    for key, value in fields.items():
        if value is None or isinstance(value, (str, int, float, bool)):
            rendered: object = value
        else:
            rendered = str(value)
        safe_fields[key] = rendered
        parts.append(f"{key}={rendered!r}")
    _PAY_PORTAL_LOGGER.info(
        "pay_lock_send_stage %s",
        " ".join(parts),
        extra={"pay_portal": safe_fields},
    )


_DEFAULT_COMMON_PLACES: tuple[dict[str, str], ...] = (
    {"label": "Union Hall", "address": "4076 Union Hall Pl, Jacksonville, FL 32205, USA"},
    {"label": "Cumberland Industrial", "address": "350 Cumberland Industrial Ct, St. Augustine, FL 32095, USA"},
    {"label": "Phillips Hwy", "address": "11700 Phillips Hwy, Jacksonville, FL 32256, USA"},
    {"label": "Penman Rd", "address": "98 Penman Rd S, Jacksonville Beach, FL 32250, USA"},
    {"label": "Ed Ball", "address": "214 N Hogan St, Jacksonville, FL 32202, USA"},
    {"label": "City Hall COJ", "address": "117 W Duval St, Jacksonville, FL 32202, USA"},
    {"label": "207", "address": "660 FL-207, St. Augustine, FL 32084, USA"},
    {"label": "S 8th St", "address": "1910 S 8Th St, Fernandina Beach, FL 32034, USA"},
    {"label": "3rd St N", "address": "1824 3rd St N, Jacksonville Beach, FL 32250, USA"},
    {"label": "Atlantic Blvd 7553", "address": "7553 Atlantic Blvd, Jacksonville, FL 32211, USA"},
    {"label": "Beach Blvd 11317", "address": "11317 Beach Blvd, Jacksonville, FL 32246, USA"},
    {"label": "N Pearl St", "address": "424 N Pearl St, Jacksonville, FL 32202, USA"},
    {"label": "Ft Caroline Rd", "address": "6654 Ft Caroline Rd, Jacksonville, FL 32277, USA"},
    {"label": "Haydon Rd Workcenter", "address": "9209 Haydon Rd, Jacksonville, FL 32218, USA"},
    {"label": "Cassat Ave", "address": "1844 Cassat Ave, Jacksonville, FL 32210, USA"},
    {"label": "Edgewood Ave W", "address": "1441 Edgewood Ave W, Jacksonville, FL 32208, USA"},
    {"label": "Old Middleburg Rd N", "address": "2200 Old Middleburg Rd N, Jacksonville, FL 32210, USA"},
    {"label": "Normandy Blvd", "address": "6602 Normandy Blvd, Jacksonville, FL 32205, USA"},
    {"label": "N Main St", "address": "11741 N Main St, Jacksonville, FL 32218, USA"},
    {"label": "Talbot Ave", "address": "1710 Talbot Ave, Jacksonville, FL 32205, USA"},
    {"label": "St Augustine Rd", "address": "6234 St Augustine Rd, Jacksonville, FL 32217, USA"},
    {"label": "Hendricks Ave", "address": "2048 Hendricks Ave, Jacksonville, FL 32207, USA"},
    {"label": "St Johns Bluff Rd N", "address": "1001 St Johns Bluff Rd N, Jacksonville, FL 32225, USA"},
    {"label": "Historic Kings Rd S", "address": "9400 Historic Kings Rd S, Jacksonville, FL 32257, USA"},
    {"label": "Jammes Rd", "address": "5532 Jammes Rd, Jacksonville, FL 32244, USA"},
    {"label": "Southside Blvd", "address": "9039 Southside Blvd, Jacksonville, FL 32256, USA"},
    {"label": "Dennis St", "address": "2096 Dennis St, Jacksonville, FL 32204, USA"},
    {"label": "River City Dr", "address": "4663 River City Dr, Jacksonville, FL 32246, USA"},
    {"label": "N Main St SOC", "address": "11741 N Main St, Jacksonville, FL 32218, USA"},
    {"label": "Crosshill Blvd", "address": "9508 Crosshill Blvd, Jacksonville, FL 32222, USA"},
    {"label": "Nautica Dr", "address": "725 Nautica Dr, Jacksonville, FL 32218, USA"},
    {"label": "San Jose Blvd", "address": "11113-102 San Jose Blvd, Jacksonville, FL 32223, USA"},
    {"label": "Atlantic Blvd 9498", "address": "9498 Atlantic Blvd, Jacksonville, FL 32225, USA"},
    {"label": "Old Saint Augustine Rd", "address": "11498 Old Saint Augustine Rd, Jacksonville, FL 32258, USA"},
    {"label": "College Dr Workcenter", "address": "74 College Dr, Orange Park, FL 32065, USA"},
    {"label": "Town Center Blvd", "address": "2000 Town Center Blvd, Orange Park, FL 32003, USA"},
    {"label": "Reed St", "address": "1929 Reed St, Orange Park, FL 32073, USA"},
    {"label": "Main St Palatka", "address": "319 Main St, Palatka, FL 32177, USA"},
    {"label": "A1A N", "address": "637 A1A N, Ponte Vedra Beach, FL 32082, USA"},
    {"label": "State Rd 312", "address": "256 State Rd 312, St. Augustine, FL 32086, USA"},
    {"label": "Cordova St", "address": "69 Cordova St, St. Augustine, FL 32084, USA"},
    {"label": "State Road 16", "address": "4875 State Road 16, St. Augustine, FL 32092, USA"},
)


@dataclass(frozen=True)
class PayActor:
    email: str
    display_name: str | None
    role: str
    can_view_all: bool
    can_edit_all: bool
    can_lock: bool
    is_guest: bool = False
    is_president: bool = False


@dataclass(frozen=True)
class DifferentialResult:
    wage_scale_id: int | None
    diff_rate: Decimal
    diff_amount: Decimal
    lost_wage_hourly_rate: Decimal
    presidential_hourly_rate: Decimal


@dataclass(frozen=True)
class CommissionCompensationResult:
    base_wage_input_type: str
    base_wage_amount: Decimal
    base_hourly_rate: Decimal
    commission_month_1_amount: Decimal
    commission_month_2_amount: Decimal
    commission_month_3_amount: Decimal
    commission_average_monthly: Decimal
    commission_hourly_rate: Decimal
    calculated_hourly_rate: Decimal


@dataclass(frozen=True)
class IrsMileageRateCandidate:
    rate_year: str
    effective_date: str
    cents_per_mile: Decimal
    rate_per_mile: Decimal
    source_url: str
    source_title: str | None = None


def normalize_email(value: object) -> str:
    return str(value or "").strip().lower()


def safe_filename(value: object, *, fallback: str = "file") -> str:
    cleaned = _SAFE_FILENAME_RE.sub("", str(value or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    return cleaned or fallback


def _money(value: object) -> Decimal:
    try:
        parsed = Decimal(str(value or "0").replace("$", "").replace(",", "").strip())
    except Exception:
        return Decimal("0")
    return parsed.quantize(_CURRENCY, rounding=ROUND_HALF_UP)


def _quantity(value: object) -> Decimal:
    try:
        parsed = Decimal(str(value or "0").replace(",", "").strip())
    except Exception:
        return Decimal("0")
    return parsed.quantize(Decimal("0.0001"), rounding=ROUND_HALF_UP)


def _mileage_rate(value: object) -> Decimal:
    try:
        parsed = Decimal(str(value or "0").replace("$", "").replace(",", "").strip())
    except Exception:
        return Decimal("0")
    return parsed.quantize(_IRS_RATE_QUANT, rounding=ROUND_HALF_UP)


def normalize_wage_input(
    *,
    input_type: object,
    amount: object,
    weekly_basis_hours: object,
) -> tuple[str, Decimal, Decimal]:
    normalized_type = str(input_type or "hourly").strip().lower()
    if normalized_type not in {"hourly", "weekly"}:
        normalized_type = "hourly"
    wage_amount = _money(amount)
    basis = _quantity(weekly_basis_hours)
    if basis <= 0:
        basis = Decimal("40")
    if wage_amount <= 0:
        return normalized_type, Decimal("0.00"), Decimal("0.00")
    if normalized_type == "weekly":
        hourly = (wage_amount / basis).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
        return normalized_type, wage_amount, hourly
    return normalized_type, wage_amount, wage_amount.quantize(_CURRENCY, rounding=ROUND_HALF_UP)


def calculate_commission_compensation(
    *,
    base_wage_input_type: object,
    base_wage_amount: object,
    weekly_basis_hours: object,
    commission_month_1_amount: object,
    commission_month_2_amount: object,
    commission_month_3_amount: object,
) -> CommissionCompensationResult:
    wage_type, wage_amount, base_hourly = normalize_wage_input(
        input_type=base_wage_input_type,
        amount=base_wage_amount,
        weekly_basis_hours=weekly_basis_hours,
    )
    month_1 = max(_money(commission_month_1_amount), Decimal("0.00"))
    month_2 = max(_money(commission_month_2_amount), Decimal("0.00"))
    month_3 = max(_money(commission_month_3_amount), Decimal("0.00"))
    average_monthly = ((month_1 + month_2 + month_3) / Decimal("3")).quantize(
        _CURRENCY,
        rounding=ROUND_HALF_UP,
    )
    commission_hourly = (average_monthly / _COMMISSION_HOUR_DIVISOR).quantize(
        _CURRENCY,
        rounding=ROUND_HALF_UP,
    )
    return CommissionCompensationResult(
        base_wage_input_type=wage_type,
        base_wage_amount=wage_amount,
        base_hourly_rate=base_hourly,
        commission_month_1_amount=month_1,
        commission_month_2_amount=month_2,
        commission_month_3_amount=month_3,
        commission_average_monthly=average_monthly,
        commission_hourly_rate=commission_hourly,
        calculated_hourly_rate=(base_hourly + commission_hourly).quantize(
            _CURRENCY,
            rounding=ROUND_HALF_UP,
        ),
    )


def required_commission_payroll_month(entry_date: date | str) -> str:
    parsed = date.fromisoformat(str(entry_date)) if not isinstance(entry_date, date) else entry_date
    previous_month_day = parsed.replace(day=1) - timedelta(days=1)
    return previous_month_day.strftime("%Y-%m")


def normalize_payroll_month(value: object, *, default_for: date | str | None = None) -> str:
    raw = str(value or "").strip()
    if not raw and default_for is not None:
        return required_commission_payroll_month(default_for)
    if not _PAYROLL_MONTH_RE.match(raw):
        raise ValueError("payroll_month must be YYYY-MM")
    return raw


def normalize_pay_basis(value: object) -> str:
    basis = str(value or "expense_only").strip().lower()
    if basis not in _PAY_PROFILE_BASIS_VALUES:
        raise ValueError("pay_basis must be hourly, weekly, commission, president, or expense_only")
    return basis


def normalize_pay_profile_status(value: object) -> str:
    status = str(value or "active").strip().lower()
    if status not in _PAY_PROFILE_STATUS_VALUES:
        raise ValueError("status must be active or disabled")
    return status


def calculate_pay_profile_snapshot(
    *,
    pay_basis: object,
    base_wage_input_type: object,
    base_wage_amount: object,
    weekly_basis_hours: object,
    commission_month_1_amount: object = 0,
    commission_month_2_amount: object = 0,
    commission_month_3_amount: object = 0,
) -> dict[str, Decimal | str]:
    basis = normalize_pay_basis(pay_basis)
    weekly_basis = _quantity(weekly_basis_hours)
    if weekly_basis <= 0:
        weekly_basis = Decimal("40")

    if basis == "expense_only":
        return {
            "pay_basis": basis,
            "base_wage_input_type": "hourly",
            "base_wage_amount": Decimal("0.00"),
            "weekly_basis_hours": weekly_basis,
            "commission_month_1_amount": Decimal("0.00"),
            "commission_month_2_amount": Decimal("0.00"),
            "commission_month_3_amount": Decimal("0.00"),
            "commission_average_monthly": Decimal("0.00"),
            "commission_hourly_rate": Decimal("0.00"),
            "calculated_hourly_rate": Decimal("0.00"),
        }

    if basis in {"hourly", "weekly"}:
        wage_type, wage_amount, hourly_rate = normalize_wage_input(
            input_type=base_wage_input_type,
            amount=base_wage_amount,
            weekly_basis_hours=weekly_basis,
        )
        commission = CommissionCompensationResult(
            base_wage_input_type=wage_type,
            base_wage_amount=wage_amount,
            base_hourly_rate=hourly_rate,
            commission_month_1_amount=Decimal("0.00"),
            commission_month_2_amount=Decimal("0.00"),
            commission_month_3_amount=Decimal("0.00"),
            commission_average_monthly=Decimal("0.00"),
            commission_hourly_rate=Decimal("0.00"),
            calculated_hourly_rate=hourly_rate,
        )
    elif basis == "president":
        commission = calculate_commission_compensation(
            base_wage_input_type=base_wage_input_type,
            base_wage_amount=base_wage_amount,
            weekly_basis_hours=weekly_basis,
            commission_month_1_amount=0,
            commission_month_2_amount=0,
            commission_month_3_amount=0,
        )
        commission = CommissionCompensationResult(
            base_wage_input_type=commission.base_wage_input_type,
            base_wage_amount=commission.base_wage_amount,
            base_hourly_rate=commission.base_hourly_rate,
            commission_month_1_amount=Decimal("0.00"),
            commission_month_2_amount=Decimal("0.00"),
            commission_month_3_amount=Decimal("0.00"),
            commission_average_monthly=Decimal("0.00"),
            commission_hourly_rate=Decimal("0.00"),
            calculated_hourly_rate=commission.base_hourly_rate,
        )
    else:
        commission = calculate_commission_compensation(
            base_wage_input_type=base_wage_input_type,
            base_wage_amount=base_wage_amount,
            weekly_basis_hours=weekly_basis,
            commission_month_1_amount=commission_month_1_amount,
            commission_month_2_amount=commission_month_2_amount,
            commission_month_3_amount=commission_month_3_amount,
        )

    return {
        "pay_basis": basis,
        "base_wage_input_type": commission.base_wage_input_type,
        "base_wage_amount": commission.base_wage_amount,
        "weekly_basis_hours": weekly_basis,
        "commission_month_1_amount": commission.commission_month_1_amount,
        "commission_month_2_amount": commission.commission_month_2_amount,
        "commission_month_3_amount": commission.commission_month_3_amount,
        "commission_average_monthly": commission.commission_average_monthly,
        "commission_hourly_rate": commission.commission_hourly_rate,
        "calculated_hourly_rate": commission.calculated_hourly_rate,
    }


def _currency_text(value: object) -> str:
    amount = _money(value)
    return "" if amount == 0 else f"{amount:.2f}"


def _currency_total_text(value: object) -> str:
    return f"$ {_money(value):.2f}"


def _decimal_text(value: object, *, places: int = 2) -> str:
    try:
        amount = Decimal(str(value or "0"))
    except Exception:
        amount = Decimal("0")
    quant = Decimal(1).scaleb(-places)
    return f"{amount.quantize(quant, rounding=ROUND_HALF_UP):f}"


def _mileage_rate_text(value: object) -> str:
    try:
        amount = Decimal(str(value or "0"))
    except Exception:
        amount = Decimal("0")
    text = f"{amount.quantize(_IRS_RATE_QUANT, rounding=ROUND_HALF_UP):f}".rstrip("0").rstrip(".")
    if "." not in text:
        return f"{text}.00"
    decimals = len(text.rsplit(".", 1)[1])
    if decimals == 1:
        return f"{text}0"
    return text


def _plain_text_from_html(value: str) -> str:
    text = html.unescape(str(value or ""))
    text = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", text)
    text = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _title_from_html(value: str) -> str | None:
    match = re.search(r"(?is)<title[^>]*>(.*?)</title>", str(value or ""))
    if not match:
        return None
    title = _plain_text_from_html(match.group(1))
    return title or None


def parse_irs_mileage_rate_candidates(*, content: str, source_url: str) -> list[IrsMileageRateCandidate]:
    title = _title_from_html(content)
    text = _plain_text_from_html(content)
    patterns = (
        r"IRS\s+sets\s+(20\d{2})\s+business\s+standard\s+mileage\s+rate\s+at\s+(\d+(?:\.\d+)?)\s+cents",
        r"Beginning\s+Jan(?:uary)?\.?\s+1,\s*(20\d{2}).{0,800}?(\d+(?:\.\d+)?)\s+cents?\s+per\s+mile\s+driven\s+for\s+business\s+use",
        r"For\s+(20\d{2}).{0,240}?standard\s+mileage\s+rate\s+is\s+(\d+(?:\.\d+)?)\s+cents?\s+per\s+mile",
        r"(20\d{2})\s+mileage\s+rates?.{0,500}?Self-employed\s+and\s+business:\s+(\d+(?:\.\d+)?)\s+cents/?mile",
    )
    candidates: list[IrsMileageRateCandidate] = []
    seen: set[tuple[str, str]] = set()
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            year = str(match.group(1))
            cents = Decimal(str(match.group(2))).quantize(Decimal("0.1"), rounding=ROUND_HALF_UP)
            key = (year, f"{cents:f}")
            if key in seen:
                continue
            seen.add(key)
            rate = (cents / Decimal("100")).quantize(_IRS_RATE_QUANT, rounding=ROUND_HALF_UP)
            candidates.append(
                IrsMileageRateCandidate(
                    rate_year=year,
                    effective_date=f"{year}-01-01",
                    cents_per_mile=cents,
                    rate_per_mile=rate,
                    source_url=source_url,
                    source_title=title,
                )
            )
    return candidates


def period_bounds_for(day: date) -> tuple[date, date]:
    delta_days = (day - _ANCHOR_PERIOD_START).days
    period_index = delta_days // 14
    start = _ANCHOR_PERIOD_START + timedelta(days=period_index * 14)
    return start, start + timedelta(days=13)


def current_period_bounds() -> tuple[date, date]:
    return period_bounds_for(date.today())


def period_id_for(start: date, end: date, revision: int = 1) -> str:
    return f"pay-{start.isoformat()}-{end.isoformat()}-r{max(1, int(revision))}"


def pay_period_folder_path(*, root_folder: str, period_start: str, period_end: str) -> str:
    year = str(period_start)[:4]
    label = f"{period_start}_to_{period_end}"
    return "/".join(part.strip("/") for part in (root_folder, year, label) if part.strip("/"))


def mileage_legacy_report_folder_path(*, root_folder: str, report_date: str) -> str:
    if not str(root_folder or "").strip():
        return ""
    parsed = datetime.strptime(str(report_date), "%Y-%m-%d").date()
    month_folder = f"{parsed.month:02d} - {parsed.strftime('%B')}"
    return "/".join(
        part.strip("/")
        for part in (root_folder, str(parsed.year), month_folder)
        if part.strip("/")
    )


async def add_pay_event(
    db: Db,
    *,
    period_id: str | None,
    event_type: str,
    actor: str | None = None,
    entry_id: str | None = None,
    packet_id: str | None = None,
    details: dict[str, object] | None = None,
) -> None:
    await db.exec(
        """INSERT INTO pay_events(period_id, entry_id, packet_id, ts_utc, event_type, actor, details_json)
           VALUES(?,?,?,?,?,?,?)""",
        (
            period_id,
            entry_id,
            packet_id,
            utcnow(),
            event_type,
            actor,
            json.dumps(details or {}, ensure_ascii=False),
        ),
    )


async def ensure_pay_period(db: Db, *, for_date: date | None = None) -> dict[str, object]:
    start, end = period_bounds_for(for_date or date.today())
    row = await db.fetchone(
        """SELECT id, period_start, period_end, status, revision, president_email,
                  sharepoint_folder_path, sharepoint_folder_web_url
           FROM pay_periods
           WHERE period_start=? AND period_end=?
           ORDER BY revision DESC
           LIMIT 1""",
        (start.isoformat(), end.isoformat()),
    )
    if row:
        return {
            "id": row[0],
            "period_start": row[1],
            "period_end": row[2],
            "status": row[3],
            "revision": int(row[4] or 1),
            "president_email": row[5],
            "sharepoint_folder_path": row[6],
            "sharepoint_folder_web_url": row[7],
        }

    now = utcnow()
    period_id = period_id_for(start, end)
    await db.exec(
        """INSERT INTO pay_periods(
             id, period_start, period_end, status, revision, created_at_utc, updated_at_utc
           ) VALUES(?,?,?,?,?,?,?)""",
        (period_id, start.isoformat(), end.isoformat(), "open", 1, now, now),
    )
    return {
        "id": period_id,
        "period_start": start.isoformat(),
        "period_end": end.isoformat(),
        "status": "open",
        "revision": 1,
        "president_email": None,
        "sharepoint_folder_path": None,
        "sharepoint_folder_web_url": None,
    }


async def get_pay_period(db: Db, period_id: str) -> dict[str, object] | None:
    row = await db.fetchone(
        """SELECT id, period_start, period_end, status, revision, locked_by, locked_at_utc,
                  completed_at_utc, president_email, sharepoint_folder_path, sharepoint_folder_web_url
           FROM pay_periods
           WHERE id=?""",
        (period_id,),
    )
    if not row:
        return None
    return {
        "id": row[0],
        "period_start": row[1],
        "period_end": row[2],
        "status": row[3],
        "revision": int(row[4] or 1),
        "locked_by": row[5],
        "locked_at_utc": row[6],
        "completed_at_utc": row[7],
        "president_email": row[8],
        "sharepoint_folder_path": row[9],
        "sharepoint_folder_web_url": row[10],
    }


def _clean_place_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def _normalize_common_place_rows(value: object) -> list[dict[str, str]]:
    if isinstance(value, dict):
        for key in ("common_places", "places", "addresses", "locations"):
            nested = value.get(key)
            if nested is not None:
                return _normalize_common_place_rows(nested)
        value = [{"label": key, "address": address} for key, address in value.items()]

    if not isinstance(value, list | tuple):
        return []

    rows: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for item in value:
        label = ""
        address = ""
        if isinstance(item, dict):
            lowered = {str(key or "").strip().lower(): raw_value for key, raw_value in item.items()}
            label = _clean_place_text(
                lowered.get("label")
                or lowered.get("name")
                or lowered.get("place")
                or lowered.get("title")
                or lowered.get("site")
            )
            address = _clean_place_text(
                lowered.get("address")
                or lowered.get("full_address")
                or lowered.get("location")
                or lowered.get("value")
            )
        else:
            address = _clean_place_text(item)
            label = address
        if address and not label:
            label = address
        if not label or not address:
            continue
        key = (label.lower(), address.lower())
        if key in seen:
            continue
        seen.add(key)
        rows.append({"label": label, "address": address})
    return rows


def _parse_common_places_csv(text: str) -> list[dict[str, str]]:
    lines = [line for line in text.splitlines() if line.strip()]
    if not lines:
        return []

    first_cells = [cell.strip().lower() for cell in next(csv.reader([lines[0]]))]
    has_header = any(cell in {"label", "name", "place", "title", "address", "location"} for cell in first_cells)
    rows: list[dict[str, str]] = []
    if has_header:
        reader = csv.DictReader(io.StringIO(text))
        for row in reader:
            rows.append(dict(row))
        return _normalize_common_place_rows(rows)

    for row in csv.reader(io.StringIO(text)):
        cleaned = [_clean_place_text(cell) for cell in row if _clean_place_text(cell)]
        if not cleaned:
            continue
        if len(cleaned) == 1:
            rows.append({"label": cleaned[0], "address": cleaned[0]})
        else:
            rows.append({"label": cleaned[0], "address": ", ".join(cleaned[1:])})
    return _normalize_common_place_rows(rows)


def _parse_common_places_text(text: str) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for raw_line in text.splitlines():
        line = _clean_place_text(raw_line)
        if not line or line.startswith("#"):
            continue
        label = ""
        address = line
        for sep in ("\t", "|", "="):
            if sep in line:
                label, address = [_clean_place_text(part) for part in line.split(sep, 1)]
                break
        if not label:
            label = address.split(",", 1)[0].strip() or address
        rows.append({"label": label, "address": address})
    return _normalize_common_place_rows(rows)


def parse_common_places_config(*, content: bytes | str, filename: str = "") -> list[dict[str, str]]:
    if isinstance(content, bytes):
        text = content.decode("utf-8-sig", errors="replace")
    else:
        text = str(content or "").lstrip("\ufeff")
    if not text.strip():
        return []

    suffix = Path(filename or "").suffix.lower()
    try:
        if suffix == ".json":
            return _normalize_common_place_rows(json.loads(text))
        if suffix in {".yaml", ".yml"}:
            import yaml

            return _normalize_common_place_rows(yaml.safe_load(text))
        if suffix == ".csv":
            return _parse_common_places_csv(text)
        if suffix == ".txt":
            return _parse_common_places_text(text)
    except Exception:
        return []

    for parser in (
        lambda value: _normalize_common_place_rows(json.loads(value)),
        lambda value: _parse_common_places_csv(value),
        _parse_common_places_text,
    ):
        try:
            rows = parser(text)
        except Exception:
            rows = []
        if rows:
            return rows
    return []


def merge_common_places(*sources: object) -> list[dict[str, str]]:
    merged: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for source in sources:
        for place in _normalize_common_place_rows(source):
            key = (place["label"].lower(), place["address"].lower())
            if key in seen:
                continue
            seen.add(key)
            merged.append(place)
    return merged


def common_places_cache_path(*, data_root: str) -> Path:
    return Path(data_root) / "pay" / "common_places.json"


def load_common_places_cache(*, data_root: str) -> list[dict[str, str]]:
    path = common_places_cache_path(data_root=data_root)
    if not path.exists():
        return []
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return []
    if isinstance(payload, dict):
        return _normalize_common_place_rows(payload.get("common_places"))
    return _normalize_common_place_rows(payload)


def write_common_places_cache(*, data_root: str, places: object, source: str = "sharepoint") -> Path:
    normalized = merge_common_places(places)
    path = common_places_cache_path(data_root=data_root)
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "source": source,
        "updated_at_utc": utcnow(),
        "common_places": normalized,
    }
    tmp_path = path.with_suffix(f"{path.suffix}.tmp")
    tmp_path.write_text(json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8")
    tmp_path.replace(path)
    return path


def _pay_graph_site(*, cfg: Any | None = None, graph_cfg: Any | None = None, pay_cfg: Any | None = None) -> tuple[str, str]:
    if cfg is not None:
        graph_cfg = getattr(cfg, "graph", graph_cfg)
        pay_cfg = getattr(cfg, "pay_portal", pay_cfg)
    hostname = str(getattr(pay_cfg, "sharepoint_site_hostname", "") or "").strip()
    site_path = str(getattr(pay_cfg, "sharepoint_site_path", "") or "").strip()
    if not hostname:
        hostname = str(getattr(graph_cfg, "site_hostname", "") or "").strip()
    if not site_path:
        site_path = str(getattr(graph_cfg, "site_path", "") or "").strip()
    return hostname, site_path


def _pay_docx_pdf_graph_temp_folder(*, cfg: Any) -> str:
    folder = str(getattr(getattr(cfg, "pay_portal", None), "docx_pdf_graph_temp_folder", "") or "").strip()
    return folder or str(getattr(cfg, "docx_pdf_graph_temp_folder", "") or "").strip()


def _sharepoint_place_config_targets(*, graph_cfg: Any, pay_cfg: Any) -> list[tuple[str, str]]:
    raw_folder_path = str(getattr(pay_cfg, "common_places_sharepoint_folder", "") or "").strip()
    if not raw_folder_path:
        return []

    default_library = str(getattr(graph_cfg, "document_library", "Documents") or "Documents").strip()
    configured_library = str(getattr(pay_cfg, "common_places_sharepoint_library", "") or "").strip()
    path_parts = [part.strip() for part in raw_folder_path.replace("\\", "/").split("/") if part.strip()]

    inferred_library = ""
    inferred_folder = ""
    for index, part in enumerate(path_parts):
        lowered = part.lower()
        if lowered == "documents" or lowered.endswith(" - documents"):
            inferred_library = part
            inferred_folder = "/".join(path_parts[index + 1 :])
            break

    if not inferred_folder:
        inferred_folder = "/".join(path_parts)

    libraries = [
        configured_library,
        inferred_library,
        default_library,
        "Officer Eboard committee - Documents",
    ]
    targets: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for library in libraries:
        clean_library = library.strip()
        clean_folder = inferred_folder.strip("/")
        if not clean_library or not clean_folder:
            continue
        key = (clean_library.lower(), clean_folder.lower())
        if key in seen:
            continue
        seen.add(key)
        targets.append((clean_library, clean_folder))
    return targets


def load_sharepoint_common_places(
    *,
    graph: Any,
    graph_cfg: Any,
    pay_cfg: Any,
) -> list[dict[str, str]]:
    targets = _sharepoint_place_config_targets(graph_cfg=graph_cfg, pay_cfg=pay_cfg)
    if not targets or graph is None:
        return []
    list_files = getattr(graph, "list_files_in_folder_path", None)
    download = getattr(graph, "download_item_bytes", None)
    if not callable(list_files) or not callable(download):
        return []

    rows: list[dict[str, str]] = []
    last_error: Exception | None = None
    for library, folder_path in targets:
        try:
            site_hostname, site_path = _pay_graph_site(graph_cfg=graph_cfg, pay_cfg=pay_cfg)
            files = list_files(
                site_hostname=site_hostname,
                site_path=site_path,
                library=library,
                folder_path=folder_path,
                recursive=True,
            )
        except Exception as exc:
            last_error = exc
            continue
        for item in files:
            name = str(getattr(item, "name", "") or "").strip()
            if Path(name).suffix.lower() not in _COMMON_PLACE_CONFIG_EXTENSIONS:
                continue
            data = download(
                drive_id=str(getattr(item, "drive_id", "") or ""),
                item_id=str(getattr(item, "item_id", "") or ""),
            )
            rows.extend(parse_common_places_config(content=data, filename=name))
        merged = merge_common_places(rows)
        if merged:
            return merged
    if last_error is not None:
        raise last_error
    return []


async def pay_settings(db: Db, *, pay_cfg: Any | None = None) -> dict[str, object]:
    default_common_places = merge_common_places(
        getattr(pay_cfg, "common_places", ()) or (),
        _DEFAULT_COMMON_PLACES,
    )
    defaults: dict[str, object] = {
        "treasurer_emails": list(getattr(pay_cfg, "treasurer_emails", ()) or ()),
        "president_email": str(getattr(pay_cfg, "president_email", "") or ""),
        "irs_rates": dict(getattr(pay_cfg, "irs_rates", {}) or {}),
        "fund_fica_rate": str(getattr(pay_cfg, "fund_fica_rate", _FUND_FICA_RATE) or _FUND_FICA_RATE),
        "common_places": default_common_places,
        "common_places_managed": False,
    }
    row = await db.app_setting("pay_portal")
    if not row:
        return defaults
    try:
        parsed = json.loads(row[0])
    except Exception:
        parsed = {}
    if not isinstance(parsed, dict):
        return defaults
    merged = dict(defaults)
    for key, value in parsed.items():
        if key == "irs_rates" and not value:
            continue
        if key == "common_places":
            if parsed.get("common_places_managed"):
                merged["common_places"] = merge_common_places(value)
            else:
                merged["common_places"] = merge_common_places(value, default_common_places)
            continue
        merged[key] = value
    return merged


async def save_pay_settings(
    db: Db,
    *,
    setting: dict[str, object],
    updated_by: str | None,
    pay_cfg: Any | None = None,
) -> dict[str, object]:
    normalized = dict(await pay_settings(db, pay_cfg=pay_cfg))
    if "common_places" in setting:
        setting = {
            **setting,
            "common_places": merge_common_places(setting.get("common_places")),
            "common_places_managed": True,
        }
    normalized.update(setting)
    await db.upsert_app_setting(setting_key="pay_portal", setting=normalized, updated_by=updated_by)
    return normalized


def normalize_pay_fund_type(value: object) -> str:
    normalized = str(value or "sif").strip().lower()
    if normalized in {"growth_fund", "growth fund"}:
        normalized = "growth"
    if normalized not in _PAY_FUND_TYPES:
        raise ValueError("fund_type must be sif or growth")
    return normalized


def normalize_pay_fund_status(value: object) -> str:
    normalized = str(value or "active").strip().lower()
    if normalized not in _PAY_FUND_STATUS_VALUES:
        raise ValueError("fund status must be active or closed")
    return normalized


def normalize_pay_fund_ledger_type(value: object) -> str:
    normalized = str(value or "advance").strip().lower()
    if normalized not in _PAY_FUND_LEDGER_TYPES:
        raise ValueError("ledger_type must be advance, reimbursement_submitted, reimbursement_received, or adjustment")
    return normalized


def fund_fica_rate_from_settings(settings: dict[str, object] | None) -> Decimal:
    raw = (settings or {}).get("fund_fica_rate", _FUND_FICA_RATE)
    rate = _quantity(raw)
    if rate < 0:
        return _FUND_FICA_RATE
    return rate.quantize(Decimal("0.0001"), rounding=ROUND_HALF_UP)


def _fund_ledger_signed_amount(ledger_type: object, amount: object) -> Decimal:
    kind = normalize_pay_fund_ledger_type(ledger_type)
    value = _money(amount)
    if kind == "reimbursement_submitted":
        return Decimal("0.00")
    return value


def _fund_row_from_db(row: Any) -> dict[str, object]:
    return {
        "id": row[0],
        "fund_type": row[1],
        "name": row[2],
        "status": row[3],
        "local_number": row[4],
        "description": row[5],
        "created_by": row[6],
        "created_at_utc": row[7],
        "updated_at_utc": row[8],
        "updated_by": row[9],
    }


async def pay_fund_by_id(db: Db, *, fund_id: str) -> dict[str, object] | None:
    row = await db.fetchone(
        """SELECT id, fund_type, name, status, local_number, description,
                  created_by, created_at_utc, updated_at_utc, updated_by
           FROM pay_funds
           WHERE id=?""",
        (fund_id,),
    )
    return _fund_row_from_db(row) if row else None


async def upsert_pay_fund(
    db: Db,
    *,
    fund_id: str | None = None,
    fund_type: object = "sif",
    name: str,
    status: object = "active",
    local_number: str | None = "3106",
    description: str | None = None,
    actor_email: str,
) -> dict[str, object]:
    normalized_name = str(name or "").strip()
    if not normalized_name:
        raise ValueError("fund name is required")
    normalized_type = normalize_pay_fund_type(fund_type)
    normalized_status = normalize_pay_fund_status(status)
    normalized_local = str(local_number or "3106").strip() or "3106"
    now = utcnow()
    requested_id = str(fund_id or "").strip()
    existing_name = await db.fetchone("SELECT id FROM pay_funds WHERE name=?", (normalized_name,))
    if existing_name and requested_id and str(existing_name[0]) != requested_id:
        raise ValueError("fund name already exists")
    saved_id = requested_id or (str(existing_name[0]) if existing_name else f"pay-fund-{uuid4().hex}")
    await db.exec(
        """
        INSERT INTO pay_funds(
          id, fund_type, name, status, local_number, description,
          created_by, created_at_utc, updated_at_utc, updated_by
        ) VALUES(?,?,?,?,?,?,?,?,?,?)
        ON CONFLICT(id) DO UPDATE SET
          fund_type=excluded.fund_type,
          name=excluded.name,
          status=excluded.status,
          local_number=excluded.local_number,
          description=excluded.description,
          updated_at_utc=excluded.updated_at_utc,
          updated_by=excluded.updated_by
        """,
        (
            saved_id,
            normalized_type,
            normalized_name,
            normalized_status,
            normalized_local,
            str(description or "").strip() or None,
            actor_email,
            now,
            now,
            actor_email,
        ),
    )
    fund = await pay_fund_by_id(db, fund_id=saved_id)
    if not fund:
        raise RuntimeError("failed to save fund")
    return fund


async def add_pay_fund_ledger_entry(
    db: Db,
    *,
    fund_id: str,
    ledger_type: object,
    amount: object,
    effective_date: str,
    reference: str | None = None,
    notes: str | None = None,
    actor_email: str,
) -> dict[str, object]:
    fund = await pay_fund_by_id(db, fund_id=fund_id)
    if not fund:
        raise ValueError("fund not found")
    normalized_type = normalize_pay_fund_ledger_type(ledger_type)
    parsed_date = date.fromisoformat(str(effective_date or "").strip())
    value = _money(amount)
    if normalized_type != "adjustment" and value < 0:
        raise ValueError("ledger amount cannot be negative")
    ledger_id = f"pay-fund-ledger-{uuid4().hex}"
    now = utcnow()
    await db.exec(
        """INSERT INTO pay_fund_ledger_entries(
             id, fund_id, ledger_type, amount, effective_date, reference, notes, created_by, created_at_utc
           ) VALUES(?,?,?,?,?,?,?,?,?)""",
        (
            ledger_id,
            fund_id,
            normalized_type,
            float(value),
            parsed_date.isoformat(),
            str(reference or "").strip() or None,
            str(notes or "").strip() or None,
            actor_email,
            now,
        ),
    )
    return {
        "id": ledger_id,
        "fund_id": fund_id,
        "ledger_type": normalized_type,
        "amount": float(value),
        "effective_date": parsed_date.isoformat(),
        "reference": str(reference or "").strip() or None,
        "notes": str(notes or "").strip() or None,
        "created_by": actor_email,
        "created_at_utc": now,
    }


def _fund_allocation_cost(*, allocation: dict[str, object], entry: dict[str, object], fica_rate: Decimal) -> dict[str, Decimal]:
    hourly_rate = _money(entry.get("lost_wage_hourly_rate") or entry.get("hourly_rate"))
    hours = _quantity(allocation.get("hours"))
    gross = (hours * hourly_rate).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
    fica = (gross * fica_rate).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
    mileage = _money(allocation.get("mileage_amount"))
    rentals = _money(allocation.get("rentals_amount"))
    meals = _money(allocation.get("meals_amount"))
    hotel = _money(allocation.get("hotel_amount"))
    misc = _money(allocation.get("miscellaneous_amount"))
    expenses = (mileage + rentals + meals + hotel + misc).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
    total = (gross + fica + expenses).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
    return {
        "lost_wage_gross": gross,
        "fica": fica,
        "mileage": mileage,
        "rentals": rentals,
        "meals": meals,
        "hotel": hotel,
        "miscellaneous": misc,
        "expenses": expenses,
        "total": total,
    }


def _allocation_from_row(row: Any) -> dict[str, object]:
    return {
        "id": row[0],
        "period_id": row[1],
        "entry_id": row[2],
        "fund_id": row[3],
        "fund_name": row[4],
        "fund_type": row[5],
        "entry_date": row[6],
        "user_email": row[7],
        "display_name": row[8],
        "hours": row[9],
        "mileage_miles": row[10],
        "mileage_amount": row[11],
        "rentals_amount": row[12],
        "meals_amount": row[13],
        "hotel_amount": row[14],
        "miscellaneous_amount": row[15],
        "notes": row[16],
        "created_by": row[17],
        "created_at_utc": row[18],
        "updated_at_utc": row[19],
    }


async def list_pay_fund_allocations(
    db: Db,
    *,
    actor: PayActor,
    period_id: str | None = None,
) -> list[dict[str, object]]:
    where = ["1=1"]
    params: list[object] = []
    if period_id:
        where.append("fa.period_id=?")
        params.append(period_id)
    if not actor.can_view_all:
        where.append("e.user_email=?")
        params.append(actor.email)
    rows = await db.fetchall(
        f"""SELECT fa.id, fa.period_id, fa.entry_id, fa.fund_id, f.name, f.fund_type,
                  e.entry_date, e.user_email, e.display_name,
                  fa.hours, fa.mileage_miles, fa.mileage_amount, fa.rentals_amount,
                  fa.meals_amount, fa.hotel_amount, fa.miscellaneous_amount, fa.notes,
                  fa.created_by, fa.created_at_utc, fa.updated_at_utc
           FROM pay_fund_allocations fa
           JOIN pay_entries e ON e.id = fa.entry_id
           JOIN pay_funds f ON f.id = fa.fund_id
           WHERE {' AND '.join(where)}
           ORDER BY e.entry_date, e.user_email, f.name""",
        tuple(params),
    )
    return [_allocation_from_row(row) for row in rows]


async def save_pay_fund_allocations_for_entry(
    db: Db,
    *,
    entry_id: str,
    actor: PayActor,
    allocations: list[dict[str, object]],
) -> list[dict[str, object]]:
    row = await db.fetchone(
        """SELECT e.period_id, e.user_email, e.entry_date, e.locked_at_utc, p.status,
                  e.hours, e.mileage_miles, e.mileage_amount, e.rentals_amount,
                  e.meals_amount, e.hotel_amount, e.miscellaneous_amount
           FROM pay_entries e
           JOIN pay_periods p ON p.id = e.period_id
           WHERE e.id=?""",
        (entry_id,),
    )
    if not row:
        raise ValueError("entry not found")
    if row[3] or str(row[4] or "") != "open":
        raise ValueError("entry is locked")
    owner_email = normalize_email(row[1])
    if not (actor.can_edit_all or actor.can_lock or owner_email == actor.email):
        raise PermissionError("cannot allocate another user's entry")

    period_id = str(row[0])
    parent_totals = {
        "hours": _quantity(row[5]),
        "mileage_miles": _quantity(row[6]),
        "mileage_amount": _money(row[7]),
        "rentals_amount": _money(row[8]),
        "meals_amount": _money(row[9]),
        "hotel_amount": _money(row[10]),
        "miscellaneous_amount": _money(row[11]),
    }
    seen_funds: set[str] = set()
    normalized_rows: list[dict[str, object]] = []
    totals = {key: Decimal("0.00") for key in parent_totals}
    for raw in allocations:
        fund_id = str(raw.get("fund_id") or "").strip()
        if not fund_id:
            continue
        if fund_id in seen_funds:
            raise ValueError("each fund can only appear once per entry")
        seen_funds.add(fund_id)
        fund = await pay_fund_by_id(db, fund_id=fund_id)
        if not fund or str(fund.get("status") or "") != "active":
            raise ValueError("active fund not found")
        normalized = {
            "fund_id": fund_id,
            "hours": _quantity(raw.get("hours")),
            "mileage_miles": _quantity(raw.get("mileage_miles")),
            "mileage_amount": _money(raw.get("mileage_amount")),
            "rentals_amount": _money(raw.get("rentals_amount")),
            "meals_amount": _money(raw.get("meals_amount")),
            "hotel_amount": _money(raw.get("hotel_amount")),
            "miscellaneous_amount": _money(raw.get("miscellaneous_amount")),
            "notes": str(raw.get("notes") or "").strip(),
        }
        for key, value in normalized.items():
            if key in totals and value < 0:  # type: ignore[operator]
                raise ValueError("fund allocation amounts cannot be negative")
        for key in totals:
            totals[key] += normalized[key]  # type: ignore[operator]
        normalized_rows.append(normalized)

    for key, total in totals.items():
        allowed = parent_totals[key]
        if key.endswith("amount") or key in {"rentals_amount", "meals_amount", "hotel_amount", "miscellaneous_amount"}:
            if total > allowed + Decimal("0.01"):
                raise ValueError(f"allocated {key} exceeds the entry total")
        elif total > allowed + Decimal("0.0001"):
            raise ValueError(f"allocated {key} exceeds the entry total")

    existing = await db.fetchall("SELECT id FROM pay_fund_allocations WHERE entry_id=?", (entry_id,))
    existing_ids = [str(item[0]) for item in existing]
    if existing_ids:
        placeholders = ",".join("?" for _ in existing_ids)
        await db.exec(
            f"DELETE FROM pay_fund_attachment_links WHERE allocation_id IN ({placeholders})",
            tuple(existing_ids),
        )
    await db.exec("DELETE FROM pay_fund_allocations WHERE entry_id=?", (entry_id,))
    now = utcnow()
    for normalized in normalized_rows:
        allocation_id = f"pay-fund-allocation-{uuid4().hex}"
        await db.exec(
            """INSERT INTO pay_fund_allocations(
                 id, period_id, entry_id, fund_id, hours, mileage_miles, mileage_amount,
                 rentals_amount, meals_amount, hotel_amount, miscellaneous_amount,
                 notes, created_by, created_at_utc, updated_at_utc
               ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                allocation_id,
                period_id,
                entry_id,
                normalized["fund_id"],
                float(normalized["hours"]),
                float(normalized["mileage_miles"]),
                float(normalized["mileage_amount"]),
                float(normalized["rentals_amount"]),
                float(normalized["meals_amount"]),
                float(normalized["hotel_amount"]),
                float(normalized["miscellaneous_amount"]),
                normalized["notes"] or None,
                actor.email,
                now,
                now,
            ),
        )
    await add_pay_event(
        db,
        period_id=period_id,
        entry_id=entry_id,
        event_type="fund_allocations_saved",
        actor=actor.email,
        details={"allocation_count": len(normalized_rows), "fund_ids": sorted(seen_funds)},
    )
    return await list_pay_fund_allocations(db, actor=actor, period_id=period_id)


async def link_pay_attachment_to_fund(
    db: Db,
    *,
    attachment_id: str,
    fund_id: str,
    actor: PayActor,
    allocation_id: str | None = None,
    notes: str | None = None,
) -> dict[str, object]:
    attachment = await attachment_for_actor(db, attachment_id=attachment_id, actor=actor)
    fund = await pay_fund_by_id(db, fund_id=fund_id)
    if not fund or str(fund.get("status") or "") != "active":
        raise ValueError("active fund not found")
    normalized_allocation_id = str(allocation_id or "").strip() or None
    if normalized_allocation_id:
        row = await db.fetchone(
            "SELECT fund_id, entry_id FROM pay_fund_allocations WHERE id=?",
            (normalized_allocation_id,),
        )
        if not row or str(row[0]) != fund_id or str(row[1]) != str(attachment.get("entry_id")):
            raise ValueError("allocation does not match this fund and attachment")
    link_id = f"pay-fund-link-{uuid4().hex}"
    now = utcnow()
    await db.exec(
        """INSERT INTO pay_fund_attachment_links(
             id, fund_id, allocation_id, attachment_id, notes, linked_by, created_at_utc
           ) VALUES(?,?,?,?,?,?,?)
           ON CONFLICT(fund_id, attachment_id) DO UPDATE SET
             allocation_id=excluded.allocation_id,
             notes=excluded.notes,
             linked_by=excluded.linked_by,
             created_at_utc=excluded.created_at_utc""",
        (
            link_id,
            fund_id,
            normalized_allocation_id,
            attachment_id,
            str(notes or "").strip() or None,
            actor.email,
            now,
        ),
    )
    await add_pay_event(
        db,
        period_id=str(attachment.get("period_id") or ""),
        entry_id=str(attachment.get("entry_id") or ""),
        event_type="fund_attachment_linked",
        actor=actor.email,
        details={"fund_id": fund_id, "attachment_id": attachment_id},
    )
    return {
        "id": link_id,
        "fund_id": fund_id,
        "allocation_id": normalized_allocation_id,
        "attachment_id": attachment_id,
        "linked_by": actor.email,
        "created_at_utc": now,
    }


async def list_pay_fund_attachment_links(
    db: Db,
    *,
    actor: PayActor,
    period_id: str | None = None,
) -> list[dict[str, object]]:
    where = ["a.removed_at_utc IS NULL"]
    params: list[object] = []
    if period_id:
        where.append("a.period_id=?")
        params.append(period_id)
    if not actor.can_view_all:
        where.append("e.user_email=?")
        params.append(actor.email)
    rows = await db.fetchall(
        f"""SELECT l.id, l.fund_id, f.name, l.allocation_id, l.attachment_id,
                  a.original_filename, a.attachment_type, e.id, e.user_email,
                  e.entry_date, l.notes, l.linked_by, l.created_at_utc
           FROM pay_fund_attachment_links l
           JOIN pay_funds f ON f.id = l.fund_id
           JOIN pay_attachments a ON a.id = l.attachment_id
           JOIN pay_entries e ON e.id = a.entry_id
           WHERE {' AND '.join(where)}
           ORDER BY e.entry_date, f.name, a.original_filename""",
        tuple(params),
    )
    return [
        {
            "id": row[0],
            "fund_id": row[1],
            "fund_name": row[2],
            "allocation_id": row[3],
            "attachment_id": row[4],
            "filename": row[5],
            "attachment_type": row[6],
            "entry_id": row[7],
            "user_email": row[8],
            "entry_date": row[9],
            "notes": row[10],
            "linked_by": row[11],
            "created_at_utc": row[12],
        }
        for row in rows
    ]


async def _fund_ledger_totals(db: Db, *, fund_id: str, through_date: str | None = None) -> dict[str, Decimal]:
    clause = "WHERE fund_id=?"
    params: list[object] = [fund_id]
    if through_date:
        clause += " AND effective_date<=?"
        params.append(through_date)
    rows = await db.fetchall(
        f"SELECT ledger_type, amount FROM pay_fund_ledger_entries {clause}",
        tuple(params),
    )
    totals = {kind: Decimal("0.00") for kind in _PAY_FUND_LEDGER_TYPES}
    for ledger_type, amount in rows:
        totals[normalize_pay_fund_ledger_type(ledger_type)] += _money(amount)
    totals["available"] = sum((_fund_ledger_signed_amount(kind, value) for kind, value in totals.items() if kind in _PAY_FUND_LEDGER_TYPES), Decimal("0.00"))
    return totals


async def _fund_allocated_total(
    db: Db,
    *,
    fund_id: str,
    fica_rate: Decimal,
    before_date: str | None = None,
    through_date: str | None = None,
) -> Decimal:
    where = ["fa.fund_id=?", "COALESCE(e.review_status, 'pending') IN ('pending', 'approved')"]
    params: list[object] = [fund_id]
    if before_date:
        where.append("e.entry_date<?")
        params.append(before_date)
    if through_date:
        where.append("e.entry_date<=?")
        params.append(through_date)
    rows = await db.fetchall(
        f"""SELECT fa.hours, fa.mileage_amount, fa.rentals_amount, fa.meals_amount,
                  fa.hotel_amount, fa.miscellaneous_amount, e.lost_wage_hourly_rate, e.hourly_rate
           FROM pay_fund_allocations fa
           JOIN pay_entries e ON e.id = fa.entry_id
           WHERE {' AND '.join(where)}""",
        tuple(params),
    )
    total = Decimal("0.00")
    for row in rows:
        allocation = {
            "hours": row[0],
            "mileage_amount": row[1],
            "rentals_amount": row[2],
            "meals_amount": row[3],
            "hotel_amount": row[4],
            "miscellaneous_amount": row[5],
        }
        entry = {"lost_wage_hourly_rate": row[6], "hourly_rate": row[7]}
        total += _fund_allocation_cost(allocation=allocation, entry=entry, fica_rate=fica_rate)["total"]
    return total.quantize(_CURRENCY, rounding=ROUND_HALF_UP)


async def list_pay_funds(
    db: Db,
    *,
    include_inactive: bool = False,
    include_financials: bool = True,
    fica_rate: Decimal = _FUND_FICA_RATE,
) -> list[dict[str, object]]:
    if include_inactive:
        rows = await db.fetchall(
            """SELECT id, fund_type, name, status, local_number, description,
                      created_by, created_at_utc, updated_at_utc, updated_by
               FROM pay_funds
               ORDER BY status, name"""
        )
    else:
        rows = await db.fetchall(
            """SELECT id, fund_type, name, status, local_number, description,
                      created_by, created_at_utc, updated_at_utc, updated_by
               FROM pay_funds
               WHERE status='active'
               ORDER BY name"""
        )
    funds = [_fund_row_from_db(row) for row in rows]
    if not include_financials:
        return funds
    for fund in funds:
        ledger = await _fund_ledger_totals(db, fund_id=str(fund["id"]))
        allocated = await _fund_allocated_total(db, fund_id=str(fund["id"]), fica_rate=fica_rate)
        available = _money(ledger.get("available"))
        submitted = _money(ledger.get("reimbursement_submitted"))
        fund.update(
            {
                "advance_amount": float(_money(ledger.get("advance"))),
                "reimbursement_submitted_amount": float(submitted),
                "reimbursement_received_amount": float(_money(ledger.get("reimbursement_received"))),
                "adjustment_amount": float(_money(ledger.get("adjustment"))),
                "allocated_amount": float(allocated),
                "remaining_balance": float((available - allocated).quantize(_CURRENCY, rounding=ROUND_HALF_UP)),
                "reimbursement_needed": float(max(allocated - submitted, Decimal("0.00"))),
            }
        )
    return funds


def _fund_packet_support_folder_name(fund: dict[str, object], period_start: str, period_end: str, packet_id: str) -> str:
    return f"Fund Packets/{safe_filename(fund.get('name'), fallback='fund')}/{period_start}_to_{period_end}_{packet_id[-8:]}"


async def _fund_packet_rows(
    db: Db,
    *,
    fund_id: str,
    period_start: str,
    period_end: str,
    fica_rate: Decimal,
) -> list[dict[str, object]]:
    rows = await db.fetchall(
        """SELECT fa.id, fa.entry_id, fa.hours, fa.mileage_miles, fa.mileage_amount,
                  fa.rentals_amount, fa.meals_amount, fa.hotel_amount, fa.miscellaneous_amount,
                  fa.notes, e.user_email, e.display_name, e.entry_date, e.local_number,
                  e.lost_wage_hourly_rate, e.hourly_rate, e.mileage_rate
           FROM pay_fund_allocations fa
           JOIN pay_entries e ON e.id = fa.entry_id
           WHERE fa.fund_id=?
             AND e.entry_date BETWEEN ? AND ?
             AND COALESCE(e.review_status, 'pending') IN ('pending', 'approved')
           ORDER BY e.entry_date, e.display_name, e.user_email""",
        (fund_id, period_start, period_end),
    )
    packet_rows: list[dict[str, object]] = []
    for row in rows:
        allocation = {
            "hours": row[2],
            "mileage_miles": row[3],
            "mileage_amount": row[4],
            "rentals_amount": row[5],
            "meals_amount": row[6],
            "hotel_amount": row[7],
            "miscellaneous_amount": row[8],
        }
        entry = {"lost_wage_hourly_rate": row[14], "hourly_rate": row[15]}
        costs = _fund_allocation_cost(allocation=allocation, entry=entry, fica_rate=fica_rate)
        packet_rows.append(
            {
                "allocation_id": row[0],
                "entry_id": row[1],
                "hours": float(_quantity(row[2])),
                "mileage_miles": float(_quantity(row[3])),
                "mileage_amount": float(_money(row[4])),
                "rentals_amount": float(_money(row[5])),
                "meals_amount": float(_money(row[6])),
                "hotel_amount": float(_money(row[7])),
                "miscellaneous_amount": float(_money(row[8])),
                "notes": row[9],
                "user_email": row[10],
                "display_name": row[11] or row[10],
                "entry_date": row[12],
                "local_number": row[13],
                "hourly_rate": float(_money(row[14] or row[15])),
                "mileage_rate": float(_mileage_rate(row[16])),
                "lost_wage_gross": float(costs["lost_wage_gross"]),
                "fica": float(costs["fica"]),
                "total": float(costs["total"]),
            }
        )
    return packet_rows


def _write_fund_packet_workbook(
    *,
    path: Path,
    fund: dict[str, object],
    period_start: str,
    period_end: str,
    beginning_balance: Decimal,
    rows: list[dict[str, object]],
    fica_rate: Decimal,
) -> tuple[Decimal, Decimal]:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
    except ImportError as exc:  # pragma: no cover - dependency is declared for the API image.
        raise RuntimeError("openpyxl is required for fund packet workbook generation") from exc

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet 1"
    ws["A1"] = f"LOCAL:      {fund.get('local_number') or '3106'}"
    ws["A1"].font = Font(bold=True)
    headers = [
        "Name",
        "Date ",
        "Hourly Rate",
        "Lost Wages (Gross) ",
        "FICA",
        "Total Wage (Wage + FICA) ",
        "Miles",
        "Mileage         ",
        "Hotel",
        "Meals",
        "Rental",
        "Misc Exp",
        "Misc",
        "Total Expense",
        "TOTAL Wage & Expense",
        "Remaining",
    ]
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=2, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="D9EAF7")
    ws["A3"] = f"Beginning {fund.get('name')} {period_start}"
    ws["P3"] = float(beginning_balance)
    for col in range(3, 16):
        ws.cell(row=3, column=col, value=0)
    current_balance = beginning_balance
    first_data_row = 4
    for offset, packet_row in enumerate(rows):
        excel_row = first_data_row + offset
        row_total = _money(packet_row.get("total"))
        current_balance = (current_balance - row_total).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
        ws.cell(row=excel_row, column=1, value=packet_row.get("display_name"))
        ws.cell(row=excel_row, column=2, value=date.fromisoformat(str(packet_row.get("entry_date"))))
        ws.cell(row=excel_row, column=3, value=float(_money(packet_row.get("hourly_rate"))))
        ws.cell(row=excel_row, column=4, value=float(_money(packet_row.get("lost_wage_gross"))))
        ws.cell(row=excel_row, column=5, value=f"=ROUND(D{excel_row}*{fica_rate},2)")
        ws.cell(row=excel_row, column=6, value=f"=SUM(D{excel_row}:E{excel_row})")
        ws.cell(row=excel_row, column=7, value=float(_quantity(packet_row.get("mileage_miles"))))
        mileage_rate = _mileage_rate(packet_row.get("mileage_rate"))
        if mileage_rate > 0 and _quantity(packet_row.get("mileage_miles")) > 0:
            ws.cell(row=excel_row, column=8, value=f"=ROUND(G{excel_row}*{mileage_rate},2)")
        else:
            ws.cell(row=excel_row, column=8, value=float(_money(packet_row.get("mileage_amount"))))
        ws.cell(row=excel_row, column=9, value=float(_money(packet_row.get("hotel_amount"))))
        ws.cell(row=excel_row, column=10, value=float(_money(packet_row.get("meals_amount"))))
        ws.cell(row=excel_row, column=11, value=float(_money(packet_row.get("rentals_amount"))))
        ws.cell(row=excel_row, column=12, value=0)
        ws.cell(row=excel_row, column=13, value=float(_money(packet_row.get("miscellaneous_amount"))))
        ws.cell(row=excel_row, column=14, value=f"=SUM(H{excel_row}:M{excel_row})")
        ws.cell(row=excel_row, column=15, value=f"=SUM(F{excel_row},N{excel_row})")
        ws.cell(row=excel_row, column=16, value=f"=P{excel_row - 1}-O{excel_row}")
    total_row = first_data_row + len(rows)
    ws.cell(row=total_row, column=1, value="TOTALS")
    ws.cell(row=total_row, column=6, value=f"=SUM(F3:F{max(total_row - 1, 3)})")
    ws.cell(row=total_row, column=8, value=f"=SUM(H3:H{max(total_row - 1, 3)})")
    ws.cell(row=total_row, column=9, value=f"=SUM(I3:I{max(total_row - 1, 3)})")
    ws.cell(row=total_row, column=10, value=f"=SUM(J3:J{max(total_row - 1, 3)})")
    ws.cell(row=total_row, column=11, value=f"=SUM(K3:K{max(total_row - 1, 3)})")
    ws.cell(row=total_row, column=13, value=f"=SUM(M3:M{max(total_row - 1, 3)})")
    ws.cell(row=total_row, column=14, value=f"=SUM(N3:N{max(total_row - 1, 3)})")
    ws.cell(row=total_row, column=15, value=f"=SUM(O3:O{max(total_row - 1, 3)})")
    for column in "ABCDEFGHIJKLMNOP":
        ws.column_dimensions[column].width = 16
    for row in ws.iter_rows(min_row=3, max_row=total_row, min_col=3, max_col=16):
        for cell in row:
            cell.number_format = '0.00'
    for cell in ws.iter_cols(min_col=2, max_col=2, min_row=4, max_row=max(total_row - 1, 4)):
        for item in cell:
            item.number_format = 'm/d/yy'
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    total_amount = sum((_money(row.get("total")) for row in rows), Decimal("0.00")).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
    return total_amount, current_balance


async def _fund_support_documents_for_packet(
    db: Db,
    *,
    fund_id: str,
    period_start: str,
    period_end: str,
) -> list[dict[str, object]]:
    rows = await db.fetchall(
        """SELECT DISTINCT a.id, a.attachment_type, a.original_filename, a.local_path,
                  a.content_type, a.sha256, e.entry_date, e.user_email
           FROM pay_attachments a
           JOIN pay_entries e ON e.id = a.entry_id
           JOIN pay_fund_allocations fa ON fa.entry_id = e.id AND fa.fund_id=?
           WHERE e.entry_date BETWEEN ? AND ?
             AND a.removed_at_utc IS NULL
             AND COALESCE(e.review_status, 'pending') IN ('pending', 'approved')
           UNION
           SELECT DISTINCT a.id, a.attachment_type, a.original_filename, a.local_path,
                  a.content_type, a.sha256, e.entry_date, e.user_email
           FROM pay_fund_attachment_links l
           JOIN pay_attachments a ON a.id = l.attachment_id
           JOIN pay_entries e ON e.id = a.entry_id
           WHERE l.fund_id=?
             AND e.entry_date BETWEEN ? AND ?
             AND a.removed_at_utc IS NULL
             AND COALESCE(e.review_status, 'pending') IN ('pending', 'approved')
           ORDER BY entry_date, original_filename""",
        (fund_id, period_start, period_end, fund_id, period_start, period_end),
    )
    return [
        {
            "id": row[0],
            "attachment_type": row[1],
            "original_filename": row[2],
            "local_path": row[3],
            "content_type": row[4],
            "sha256": row[5],
            "entry_date": row[6],
            "user_email": row[7],
        }
        for row in rows
    ]


async def generate_pay_fund_packet(
    db: Db,
    *,
    cfg: Any,
    fund_id: str,
    actor: PayActor,
    period_start: str,
    period_end: str,
    graph: Any | None = None,
) -> dict[str, object]:
    fund = await pay_fund_by_id(db, fund_id=fund_id)
    if not fund:
        raise ValueError("fund not found")
    start = date.fromisoformat(str(period_start or "").strip())
    end = date.fromisoformat(str(period_end or "").strip())
    if end < start:
        raise ValueError("period_end must be on or after period_start")
    settings = await pay_settings(db, pay_cfg=getattr(cfg, "pay_portal", None))
    fica_rate = fund_fica_rate_from_settings(settings)
    rows = await _fund_packet_rows(
        db,
        fund_id=fund_id,
        period_start=start.isoformat(),
        period_end=end.isoformat(),
        fica_rate=fica_rate,
    )
    ledger = await _fund_ledger_totals(db, fund_id=fund_id, through_date=end.isoformat())
    prior_allocated = await _fund_allocated_total(
        db,
        fund_id=fund_id,
        fica_rate=fica_rate,
        before_date=start.isoformat(),
    )
    beginning_balance = (_money(ledger.get("available")) - prior_allocated).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
    packet_id = f"pay-fund-packet-{uuid4().hex}"
    packet_dir = Path(cfg.data_root) / "pay" / "funds" / safe_filename(fund.get("name"), fallback=fund_id) / "packets" / packet_id
    workbook_path = packet_dir / f"{safe_filename(fund.get('name'), fallback='fund')}_{start.isoformat()}_to_{end.isoformat()}.xlsx"
    total_amount, ending_balance = _write_fund_packet_workbook(
        path=workbook_path,
        fund=fund,
        period_start=start.isoformat(),
        period_end=end.isoformat(),
        beginning_balance=beginning_balance,
        rows=rows,
        fica_rate=fica_rate,
    )
    support_original_dir = packet_dir / "support" / "original"
    support_pdf_dir = packet_dir / "support" / "pdf"
    support_original_dir.mkdir(parents=True, exist_ok=True)
    support_pdf_dir.mkdir(parents=True, exist_ok=True)
    support_docs = await _fund_support_documents_for_packet(
        db,
        fund_id=fund_id,
        period_start=start.isoformat(),
        period_end=end.isoformat(),
    )
    manifest_docs: list[dict[str, object]] = []
    for index, doc in enumerate(support_docs, start=1):
        source = Path(str(doc.get("local_path") or ""))
        if not source.exists():
            continue
        original_name = f"{index:02d}-{safe_filename(doc.get('entry_date'), fallback='date')}-{safe_filename(doc.get('original_filename'), fallback='support')}"
        original_path = support_original_dir / original_name
        shutil.copy2(source, original_path)
        pdf_path = None
        if str(doc.get("content_type") or "") == "application/pdf":
            pdf_path = support_pdf_dir / f"{Path(original_name).stem}.pdf"
            shutil.copy2(source, pdf_path)
        elif str(doc.get("content_type") or "").startswith("image/"):
            pdf_path = support_pdf_dir / f"{Path(original_name).stem}.pdf"
            image_to_pdf(str(source), str(pdf_path))
        manifest_docs.append(
            {
                **doc,
                "original_path": str(original_path),
                "pdf_path": str(pdf_path) if pdf_path else None,
            }
        )
    manifest = {
        "packet_id": packet_id,
        "fund": fund,
        "period_start": start.isoformat(),
        "period_end": end.isoformat(),
        "fica_rate": str(fica_rate),
        "beginning_balance": float(beginning_balance),
        "total_amount": float(total_amount),
        "ending_balance": float(ending_balance),
        "rows": rows,
        "support_documents": manifest_docs,
        "created_by": actor.email,
        "created_at_utc": utcnow(),
    }
    manifest_path = packet_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2, sort_keys=True), encoding="utf-8")
    workbook_sha = hashlib.sha256(workbook_path.read_bytes()).hexdigest()
    sharepoint_folder_path = None
    sharepoint_folder_web_url = None
    upload_folder = _fund_packet_support_folder_name(fund, start.isoformat(), end.isoformat(), packet_id)
    workbook_url, workbook_sp_path = await _upload_if_configured(
        cfg=cfg,
        graph=graph,
        folder_path=upload_folder,
        filename=workbook_path.name,
        local_path=str(workbook_path),
    )
    if workbook_url:
        sharepoint_folder_web_url = workbook_url
        sharepoint_folder_path = workbook_sp_path
        await _upload_if_configured(
            cfg=cfg,
            graph=graph,
            folder_path=upload_folder,
            filename=manifest_path.name,
            local_path=str(manifest_path),
        )
        for doc in manifest_docs:
            for key in ("original_path", "pdf_path"):
                value = doc.get(key)
                if value:
                    await _upload_if_configured(
                        cfg=cfg,
                        graph=graph,
                        folder_path=f"{upload_folder}/support",
                        filename=Path(str(value)).name,
                        local_path=str(value),
                    )
    now = utcnow()
    await db.exec(
        """INSERT INTO pay_fund_packets(
             id, fund_id, period_start, period_end, status, packet_dir_path,
             workbook_path, manifest_path, workbook_sha256, total_amount,
             beginning_balance, ending_balance, sharepoint_folder_path,
             sharepoint_folder_web_url, created_by, created_at_utc, updated_at_utc
           ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            packet_id,
            fund_id,
            start.isoformat(),
            end.isoformat(),
            "generated",
            str(packet_dir),
            str(workbook_path),
            str(manifest_path),
            workbook_sha,
            float(total_amount),
            float(beginning_balance),
            float(ending_balance),
            sharepoint_folder_path,
            sharepoint_folder_web_url,
            actor.email,
            now,
            now,
        ),
    )
    await add_pay_event(
        db,
        period_id=None,
        packet_id=None,
        event_type="fund_packet_generated",
        actor=actor.email,
        details={"fund_id": fund_id, "fund_packet_id": packet_id, "total_amount": float(total_amount)},
    )
    return {
        "id": packet_id,
        "fund_id": fund_id,
        "period_start": start.isoformat(),
        "period_end": end.isoformat(),
        "status": "generated",
        "packet_dir_path": str(packet_dir),
        "workbook_path": str(workbook_path),
        "manifest_path": str(manifest_path),
        "workbook_sha256": workbook_sha,
        "total_amount": float(total_amount),
        "beginning_balance": float(beginning_balance),
        "ending_balance": float(ending_balance),
        "support_document_count": len(manifest_docs),
        "sharepoint_folder_path": sharepoint_folder_path,
        "sharepoint_folder_web_url": sharepoint_folder_web_url,
        "workbook_download_url": f"/pay/api/funds/packets/{packet_id}/workbook",
    }


async def list_pay_fund_packets(db: Db, *, fund_id: str | None = None) -> list[dict[str, object]]:
    params: tuple[object, ...] = ()
    where = ""
    if fund_id:
        where = "WHERE p.fund_id=?"
        params = (fund_id,)
    rows = await db.fetchall(
        f"""SELECT p.id, p.fund_id, f.name, p.period_start, p.period_end, p.status,
                  p.workbook_path, p.manifest_path, p.workbook_sha256, p.total_amount,
                  p.beginning_balance, p.ending_balance, p.sharepoint_folder_path,
                  p.sharepoint_folder_web_url, p.created_by, p.created_at_utc, p.updated_at_utc
           FROM pay_fund_packets p
           JOIN pay_funds f ON f.id = p.fund_id
           {where}
           ORDER BY p.created_at_utc DESC""",
        params,
    )
    return [
        {
            "id": row[0],
            "fund_id": row[1],
            "fund_name": row[2],
            "period_start": row[3],
            "period_end": row[4],
            "status": row[5],
            "workbook_path": row[6],
            "manifest_path": row[7],
            "workbook_sha256": row[8],
            "total_amount": row[9],
            "beginning_balance": row[10],
            "ending_balance": row[11],
            "sharepoint_folder_path": row[12],
            "sharepoint_folder_web_url": row[13],
            "created_by": row[14],
            "created_at_utc": row[15],
            "updated_at_utc": row[16],
            "workbook_download_url": f"/pay/api/funds/packets/{row[0]}/workbook",
        }
        for row in rows
    ]


async def pay_fund_packet_by_id(db: Db, *, packet_id: str) -> dict[str, object]:
    packets = await db.fetchall(
        """SELECT p.id, p.fund_id, f.name, p.period_start, p.period_end, p.status,
                  p.workbook_path, p.manifest_path, p.workbook_sha256, p.total_amount,
                  p.beginning_balance, p.ending_balance, p.sharepoint_folder_path,
                  p.sharepoint_folder_web_url, p.created_by, p.created_at_utc, p.updated_at_utc
           FROM pay_fund_packets p
           JOIN pay_funds f ON f.id = p.fund_id
           WHERE p.id=?""",
        (packet_id,),
    )
    if not packets:
        raise ValueError("fund packet not found")
    row = packets[0]
    return {
        "id": row[0],
        "fund_id": row[1],
        "fund_name": row[2],
        "period_start": row[3],
        "period_end": row[4],
        "status": row[5],
        "workbook_path": row[6],
        "manifest_path": row[7],
        "workbook_sha256": row[8],
        "total_amount": row[9],
        "beginning_balance": row[10],
        "ending_balance": row[11],
        "sharepoint_folder_path": row[12],
        "sharepoint_folder_web_url": row[13],
        "created_by": row[14],
        "created_at_utc": row[15],
        "updated_at_utc": row[16],
    }


def _default_pay_demo_settings() -> dict[str, object]:
    return {
        "demo_mode_enabled": True,
        "demo_cycle_title": "Training Demo Cycle",
        "demo_cycle_notes": "",
    }


async def pay_demo_settings(db: Db) -> dict[str, object]:
    defaults = _default_pay_demo_settings()
    row = await db.app_setting("pay_demo")
    if not row:
        return defaults
    try:
        parsed = json.loads(row[0])
    except Exception:
        parsed = {}
    if not isinstance(parsed, dict):
        return defaults
    merged = dict(defaults)
    if "demo_mode_enabled" in parsed:
        value = parsed.get("demo_mode_enabled")
        merged["demo_mode_enabled"] = not (value is False or str(value).lower() == "false")
    if "demo_cycle_title" in parsed:
        merged["demo_cycle_title"] = str(parsed.get("demo_cycle_title") or "").strip() or "Training Demo Cycle"
    if "demo_cycle_notes" in parsed:
        merged["demo_cycle_notes"] = str(parsed.get("demo_cycle_notes") or "").strip()
    return merged


async def save_pay_demo_settings(
    db: Db,
    *,
    setting: dict[str, object],
    updated_by: str | None,
) -> dict[str, object]:
    normalized = dict(await pay_demo_settings(db))
    if "demo_mode_enabled" in setting:
        value = setting.get("demo_mode_enabled")
        normalized["demo_mode_enabled"] = not (value is False or str(value).lower() == "false")
    if "demo_cycle_title" in setting:
        normalized["demo_cycle_title"] = str(setting.get("demo_cycle_title") or "").strip() or "Training Demo Cycle"
    if "demo_cycle_notes" in setting:
        normalized["demo_cycle_notes"] = str(setting.get("demo_cycle_notes") or "").strip()
    await db.upsert_app_setting(setting_key="pay_demo", setting=normalized, updated_by=updated_by)
    return normalized


def pay_demo_artifact_dir(*, data_root: str, actor: PayActor) -> Path:
    label = safe_filename(actor.email or actor.display_name, fallback="demo-user")
    return Path(data_root) / "pay" / "demo" / label


def _demo_artifact_metadata(path: Path) -> dict[str, object]:
    stat = path.stat()
    return {
        "filename": path.name,
        "size_bytes": int(stat.st_size),
        "updated_at_utc": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }


def list_pay_demo_artifacts(*, data_root: str, actor: PayActor) -> list[dict[str, object]]:
    root = pay_demo_artifact_dir(data_root=data_root, actor=actor)
    if not root.exists():
        return []
    rows = [
        _demo_artifact_metadata(path)
        for path in root.iterdir()
        if path.is_file() and path.suffix.lower() == ".pdf"
    ]
    rows.sort(key=lambda row: str(row["filename"]))
    return rows


def pay_demo_artifact_path(*, data_root: str, actor: PayActor, filename: str) -> Path:
    safe_name = safe_filename(filename, fallback="")
    if not safe_name or safe_name != filename:
        raise ValueError("invalid demo artifact filename")
    root = pay_demo_artifact_dir(data_root=data_root, actor=actor)
    path = root / safe_name
    if path.suffix.lower() != ".pdf" or not path.exists() or not path.is_file() or path.parent != root:
        raise ValueError("demo artifact not found")
    return path


_DEMO_PACKET_PEOPLE: tuple[dict[str, str], ...] = (
    {
        "display_name": "Demo President",
        "email": "demo.president@cwa3106.local",
        "address": "117 W Duval St, Jacksonville, FL 32202, USA",
        "role": "president",
    },
    {
        "display_name": "Demo Treasurer",
        "email": "demo.treasurer@cwa3106.local",
        "address": "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
        "role": "treasurer",
    },
    {
        "display_name": "Demo Steward",
        "email": "demo.steward@cwa3106.local",
        "address": "11700 Phillips Hwy, Jacksonville, FL 32256, USA",
        "role": "steward",
    },
    {
        "display_name": "Demo Area Rep",
        "email": "demo.area.rep@cwa3106.local",
        "address": "98 Penman Rd S, Jacksonville Beach, FL 32250, USA",
        "role": "officer",
    },
    {
        "display_name": "Demo Chief Steward",
        "email": "demo.chief.steward@cwa3106.local",
        "address": "214 N Hogan St, Jacksonville, FL 32202, USA",
        "role": "steward",
    },
    {
        "display_name": "Demo Mobilization Lead",
        "email": "demo.mobilization@cwa3106.local",
        "address": "9209 Haydon Rd, Jacksonville, FL 32218, USA",
        "role": "officer",
    },
    {
        "display_name": "Demo Unit Steward",
        "email": "demo.unit.steward@cwa3106.local",
        "address": "6234 St Augustine Rd, Jacksonville, FL 32217, USA",
        "role": "steward",
    },
    {
        "display_name": "Demo Executive Board",
        "email": "demo.eboard@cwa3106.local",
        "address": "2048 Hendricks Ave, Jacksonville, FL 32207, USA",
        "role": "officer",
    },
    {
        "display_name": "Demo Safety Rep",
        "email": "demo.safety@cwa3106.local",
        "address": "1001 St Johns Bluff Rd N, Jacksonville, FL 32225, USA",
        "role": "steward",
    },
)
_DEMO_NARRATIVE_PHRASES: tuple[str, ...] = (
    "met with member about payroll correction and documented next steps",
    "reviewed route mileage, receipts, and pay profile rate for officer practice",
    "prepared grievance packet notes for treasurer review",
    "confirmed president differential example and president-only packet signature path",
    "reconciled mileage attachment with daily expense voucher totals",
    "recorded officer training feedback for the demo cycle",
    "checked supporting documents before the demo lock step",
)
_DEMO_MILEAGE_ROUTES: tuple[tuple[str, ...], ...] = (
    (
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
        "11700 Phillips Hwy, Jacksonville, FL 32256, USA",
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
    ),
    (
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
        "214 N Hogan St, Jacksonville, FL 32202, USA",
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
    ),
    (
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
        "98 Penman Rd S, Jacksonville Beach, FL 32250, USA",
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
    ),
    (
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
        "9209 Haydon Rd, Jacksonville, FL 32218, USA",
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
    ),
    (
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
        "6234 St Augustine Rd, Jacksonville, FL 32217, USA",
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
    ),
    (
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
        "1001 St Johns Bluff Rd N, Jacksonville, FL 32225, USA",
        "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
    ),
)


def _demo_people_for_packet(actor: PayActor) -> list[dict[str, str]]:
    actor_email = normalize_email(actor.email) or "demo.officer@cwa3106.local"
    people = [
        {
            "display_name": "Nick Craig",
            "email": actor_email,
            "address": "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
            "role": "officer",
        }
    ]
    people.extend(dict(row) for row in _DEMO_PACKET_PEOPLE)
    return people


def _demo_day_offsets(rng: random.Random, *, demo_step: int, person_index: int) -> list[int]:
    if demo_step < 1:
        return []
    week_one = list(range(1, 6))
    week_two = list(range(8, 13))
    if person_index < 3:
        return sorted(rng.sample(week_one, 4) + rng.sample(week_two, 4))
    return sorted([rng.choice(week_one), rng.choice(week_two)])


def _demo_hourly_rate(rng: random.Random, *, person_index: int) -> Decimal:
    if person_index == 0:
        return Decimal("250.00")
    if person_index == 1:
        return Decimal("45.00")
    base_cents = rng.randrange(3200, 5900, 25)
    return (Decimal(base_cents) / Decimal("100")).quantize(_CURRENCY, rounding=ROUND_HALF_UP)


def _demo_entries_for_packet(
    *,
    period_start: str,
    demo_step: int,
    person: dict[str, str],
    day_offsets: list[int],
    hourly_rate: Decimal,
    mileage_miles: Decimal,
    mileage_rate: Decimal,
    mileage_amount: Decimal,
    rng: random.Random,
) -> list[dict[str, object]]:
    try:
        start = date.fromisoformat(str(period_start))
    except Exception:
        start = date.today()
    if demo_step < 1:
        return []
    rows: list[dict[str, object]] = []
    base_miles = _money(mileage_miles)
    if base_miles <= 0:
        base_miles = Decimal("24.00")
    for entry_index, day_offset in enumerate(day_offsets, start=1):
        has_receipt = demo_step >= 3 and entry_index == len(day_offsets)
        if person.get("role") == "president":
            hours = Decimal(str(rng.choice(["2.00", "3.00", "5.00", "6.00"])))
        else:
            hours = Decimal(str(rng.choice(["2.00", "2.50", "3.00", "4.00", "6.00", "8.00"])))
        misc_amount = Decimal(str(rng.choice(["12.50", "18.75", "24.00", "31.25"]))) if has_receipt else Decimal("0")
        meals_amount = Decimal(str(rng.choice(["15.00", "22.50", "28.00"]))) if has_receipt and entry_index % 2 == 0 else Decimal("0")
        normal_hourly_rate = Decimal("45.00") if person.get("role") == "president" else hourly_rate
        presidential_hourly_rate = Decimal("62.00") if person.get("role") == "president" else normal_hourly_rate
        scheduled_hours = Decimal("8.00")
        president_diff_hours = _quantity(scheduled_hours - hours) if person.get("role") == "president" and demo_step >= 3 else Decimal("0")
        if president_diff_hours < 0:
            president_diff_hours = Decimal("0")
        president_diff_rate = (presidential_hourly_rate - normal_hourly_rate).quantize(_CURRENCY, rounding=ROUND_HALF_UP) if president_diff_hours > 0 else Decimal("0")
        if president_diff_rate < 0:
            president_diff_rate = Decimal("0")
        president_diff = (president_diff_hours * president_diff_rate).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
        note = rng.choice(_DEMO_NARRATIVE_PHRASES)
        if entry_index == 1:
            note = "reviewed route mileage, receipts, and pay profile rate for officer practice"
        if person.get("role") == "president" and president_diff_hours > 0:
            note = f"president demo split {hours:f} union hours from {president_diff_hours:f} scheduled employer-work differential hours"
        miles_multiplier = Decimal(str(rng.choice(["0.70", "0.85", "1.00", "1.15", "1.30"])))
        entry_miles = (base_miles * miles_multiplier).quantize(_MILES, rounding=ROUND_HALF_UP)
        entry_mileage_amount = (entry_miles * mileage_rate).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
        rows.append(
            {
                "id": f"demo-entry-{safe_filename(person.get('email'), fallback='person')}-{entry_index}",
                "user_email": person.get("email") or "demo.officer@cwa3106.local",
                "entry_date": (start + timedelta(days=day_offset)).isoformat(),
                "display_name": person.get("display_name") or "Demo Officer",
                "address": person.get("address") or "4076 Union Hall Pl, Jacksonville, FL 32205, USA",
                "local_number": "3106",
                "hours": float(hours),
                "hourly_rate": float(presidential_hourly_rate),
                "lost_wage_hourly_rate": float(presidential_hourly_rate),
                "mileage_miles": float(entry_miles),
                "mileage_rate": float(mileage_rate),
                "mileage_amount": float(entry_mileage_amount),
                "rentals_amount": 0,
                "meals_amount": float(meals_amount),
                "hotel_amount": 0,
                "miscellaneous_amount": float(misc_amount),
                "president_diff_hours": float(president_diff_hours),
                "president_diff_rate": float(president_diff_rate),
                "president_diff_amount": float(president_diff),
                "submitter_certified_at_utc": f"{(start + timedelta(days=day_offset)).isoformat()}T12:00:00Z",
                "submitter_certified_by": person.get("email") or "demo.officer@cwa3106.local",
                "submitter_certification_text": "DEMO TRAINING - submitter signed off on this daily entry in Pay Portal.",
                "notes": f"DEMO TRAINING - {note}.",
            }
        )
    return rows


def watermark_pdf(input_path: str, output_path: str | None = None, *, text: str = "DEMO") -> None:
    try:
        from pypdf import PdfReader, PdfWriter  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except ImportError as exc:  # pragma: no cover - runtime packaging guard
        raise RuntimeError("pypdf and reportlab are required for PDF watermarking") from exc

    source = Path(input_path)
    target = Path(output_path or input_path)
    tmp_target = target.with_name(f"{target.stem}.watermarking{target.suffix}")
    reader = PdfReader(str(source))
    writer = PdfWriter()
    for page in reader.pages:
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        overlay = io.BytesIO()
        c = canvas.Canvas(overlay, pagesize=(width, height))
        c.saveState()
        try:
            c.setFillAlpha(0.16)
        except Exception:
            pass
        c.setFillColorRGB(0.72, 0.72, 0.72)
        c.translate(width / 2, height / 2)
        c.rotate(35)
        c.setFont("Helvetica-Bold", max(72, min(width, height) / 4))
        c.drawCentredString(0, 0, text)
        c.restoreState()
        c.saveState()
        try:
            c.setFillAlpha(0.45)
        except Exception:
            pass
        c.setFillColorRGB(0.5, 0.5, 0.5)
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(width / 2, 24, "DEMO TRAINING PACKET - NOT FOR PAYROLL")
        c.restoreState()
        c.save()
        overlay.seek(0)
        watermark_page = PdfReader(overlay).pages[0]
        page.merge_page(watermark_page)
        writer.add_page(page)
    with open(tmp_target, "wb") as out:
        writer.write(out)
    tmp_target.replace(target)


def _demo_mileage_locations(person_index: int, rng: random.Random) -> list[str]:
    route_index = (person_index + rng.randrange(len(_DEMO_MILEAGE_ROUTES))) % len(_DEMO_MILEAGE_ROUTES)
    return list(_DEMO_MILEAGE_ROUTES[route_index])


def _emit_progress(progress_callback: Any, **payload: object) -> None:
    if progress_callback is None:
        return
    try:
        progress_callback(dict(payload))
    except Exception:
        return


def generate_pay_demo_artifacts(
    *,
    cfg: Any,
    settings: dict[str, object],
    actor: PayActor,
    demo_step: int,
    demo_cycle_title: str,
    period_start: str,
    period_end: str,
    docx_to_pdf_func: Any,
    graph: Any = None,
    progress_callback: Any = None,
) -> list[dict[str, object]]:
    try:
        step = max(0, min(int(demo_step), 20))
    except Exception:
        step = 0
    title = str(demo_cycle_title or "").strip() or "Training Demo Cycle"
    root = pay_demo_artifact_dir(data_root=cfg.data_root, actor=actor)
    root.mkdir(parents=True, exist_ok=True)
    for old_path in root.glob("demo-*"):
        if old_path.is_file() and old_path.suffix.lower() in {".html", ".txt", ".pdf", ".docx"}:
            old_path.unlink()
    work_dir = root / "work"
    work_dir.mkdir(parents=True, exist_ok=True)
    for old_path in work_dir.glob("demo-*"):
        if old_path.is_file():
            old_path.unlink()

    template_path = Path(str(cfg.pay_portal.voucher_template_path or ""))
    if not template_path.exists():
        raise RuntimeError("pay voucher template is not configured or does not exist")

    try:
        period_start_date = date.fromisoformat(str(period_start))
    except Exception:
        period_start_date = date.today()
    packet_step = max(step, 4)
    seed = f"{title}|{period_start}|{period_end}|{actor.email}|{packet_step}"
    people = _demo_people_for_packet(actor)
    packet_pdf_paths: list[str] = []
    plans: list[dict[str, object]] = []
    for person_index, person in enumerate(people, start=1):
        person_rng = random.Random(f"{seed}|{person_index}|{person.get('email', '')}")
        day_offsets = _demo_day_offsets(person_rng, demo_step=packet_step, person_index=person_index - 1)
        hourly_rate = _demo_hourly_rate(person_rng, person_index=person_index - 1)
        mileage_day_offset = day_offsets[0] if day_offsets else 1
        demo_day = (period_start_date + timedelta(days=mileage_day_offset)).isoformat()
        rate, rate_display = mileage_rate_from_settings(settings, int(demo_day[:4]))
        locations = _demo_mileage_locations(person_index - 1, person_rng)
        plans.append(
            {
                "person_index": person_index,
                "person": person,
                "rng": person_rng,
                "day_offsets": day_offsets,
                "hourly_rate": hourly_rate,
                "demo_day": demo_day,
                "rate": rate,
                "rate_display": rate_display,
                "locations": locations,
            }
        )

    _emit_progress(
        progress_callback,
        stage="mileage",
        current=0,
        total=len(plans),
        message=f"Prefetching mileage routes for {len(plans)} people",
    )
    _prefetch_google_legs(
        api_key=cfg.pay_portal.google_maps_api_key,
        routes=[list(plan["locations"]) for plan in plans],
        progress_callback=progress_callback,
    )

    for plan in plans:
        person_index = int(plan["person_index"])
        person = plan["person"]  # type: ignore[assignment]
        person_rng = plan["rng"]  # type: ignore[assignment]
        day_offsets = plan["day_offsets"]  # type: ignore[assignment]
        hourly_rate = plan["hourly_rate"]  # type: ignore[assignment]
        demo_day = str(plan["demo_day"])
        rate = plan["rate"]  # type: ignore[assignment]
        rate_display = str(plan["rate_display"])
        locations = plan["locations"]  # type: ignore[assignment]
        _emit_progress(
            progress_callback,
            stage="mileage",
            current=person_index,
            total=len(plans),
            message=f"Building mileage PDF {person_index}/{len(plans)}",
        )
        mileage_pdf_bytes, reimbursement, total_miles = build_mileage_pdf(
            name=person.get("display_name") or "Demo Officer",
            local_number="3106",
            date_str=demo_day,
            description=f"DEMO TRAINING - {title} mileage route for {person.get('display_name') or 'Demo Officer'}",
            rate=rate,
            rate_display=rate_display,
            locations=locations,
            google_maps_api_key=cfg.pay_portal.google_maps_api_key,
        )
        label = safe_filename(person.get("display_name") or person.get("email"), fallback=f"person-{person_index}")
        mileage_pdf_path = work_dir / f"demo-mileage-{person_index:02d}-{label}.pdf"
        mileage_pdf_path.write_bytes(mileage_pdf_bytes)

        entries = _demo_entries_for_packet(
            period_start=period_start,
            demo_step=packet_step,
            person=person,
            day_offsets=day_offsets,
            hourly_rate=hourly_rate,
            mileage_miles=total_miles,
            mileage_rate=rate,
            mileage_amount=reimbursement,
            rng=person_rng,
        )

        voucher_docx_path = work_dir / f"demo-voucher-{person_index:02d}-{label}.docx"
        _emit_progress(
            progress_callback,
            stage="voucher",
            current=person_index,
            total=len(plans),
            message=f"Filling voucher {person_index}/{len(plans)}",
        )
        fill_pay_voucher_docx(
            template_path=str(template_path),
            output_path=str(voucher_docx_path),
            period_start=period_start,
            period_end=period_end,
            entries=entries,
            include_signature_placeholders=False,
        )
        _emit_progress(
            progress_callback,
            stage="convert",
            current=person_index,
            total=len(plans),
            message=f"Converting voucher {person_index}/{len(plans)} to PDF",
        )
        voucher_pdf_path = Path(
            docx_to_pdf_func(
                str(voucher_docx_path),
                str(work_dir),
                int(getattr(cfg, "libreoffice_timeout_seconds", 45) or 45),
                engine=getattr(cfg, "docx_pdf_engine", "libreoffice"),
                graph_uploader=graph,
                graph_site_hostname=_pay_graph_site(cfg=cfg)[0],
                graph_site_path=_pay_graph_site(cfg=cfg)[1],
                graph_library=getattr(cfg.graph, "document_library", ""),
                graph_temp_folder_path=_pay_docx_pdf_graph_temp_folder(cfg=cfg),
            )
        )
        packet_pdf_paths.extend([str(voucher_pdf_path), str(mileage_pdf_path)])

    packet_path = root / "demo-payroll-packet.pdf"
    _emit_progress(progress_callback, stage="merge", current=0, total=1, message="Merging demo packet")
    merge_pdfs(packet_pdf_paths, str(packet_path))
    _emit_progress(progress_callback, stage="watermark", current=0, total=1, message="Watermarking demo packet")
    watermark_pdf(str(packet_path), text="DEMO")
    _emit_progress(progress_callback, stage="complete", current=1, total=1, message="Demo packet ready")
    return list_pay_demo_artifacts(data_root=cfg.data_root, actor=actor)


def _active_irs_rate_matches(settings: dict[str, object], *, year: str, rate_per_mile: Decimal) -> bool:
    rates = settings.get("irs_rates")
    if not isinstance(rates, dict):
        return False
    active = rates.get(str(year))
    if active is None:
        return False
    try:
        active_rate = Decimal(str(active))
    except Exception:
        return False
    return active_rate.quantize(_IRS_RATE_QUANT, rounding=ROUND_HALF_UP) == rate_per_mile


def _irs_rate_candidate_from_row(row: Any) -> dict[str, object]:
    return {
        "id": int(row[0]),
        "rate_year": row[1],
        "effective_date": row[2],
        "cents_per_mile": row[3],
        "rate_per_mile": _mileage_rate_text(row[4]),
        "source_url": row[5],
        "source_title": row[6],
        "detected_at_utc": row[7],
        "status": row[8],
        "approved_by": row[9],
        "approved_at_utc": row[10],
        "updated_at_utc": row[11],
    }


async def list_irs_rate_candidates(db: Db, *, status: str | None = None) -> list[dict[str, object]]:
    select_sql = """SELECT id, rate_year, effective_date, cents_per_mile, rate_per_mile,
                           source_url, source_title, detected_at_utc, status,
                           approved_by, approved_at_utc, updated_at_utc
                    FROM pay_irs_rate_candidates"""
    if status:
        rows = await db.fetchall(
            f"{select_sql} WHERE status=? ORDER BY effective_date DESC, id DESC",
            (status,),
        )
    else:
        rows = await db.fetchall(
            f"""{select_sql}
                ORDER BY CASE status WHEN 'pending' THEN 0 WHEN 'approved' THEN 1 ELSE 2 END,
                         effective_date DESC, id DESC"""
        )
    return [_irs_rate_candidate_from_row(row) for row in rows]


async def sync_irs_mileage_rate_candidates(
    db: Db,
    *,
    pay_cfg: Any | None = None,
    http_get: Any | None = None,
    source_urls: tuple[str, ...] | list[str] | None = None,
) -> dict[str, object]:
    if pay_cfg is not None and not bool(getattr(pay_cfg, "irs_rate_sync_enabled", True)):
        return {"detected": [], "skipped_existing": 0, "skipped_duplicates": 0, "failures": []}

    configured_urls = tuple(source_urls or getattr(pay_cfg, "irs_rate_source_urls", ()) or _DEFAULT_IRS_RATE_SOURCE_URLS)
    get = http_get or requests.get
    settings = await pay_settings(db, pay_cfg=pay_cfg)
    detected: list[dict[str, object]] = []
    failures: list[dict[str, str]] = []
    skipped_existing = 0
    skipped_duplicates = 0
    now = utcnow()

    for source_url in configured_urls:
        url = str(source_url or "").strip()
        if not url:
            continue
        try:
            def _fetch():  # noqa: ANN202
                try:
                    return get(url, timeout=15)
                except TypeError:
                    return get(url)

            response = await asyncio.to_thread(_fetch)
            status_code = int(getattr(response, "status_code", 200) or 200)
            if status_code >= 400:
                raise RuntimeError(f"IRS source returned HTTP {status_code}")
            content = str(getattr(response, "text", "") or "")
            candidates = parse_irs_mileage_rate_candidates(content=content, source_url=url)
            if not candidates:
                raise RuntimeError("no IRS business mileage rate found")
        except Exception as exc:
            failures.append({"source_url": url, "error": str(exc)})
            continue

        for candidate in candidates:
            if _active_irs_rate_matches(settings, year=candidate.rate_year, rate_per_mile=candidate.rate_per_mile):
                skipped_existing += 1
                continue
            existing = await db.fetchone(
                """SELECT id
                   FROM pay_irs_rate_candidates
                   WHERE rate_year=? AND effective_date=? AND cents_per_mile=? AND source_url=?""",
                (
                    candidate.rate_year,
                    candidate.effective_date,
                    float(candidate.cents_per_mile),
                    candidate.source_url,
                ),
            )
            if existing:
                skipped_duplicates += 1
                continue
            candidate_id = await db.insert(
                """INSERT INTO pay_irs_rate_candidates(
                     rate_year, effective_date, cents_per_mile, rate_per_mile,
                     source_url, source_title, detected_at_utc, status, updated_at_utc
                   ) VALUES(?,?,?,?,?,?,?,?,?)""",
                (
                    candidate.rate_year,
                    candidate.effective_date,
                    float(candidate.cents_per_mile),
                    float(candidate.rate_per_mile),
                    candidate.source_url,
                    candidate.source_title,
                    now,
                    "pending",
                    now,
                ),
            )
            row = await db.fetchone(
                """SELECT id, rate_year, effective_date, cents_per_mile, rate_per_mile,
                          source_url, source_title, detected_at_utc, status,
                          approved_by, approved_at_utc, updated_at_utc
                   FROM pay_irs_rate_candidates
                   WHERE id=?""",
                (candidate_id,),
            )
            if row:
                detected.append(_irs_rate_candidate_from_row(row))

    if failures:
        await add_pay_event(
            db,
            period_id=None,
            event_type="irs_rate_sync_failed",
            details={"failures": failures},
        )
    if detected:
        await add_pay_event(
            db,
            period_id=None,
            event_type="irs_rate_candidates_detected",
            details={"candidate_ids": [row["id"] for row in detected]},
        )
    return {
        "detected": detected,
        "skipped_existing": skipped_existing,
        "skipped_duplicates": skipped_duplicates,
        "failures": failures,
    }


async def approve_irs_rate_candidate(
    db: Db,
    *,
    candidate_id: int,
    actor: str,
    pay_cfg: Any | None = None,
) -> dict[str, object]:
    row = await db.fetchone(
        """SELECT id, rate_year, effective_date, cents_per_mile, rate_per_mile,
                  source_url, source_title, detected_at_utc, status,
                  approved_by, approved_at_utc, updated_at_utc
           FROM pay_irs_rate_candidates
           WHERE id=?""",
        (int(candidate_id),),
    )
    if not row:
        raise ValueError("IRS rate candidate not found")
    current = _irs_rate_candidate_from_row(row)
    if current["status"] not in {"pending", "approved"}:
        raise ValueError("IRS rate candidate is not pending")

    rate_year = str(current["rate_year"])
    rate_text = _mileage_rate_text(current["rate_per_mile"])
    settings = await pay_settings(db, pay_cfg=pay_cfg)
    rates = settings.get("irs_rates")
    active_rates = dict(rates) if isinstance(rates, dict) else {}
    active_rates[rate_year] = rate_text
    await save_pay_settings(
        db,
        setting={"irs_rates": active_rates},
        updated_by=actor,
        pay_cfg=pay_cfg,
    )
    now = utcnow()
    await db.exec(
        """UPDATE pay_irs_rate_candidates
           SET status='approved', approved_by=?, approved_at_utc=?, updated_at_utc=?
           WHERE id=?""",
        (actor, now, now, int(candidate_id)),
    )
    await add_pay_event(
        db,
        period_id=None,
        event_type="irs_rate_candidate_approved",
        actor=actor,
        details={"candidate_id": int(candidate_id), "rate_year": rate_year, "rate_per_mile": rate_text},
    )
    approved = await db.fetchone(
        """SELECT id, rate_year, effective_date, cents_per_mile, rate_per_mile,
                  source_url, source_title, detected_at_utc, status,
                  approved_by, approved_at_utc, updated_at_utc
           FROM pay_irs_rate_candidates
           WHERE id=?""",
        (int(candidate_id),),
    )
    result = _irs_rate_candidate_from_row(approved)
    result["active_rate"] = rate_text
    return result


async def treasurer_recipients(db: Db, *, fallback: tuple[str, ...], pay_cfg: Any | None = None) -> list[str]:
    settings = await pay_settings(db, pay_cfg=pay_cfg)
    configured = settings.get("treasurer_emails")
    recipients: list[str] = []
    if isinstance(configured, list):
        recipients.extend(str(v).strip() for v in configured if str(v).strip())
    elif isinstance(configured, str):
        recipients.extend(v.strip() for v in configured.split(",") if v.strip())
    rows = await db.fetchall("SELECT email FROM pay_users WHERE status='active' AND role='treasurer'")
    recipients.extend(str(row[0] or "").strip() for row in rows if str(row[0] or "").strip())
    internal_rows = await db.fetchall(
        """
        SELECT principal_email
        FROM internal_role_assignments
        WHERE status='active' AND role='treasurer'
        """
    )
    recipients.extend(str(row[0] or "").strip() for row in internal_rows if str(row[0] or "").strip())
    if not recipients:
        recipients.extend(fallback)
    out: list[str] = []
    seen: set[str] = set()
    for recipient in recipients:
        lowered = recipient.lower()
        if lowered and lowered not in seen:
            seen.add(lowered)
            out.append(recipient)
    return out


async def president_email(db: Db, *, explicit: str | None = None, pay_cfg: Any | None = None) -> str:
    if explicit and explicit.strip():
        return explicit.strip()
    settings = await pay_settings(db, pay_cfg=pay_cfg)
    return str(settings.get("president_email") or "").strip()


async def upsert_pay_user(
    db: Db,
    *,
    email: str,
    display_name: str | None,
    role: str,
    status: str,
    actor: str,
) -> dict[str, object]:
    normalized_email = normalize_email(email)
    if not normalized_email:
        raise ValueError("email is required")
    normalized_role = str(role or "guest").strip().lower()
    if normalized_role not in {"guest", "treasurer"}:
        normalized_role = "guest"
    normalized_status = str(status or "active").strip().lower()
    if normalized_status not in {"active", "disabled"}:
        normalized_status = "active"
    now = utcnow()
    await db.exec(
        """
        INSERT INTO pay_users(email, display_name, role, status, created_at_utc, updated_at_utc, invited_by)
        VALUES(?,?,?,?,?,?,?)
        ON CONFLICT(email) DO UPDATE SET
          display_name=excluded.display_name,
          role=excluded.role,
          status=excluded.status,
          updated_at_utc=excluded.updated_at_utc
        """,
        (normalized_email, display_name, normalized_role, normalized_status, now, now, actor),
    )
    return {
        "email": normalized_email,
        "display_name": display_name,
        "role": normalized_role,
        "status": normalized_status,
    }


async def pay_user_by_email(db: Db, email: str) -> dict[str, object] | None:
    row = await db.fetchone(
        "SELECT email, display_name, role, status FROM pay_users WHERE email=?",
        (normalize_email(email),),
    )
    if not row:
        return None
    return {"email": row[0], "display_name": row[1], "role": row[2], "status": row[3]}


async def list_pay_users(db: Db) -> list[dict[str, object]]:
    rows = await db.fetchall(
        "SELECT email, display_name, role, status, updated_at_utc FROM pay_users ORDER BY email"
    )
    return [
        {
            "email": row[0],
            "display_name": row[1],
            "role": row[2],
            "status": row[3],
            "updated_at_utc": row[4],
        }
        for row in rows
    ]


_DEMO_FEEDBACK_CATEGORIES = {"suggestion", "confusing", "missing", "bug", "training"}


def normalize_demo_feedback_category(value: object) -> str:
    category = str(value or "suggestion").strip().lower()
    return category if category in _DEMO_FEEDBACK_CATEGORIES else "suggestion"


def _demo_feedback_from_row(row: Any) -> dict[str, object]:
    return {
        "id": int(row[0]),
        "created_at_utc": row[1],
        "actor_email": row[2],
        "actor_display_name": row[3],
        "actor_role": row[4],
        "demo_step": int(row[5] or 0),
        "demo_cycle_title": row[6],
        "screen": row[7],
        "category": row[8],
        "comment": row[9],
        "status": row[10],
    }


async def list_pay_demo_feedback(
    db: Db,
    *,
    status: str | None = None,
    limit: int = 100,
) -> list[dict[str, object]]:
    bounded_limit = max(1, min(int(limit or 100), 500))
    params: list[object] = []
    where = ""
    clean_status = str(status or "").strip().lower()
    if clean_status:
        where = "WHERE status=?"
        params.append(clean_status)
    params.append(bounded_limit)
    rows = await db.fetchall(
        f"""SELECT id, created_at_utc, actor_email, actor_display_name, actor_role,
                  demo_step, demo_cycle_title, screen, category, comment, status
            FROM pay_demo_feedback
            {where}
            ORDER BY id DESC
            LIMIT ?""",
        tuple(params),
    )
    return [_demo_feedback_from_row(row) for row in rows]


async def create_pay_demo_feedback(
    db: Db,
    *,
    actor: PayActor,
    screen: str | None,
    category: str | None,
    comment: str | None,
    demo_step: int | None = None,
    demo_cycle_title: str | None = None,
) -> dict[str, object]:
    text = str(comment or "").strip()
    if not text:
        raise ValueError("suggestion is required")
    if len(text) > 4000:
        raise ValueError("suggestion must be 4000 characters or less")
    clean_screen = re.sub(r"[^a-z0-9_-]+", "_", str(screen or "demo").strip().lower()).strip("_") or "demo"
    clean_screen = clean_screen[:80]
    clean_category = normalize_demo_feedback_category(category)
    try:
        clean_demo_step = max(0, min(int(demo_step or 0), 20))
    except Exception:
        clean_demo_step = 0
    clean_cycle_title = str(demo_cycle_title or "").strip()[:200] or "Training Demo Cycle"
    created_at = utcnow()
    feedback_id = await db.insert(
        """INSERT INTO pay_demo_feedback(
             created_at_utc, actor_email, actor_display_name, actor_role,
             demo_step, demo_cycle_title, screen, category, comment, status
           )
           VALUES(?,?,?,?,?,?,?,?,?,'open')""",
        (
            created_at,
            actor.email,
            actor.display_name,
            actor.role,
            clean_demo_step,
            clean_cycle_title,
            clean_screen,
            clean_category,
            text,
        ),
    )
    row = await db.fetchone(
        """SELECT id, created_at_utc, actor_email, actor_display_name, actor_role,
                  demo_step, demo_cycle_title, screen, category, comment, status
           FROM pay_demo_feedback
           WHERE id=?""",
        (feedback_id,),
    )
    return _demo_feedback_from_row(row)


async def update_pay_demo_feedback_status(
    db: Db,
    *,
    feedback_id: int,
    status: str,
) -> dict[str, object]:
    clean_status = str(status or "").strip().lower()
    if clean_status not in {"open", "closed"}:
        raise ValueError("status must be open or closed")
    await db.exec("UPDATE pay_demo_feedback SET status=? WHERE id=?", (clean_status, int(feedback_id)))
    row = await db.fetchone(
        """SELECT id, created_at_utc, actor_email, actor_display_name, actor_role,
                  demo_step, demo_cycle_title, screen, category, comment, status
           FROM pay_demo_feedback
           WHERE id=?""",
        (int(feedback_id),),
    )
    if not row:
        raise ValueError("demo feedback not found")
    return _demo_feedback_from_row(row)


def _pay_profile_from_row(row: Any) -> dict[str, object]:
    return {
        "id": int(row[0]),
        "principal_id": row[1],
        "principal_email": row[2],
        "principal_display_name": row[3],
        "pay_basis": row[4],
        "base_wage_input_type": row[5],
        "base_wage_amount": float(row[6] or 0),
        "weekly_basis_hours": float(row[7] or 40),
        "commission_month_1_amount": float(row[8] or 0),
        "commission_month_2_amount": float(row[9] or 0),
        "commission_month_3_amount": float(row[10] or 0),
        "commission_average_monthly": float(row[11] or 0),
        "commission_hourly_rate": float(row[12] or 0),
        "calculated_hourly_rate": float(row[13] or 0),
        "default_address": row[14],
        "status": row[15],
        "notes": row[16],
        "created_at_utc": row[17],
        "updated_at_utc": row[18],
        "updated_by": row[19],
    }


_PAY_PROFILE_SELECT = """
    SELECT id, principal_id, principal_email, principal_display_name, pay_basis,
           base_wage_input_type, base_wage_amount, weekly_basis_hours,
           commission_month_1_amount, commission_month_2_amount, commission_month_3_amount,
           commission_average_monthly, commission_hourly_rate, calculated_hourly_rate,
           default_address, status, notes, created_at_utc, updated_at_utc, updated_by
    FROM pay_profiles
"""


async def list_pay_profiles(db: Db) -> list[dict[str, object]]:
    rows = await db.fetchall(
        f"{_PAY_PROFILE_SELECT} ORDER BY lower(principal_display_name), lower(principal_email)"
    )
    return [_pay_profile_from_row(row) for row in rows]


async def pay_profile_by_email(
    db: Db,
    *,
    email: str,
    active_only: bool = False,
) -> dict[str, object] | None:
    normalized_email = normalize_email(email)
    if not normalized_email:
        return None
    where = "WHERE principal_email=?"
    params: tuple[object, ...] = (normalized_email,)
    if active_only:
        where += " AND status='active'"
    row = await db.fetchone(f"{_PAY_PROFILE_SELECT} {where} LIMIT 1", params)
    return _pay_profile_from_row(row) if row else None


async def delete_pay_profile(
    db: Db,
    *,
    email: str,
    actor: str | None = None,
) -> dict[str, object]:
    normalized_email = normalize_email(email)
    if not normalized_email:
        raise ValueError("principal_email is required")
    existing = await pay_profile_by_email(db, email=normalized_email)
    if not existing:
        raise ValueError("pay profile not found")
    now = utcnow()
    await db.exec(
        "UPDATE pay_profiles SET status='disabled', updated_at_utc=?, updated_by=? WHERE principal_email=?",
        (now, actor, normalized_email),
    )
    await add_pay_event(
        db,
        period_id=None,
        event_type="pay_profile_disabled",
        actor=actor,
        details={"principal_email": normalized_email, "previous_status": existing.get("status")},
    )
    disabled = await pay_profile_by_email(db, email=normalized_email)
    return disabled or {**existing, "status": "disabled", "updated_at_utc": now, "updated_by": actor}


def _pay_profile_wage_fields(profile: dict[str, object] | None) -> dict[str, object]:
    profile = profile or {}
    return {
        "pay_basis": str(profile.get("pay_basis") or "expense_only"),
        "base_wage_input_type": str(profile.get("base_wage_input_type") or "hourly"),
        "base_wage_amount": float(_money(profile.get("base_wage_amount"))),
        "weekly_basis_hours": float(_quantity(profile.get("weekly_basis_hours") or 40)),
        "commission_month_1_amount": float(_money(profile.get("commission_month_1_amount"))),
        "commission_month_2_amount": float(_money(profile.get("commission_month_2_amount"))),
        "commission_month_3_amount": float(_money(profile.get("commission_month_3_amount"))),
    }


def pay_profile_wage_fields_changed(existing: dict[str, object] | None, requested: dict[str, object]) -> bool:
    return _pay_profile_wage_fields(existing) != _pay_profile_wage_fields(requested)


def _profile_change_request_from_row(row: Any) -> dict[str, object]:
    return {
        "id": row[0],
        "principal_id": row[1],
        "principal_email": row[2],
        "principal_display_name": row[3],
        "pay_basis": row[4],
        "base_wage_input_type": row[5],
        "base_wage_amount": row[6],
        "weekly_basis_hours": row[7],
        "commission_month_1_amount": row[8],
        "commission_month_2_amount": row[9],
        "commission_month_3_amount": row[10],
        "commission_average_monthly": row[11],
        "commission_hourly_rate": row[12],
        "calculated_hourly_rate": row[13],
        "default_address": row[14],
        "profile_status": row[15],
        "notes": row[16],
        "requested_by": row[17],
        "requested_at_utc": row[18],
        "status": row[19],
        "reviewed_by": row[20],
        "reviewed_at_utc": row[21],
        "review_note": row[22],
        "current_profile": json.loads(row[23] or "{}"),
        "requested_profile": json.loads(row[24] or "{}"),
    }


_PAY_PROFILE_CHANGE_SELECT = """
    SELECT id, principal_id, principal_email, principal_display_name, pay_basis,
           base_wage_input_type, base_wage_amount, weekly_basis_hours,
           commission_month_1_amount, commission_month_2_amount, commission_month_3_amount,
           commission_average_monthly, commission_hourly_rate, calculated_hourly_rate,
           default_address, profile_status, notes, requested_by, requested_at_utc, status,
           reviewed_by, reviewed_at_utc, review_note, current_profile_json, requested_profile_json
    FROM pay_profile_change_requests
"""


async def list_pay_profile_change_requests(
    db: Db,
    *,
    email: str | None = None,
    pending_only: bool = False,
) -> list[dict[str, object]]:
    where: list[str] = []
    params: list[object] = []
    normalized_email = normalize_email(email) if email else ""
    if normalized_email:
        where.append("principal_email=?")
        params.append(normalized_email)
    if pending_only:
        where.append("status='pending'")
    sql = _PAY_PROFILE_CHANGE_SELECT
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY requested_at_utc DESC"
    rows = await db.fetchall(sql, tuple(params))
    return [_profile_change_request_from_row(row) for row in rows]


async def request_pay_profile_change(
    db: Db,
    *,
    principal_id: str | None,
    principal_email: str,
    principal_display_name: str | None,
    pay_basis: str,
    base_wage_input_type: str,
    base_wage_amount: float,
    weekly_basis_hours: float,
    commission_month_1_amount: float,
    commission_month_2_amount: float,
    commission_month_3_amount: float,
    status: str,
    notes: str | None,
    requested_by: str,
    default_address: str | None = None,
) -> dict[str, object]:
    normalized_email = normalize_email(principal_email)
    if not normalized_email:
        raise ValueError("principal_email is required")
    normalized_status = normalize_pay_profile_status(status)
    snapshot = calculate_pay_profile_snapshot(
        pay_basis=pay_basis,
        base_wage_input_type=base_wage_input_type,
        base_wage_amount=base_wage_amount,
        weekly_basis_hours=weekly_basis_hours,
        commission_month_1_amount=commission_month_1_amount,
        commission_month_2_amount=commission_month_2_amount,
        commission_month_3_amount=commission_month_3_amount,
    )
    if snapshot["pay_basis"] != "expense_only" and Decimal(str(snapshot["calculated_hourly_rate"])) <= 0:
        raise ValueError("base wage amount is required for this pay profile")
    existing = await pay_profile_by_email(db, email=normalized_email)
    request_id = f"pay-profile-change-{uuid4().hex}"
    now = utcnow()
    requested_profile = {
        "principal_id": str(principal_id or "").strip() or None,
        "principal_email": normalized_email,
        "principal_display_name": str(principal_display_name or "").strip() or None,
        "pay_basis": snapshot["pay_basis"],
        "base_wage_input_type": snapshot["base_wage_input_type"],
        "base_wage_amount": float(snapshot["base_wage_amount"]),
        "weekly_basis_hours": float(snapshot["weekly_basis_hours"]),
        "commission_month_1_amount": float(snapshot["commission_month_1_amount"]),
        "commission_month_2_amount": float(snapshot["commission_month_2_amount"]),
        "commission_month_3_amount": float(snapshot["commission_month_3_amount"]),
        "commission_average_monthly": float(snapshot["commission_average_monthly"]),
        "commission_hourly_rate": float(snapshot["commission_hourly_rate"]),
        "calculated_hourly_rate": float(snapshot["calculated_hourly_rate"]),
        "default_address": str(default_address or "").strip() or None,
        "status": normalized_status,
        "notes": str(notes or "").strip() or None,
    }
    await db.exec(
        """UPDATE pay_profile_change_requests
              SET status='superseded', reviewed_by=?, reviewed_at_utc=?, review_note=?
            WHERE principal_email=? AND status='pending'""",
        (requested_by, now, "Replaced by a newer wage change request.", normalized_email),
    )
    await db.exec(
        """INSERT INTO pay_profile_change_requests(
             id, principal_id, principal_email, principal_display_name, pay_basis,
             base_wage_input_type, base_wage_amount, weekly_basis_hours,
             commission_month_1_amount, commission_month_2_amount, commission_month_3_amount,
             commission_average_monthly, commission_hourly_rate, calculated_hourly_rate,
             default_address, profile_status, notes, requested_by, requested_at_utc, status,
             current_profile_json, requested_profile_json
           ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            request_id,
            requested_profile["principal_id"],
            normalized_email,
            requested_profile["principal_display_name"],
            requested_profile["pay_basis"],
            requested_profile["base_wage_input_type"],
            requested_profile["base_wage_amount"],
            requested_profile["weekly_basis_hours"],
            requested_profile["commission_month_1_amount"],
            requested_profile["commission_month_2_amount"],
            requested_profile["commission_month_3_amount"],
            requested_profile["commission_average_monthly"],
            requested_profile["commission_hourly_rate"],
            requested_profile["calculated_hourly_rate"],
            requested_profile["default_address"],
            requested_profile["status"],
            requested_profile["notes"],
            requested_by,
            now,
            "pending",
            json.dumps(existing or {}, ensure_ascii=False),
            json.dumps(requested_profile, ensure_ascii=False),
        ),
    )
    await add_pay_event(
        db,
        period_id=None,
        event_type="pay_profile_change_requested",
        actor=requested_by,
        details={"request_id": request_id, "principal_email": normalized_email},
    )
    rows = await list_pay_profile_change_requests(db, email=normalized_email, pending_only=True)
    return next(row for row in rows if row["id"] == request_id)


async def review_pay_profile_change_request(
    db: Db,
    *,
    request_id: str,
    actor: str,
    approved: bool,
    review_note: str | None = None,
) -> dict[str, object]:
    row = await db.fetchone(f"{_PAY_PROFILE_CHANGE_SELECT} WHERE id=?", (request_id,))
    if not row:
        raise ValueError("pay profile change request not found")
    request = _profile_change_request_from_row(row)
    if str(request.get("status") or "") != "pending":
        raise ValueError("pay profile change request is not pending")
    now = utcnow()
    if approved:
        saved = await upsert_pay_profile(
            db,
            principal_id=str(request.get("principal_id") or "").strip() or None,
            principal_email=str(request["principal_email"]),
            principal_display_name=str(request.get("principal_display_name") or "").strip() or None,
            pay_basis=str(request["pay_basis"]),
            base_wage_input_type=str(request["base_wage_input_type"]),
            base_wage_amount=float(request.get("base_wage_amount") or 0),
            weekly_basis_hours=float(request.get("weekly_basis_hours") or 40),
            commission_month_1_amount=float(request.get("commission_month_1_amount") or 0),
            commission_month_2_amount=float(request.get("commission_month_2_amount") or 0),
            commission_month_3_amount=float(request.get("commission_month_3_amount") or 0),
            status=str(request.get("profile_status") or "active"),
            notes=request.get("notes"),
            updated_by=actor,
            default_address=request.get("default_address"),
        )
        status_value = "approved"
    else:
        saved = None
        status_value = "rejected"
    await db.exec(
        """UPDATE pay_profile_change_requests
              SET status=?, reviewed_by=?, reviewed_at_utc=?, review_note=?
            WHERE id=?""",
        (status_value, actor, now, str(review_note or "").strip() or None, request_id),
    )
    await add_pay_event(
        db,
        period_id=None,
        event_type="pay_profile_change_approved" if approved else "pay_profile_change_rejected",
        actor=actor,
        details={"request_id": request_id, "principal_email": request["principal_email"], "review_note": review_note or ""},
    )
    updated_row = await db.fetchone(f"{_PAY_PROFILE_CHANGE_SELECT} WHERE id=?", (request_id,))
    result = _profile_change_request_from_row(updated_row)
    if saved is not None:
        result["saved_profile"] = saved
    return result


async def upsert_pay_profile(
    db: Db,
    *,
    principal_id: str | None,
    principal_email: str,
    principal_display_name: str | None,
    pay_basis: str,
    base_wage_input_type: str,
    base_wage_amount: float,
    weekly_basis_hours: float,
    commission_month_1_amount: float,
    commission_month_2_amount: float,
    commission_month_3_amount: float,
    status: str,
    notes: str | None,
    updated_by: str,
    default_address: str | None = None,
) -> dict[str, object]:
    normalized_email = normalize_email(principal_email)
    if not normalized_email:
        raise ValueError("principal_email is required")
    normalized_status = normalize_pay_profile_status(status)
    snapshot = calculate_pay_profile_snapshot(
        pay_basis=pay_basis,
        base_wage_input_type=base_wage_input_type,
        base_wage_amount=base_wage_amount,
        weekly_basis_hours=weekly_basis_hours,
        commission_month_1_amount=commission_month_1_amount,
        commission_month_2_amount=commission_month_2_amount,
        commission_month_3_amount=commission_month_3_amount,
    )
    if snapshot["pay_basis"] != "expense_only" and Decimal(str(snapshot["calculated_hourly_rate"])) <= 0:
        raise ValueError("base wage amount is required for this pay profile")
    now = utcnow()
    await db.exec(
        """
        INSERT INTO pay_profiles(
          principal_id, principal_email, principal_display_name, pay_basis,
          base_wage_input_type, base_wage_amount, weekly_basis_hours,
          commission_month_1_amount, commission_month_2_amount, commission_month_3_amount,
          commission_average_monthly, commission_hourly_rate, calculated_hourly_rate,
          default_address, status, notes, created_at_utc, updated_at_utc, updated_by
        ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        ON CONFLICT(principal_email) DO UPDATE SET
          principal_id=excluded.principal_id,
          principal_display_name=excluded.principal_display_name,
          pay_basis=excluded.pay_basis,
          base_wage_input_type=excluded.base_wage_input_type,
          base_wage_amount=excluded.base_wage_amount,
          weekly_basis_hours=excluded.weekly_basis_hours,
          commission_month_1_amount=excluded.commission_month_1_amount,
          commission_month_2_amount=excluded.commission_month_2_amount,
          commission_month_3_amount=excluded.commission_month_3_amount,
          commission_average_monthly=excluded.commission_average_monthly,
          commission_hourly_rate=excluded.commission_hourly_rate,
          calculated_hourly_rate=excluded.calculated_hourly_rate,
          default_address=excluded.default_address,
          status=excluded.status,
          notes=excluded.notes,
          updated_at_utc=excluded.updated_at_utc,
          updated_by=excluded.updated_by
        """,
        (
            str(principal_id or "").strip() or None,
            normalized_email,
            str(principal_display_name or "").strip() or None,
            snapshot["pay_basis"],
            snapshot["base_wage_input_type"],
            float(snapshot["base_wage_amount"]),
            float(snapshot["weekly_basis_hours"]),
            float(snapshot["commission_month_1_amount"]),
            float(snapshot["commission_month_2_amount"]),
            float(snapshot["commission_month_3_amount"]),
            float(snapshot["commission_average_monthly"]),
            float(snapshot["commission_hourly_rate"]),
            float(snapshot["calculated_hourly_rate"]),
            str(default_address or "").strip() or None,
            normalized_status,
            str(notes or "").strip() or None,
            now,
            now,
            updated_by,
        ),
    )
    saved = await pay_profile_by_email(db, email=normalized_email)
    if not saved:
        raise RuntimeError("failed to save pay profile")
    return saved


async def list_wage_scales(db: Db) -> list[dict[str, object]]:
    rows = await db.fetchall(
        """SELECT id, effective_date, weekly_basis_hours, target_scale, actual_scale,
                  target_weekly_amount, actual_weekly_amount, target_multiplier, notes,
                  updated_at_utc, updated_by
           FROM pay_wage_scales
           ORDER BY effective_date DESC, weekly_basis_hours"""
    )
    return [
        {
            "id": int(row[0]),
            "effective_date": row[1],
            "weekly_basis_hours": float(row[2]),
            "target_scale": row[3],
            "actual_scale": row[4],
            "target_weekly_amount": row[5],
            "actual_weekly_amount": row[6],
            "target_multiplier": row[7],
            "notes": row[8],
            "updated_at_utc": row[9],
            "updated_by": row[10],
        }
        for row in rows
    ]


def _compensation_stub_from_row(row: Any) -> dict[str, object]:
    return {
        "id": row[0],
        "user_email": row[1],
        "uploaded_by": row[2],
        "base_wage_input_type": row[3],
        "base_wage_amount": row[4],
        "weekly_basis_hours": row[5],
        "commission_month_1_amount": row[6],
        "commission_month_2_amount": row[7],
        "commission_month_3_amount": row[8],
        "commission_average_monthly": row[9],
        "commission_hourly_rate": row[10],
        "calculated_hourly_rate": row[11],
        "payroll_month": row[12],
        "filename": row[13],
        "content_type": row[14],
        "size_bytes": row[15],
        "sha256": row[16],
        "scan_status": row[17],
        "sharepoint_url": row[18],
        "notes": row[19],
        "created_at_utc": row[20],
    }


async def list_compensation_stubs(db: Db, *, actor: PayActor) -> list[dict[str, object]]:
    select_sql = """SELECT id, user_email, uploaded_by, base_wage_input_type, base_wage_amount,
                           weekly_basis_hours, commission_month_1_amount, commission_month_2_amount,
                           commission_month_3_amount, commission_average_monthly,
                           commission_hourly_rate, calculated_hourly_rate, payroll_month, original_filename,
                           content_type, size_bytes, sha256, scan_status, sharepoint_url,
                           notes, created_at_utc
                    FROM pay_compensation_stubs"""
    if actor.can_view_all:
        rows = await db.fetchall(f"{select_sql} ORDER BY user_email, created_at_utc DESC")
    else:
        rows = await db.fetchall(
            f"{select_sql} WHERE user_email=? ORDER BY created_at_utc DESC",
            (actor.email,),
        )
    return [_compensation_stub_from_row(row) for row in rows]


async def latest_compensation_stub(db: Db, *, user_email: str) -> dict[str, object] | None:
    row = await db.fetchone(
        """SELECT id, user_email, uploaded_by, base_wage_input_type, base_wage_amount,
                  weekly_basis_hours, commission_month_1_amount, commission_month_2_amount,
                  commission_month_3_amount, commission_average_monthly,
                  commission_hourly_rate, calculated_hourly_rate, payroll_month, original_filename,
                  content_type, size_bytes, sha256, scan_status, sharepoint_url,
                  notes, created_at_utc
           FROM pay_compensation_stubs
           WHERE user_email=?
           ORDER BY created_at_utc DESC, id DESC
           LIMIT 1""",
        (normalize_email(user_email),),
    )
    return _compensation_stub_from_row(row) if row else None


async def compensation_stub_for_payroll_month(
    db: Db,
    *,
    user_email: str,
    payroll_month: str,
) -> dict[str, object] | None:
    normalized_month = normalize_payroll_month(payroll_month)
    row = await db.fetchone(
        """SELECT id, user_email, uploaded_by, base_wage_input_type, base_wage_amount,
                  weekly_basis_hours, commission_month_1_amount, commission_month_2_amount,
                  commission_month_3_amount, commission_average_monthly,
                  commission_hourly_rate, calculated_hourly_rate, payroll_month, original_filename,
                  content_type, size_bytes, sha256, scan_status, sharepoint_url,
                  notes, created_at_utc
           FROM pay_compensation_stubs
           WHERE user_email=? AND payroll_month=?
           ORDER BY created_at_utc DESC, id DESC
           LIMIT 1""",
        (normalize_email(user_email), normalized_month),
    )
    return _compensation_stub_from_row(row) if row else None


async def upsert_wage_scale(
    db: Db,
    *,
    effective_date: str,
    weekly_basis_hours: float,
    target_weekly_amount: float,
    actual_weekly_amount: float | None,
    target_multiplier: float,
    target_scale: str,
    actual_scale: str,
    notes: str | None,
    updated_by: str,
) -> dict[str, object]:
    now = utcnow()
    await db.exec(
        """
        INSERT INTO pay_wage_scales(
          effective_date, weekly_basis_hours, target_scale, actual_scale,
          target_weekly_amount, actual_weekly_amount, target_multiplier,
          notes, created_at_utc, updated_at_utc, updated_by
        ) VALUES(?,?,?,?,?,?,?,?,?,?,?)
        ON CONFLICT(effective_date, weekly_basis_hours, target_scale, actual_scale) DO UPDATE SET
          target_weekly_amount=excluded.target_weekly_amount,
          actual_weekly_amount=excluded.actual_weekly_amount,
          target_multiplier=excluded.target_multiplier,
          notes=excluded.notes,
          updated_at_utc=excluded.updated_at_utc,
          updated_by=excluded.updated_by
        """,
        (
            effective_date,
            float(weekly_basis_hours),
            str(target_scale or "36"),
            str(actual_scale or "base"),
            float(target_weekly_amount),
            None if actual_weekly_amount is None else float(actual_weekly_amount),
            float(target_multiplier),
            notes,
            now,
            now,
            updated_by,
        ),
    )
    row = await db.fetchone(
        """SELECT id FROM pay_wage_scales
           WHERE effective_date=? AND weekly_basis_hours=? AND target_scale=? AND actual_scale=?""",
        (effective_date, float(weekly_basis_hours), str(target_scale or "36"), str(actual_scale or "base")),
    )
    return {"id": int(row[0]) if row else None}


async def president_week_scheduled_hours(
    db: Db,
    *,
    period_start: date,
    entry_date: date,
    period_id: str,
    user_email: str,
    exclude_entry_date: str | None = None,
) -> Decimal:
    week_start = period_start if (entry_date - period_start).days < 7 else period_start + timedelta(days=7)
    week_end = min(week_start + timedelta(days=6), period_start + timedelta(days=13))
    params: list[object] = [period_id, user_email, week_start.isoformat(), week_end.isoformat()]
    exclude_clause = ""
    if exclude_entry_date:
        exclude_clause = " AND entry_date <> ?"
        params.append(exclude_entry_date)
    row = await db.fetchone(
        f"""SELECT COALESCE(SUM(hours + president_diff_hours), 0)
             FROM pay_entries
             WHERE period_id=? AND user_email=? AND entry_date BETWEEN ? AND ?{exclude_clause}""",
        tuple(params),
    )
    return _quantity(row[0] if row else 0)


def president_daily_differential_hours(*, union_hours: Decimal, requested_diff_hours: Decimal) -> Decimal:
    daily_cap = Decimal("8.00")
    if union_hours < 0 or requested_diff_hours < 0:
        raise ValueError("president hours cannot be negative")
    if union_hours == 0 and requested_diff_hours == 0:
        return Decimal("0.00")
    if union_hours > daily_cap:
        raise ValueError("president union hours cannot exceed 8 scheduled hours in a day")
    return _quantity(daily_cap - union_hours)


def validate_president_week_scheduled_hours(
    *,
    existing_week_hours: Decimal,
    union_hours: Decimal,
    requested_diff_hours: Decimal,
) -> Decimal:
    cap = Decimal("40.00")
    diff_hours = president_daily_differential_hours(
        union_hours=union_hours,
        requested_diff_hours=requested_diff_hours,
    )
    if existing_week_hours + union_hours + diff_hours > cap:
        raise ValueError("president scheduled lost-wage hours cannot exceed 40 hours in a week")
    return diff_hours


async def calculate_president_differential(
    db: Db,
    *,
    entry_date: str,
    weekly_basis_hours: float,
    president_diff_hours: object,
    target_scale: str,
    target_multiplier: float,
    lost_wage_input_type: object,
    lost_wage_amount: object,
) -> DifferentialResult:
    hours = _quantity(president_diff_hours)
    _, _, lost_wage_hourly = normalize_wage_input(
        input_type=lost_wage_input_type,
        amount=lost_wage_amount,
        weekly_basis_hours=weekly_basis_hours,
    )
    if lost_wage_hourly <= 0:
        return DifferentialResult(
            wage_scale_id=None,
            diff_rate=Decimal("0.00"),
            diff_amount=Decimal("0.00"),
            lost_wage_hourly_rate=lost_wage_hourly,
            presidential_hourly_rate=lost_wage_hourly,
        )

    row = await db.fetchone(
        """SELECT id, weekly_basis_hours, target_weekly_amount, target_multiplier
           FROM pay_wage_scales
           WHERE effective_date <= ?
             AND ABS(weekly_basis_hours - ?) < 0.0001
             AND target_scale=?
           ORDER BY effective_date DESC, id DESC
           LIMIT 1""",
        (entry_date, float(weekly_basis_hours), str(target_scale)),
    )
    if not row:
        return DifferentialResult(
            wage_scale_id=None,
            diff_rate=Decimal("0.00"),
            diff_amount=Decimal("0.00"),
            lost_wage_hourly_rate=lost_wage_hourly,
            presidential_hourly_rate=lost_wage_hourly,
        )

    basis = Decimal(str(row[1]))
    if basis <= 0:
        return DifferentialResult(
            wage_scale_id=int(row[0]),
            diff_rate=Decimal("0.00"),
            diff_amount=Decimal("0.00"),
            lost_wage_hourly_rate=lost_wage_hourly,
            presidential_hourly_rate=lost_wage_hourly,
        )
    # Wage scale rows store the scale 36 base; the presidential target is scale 36 plus 20%.
    multiplier = Decimal(str(row[3] if row[3] is not None else target_multiplier))
    target_hourly = ((Decimal(str(row[2])) / basis) * multiplier).quantize(Decimal("1"), rounding=ROUND_CEILING)
    diff_rate = max(target_hourly - lost_wage_hourly, Decimal("0")).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
    diff_amount = (diff_rate * hours).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
    return DifferentialResult(
        wage_scale_id=int(row[0]),
        diff_rate=diff_rate,
        diff_amount=diff_amount,
        lost_wage_hourly_rate=lost_wage_hourly,
        presidential_hourly_rate=target_hourly,
    )


async def list_entries(db: Db, *, period_id: str, actor: PayActor) -> list[dict[str, object]]:
    if actor.can_view_all:
        rows = await db.fetchall(
            """SELECT id, user_email, display_name, entry_date, local_number, address, hourly_rate,
                      lost_wage_input_type, lost_wage_amount, lost_wage_hourly_rate,
                      compensation_stub_id, hours, mileage_miles, mileage_rate, mileage_amount,
                      rentals_amount, meals_amount, hotel_amount, miscellaneous_amount,
                      president_diff_hours, president_diff_rate, president_diff_amount,
                      wage_scale_id, notes, review_status, review_note, reviewed_by, reviewed_at_utc,
                      submitter_certified_at_utc, submitter_certified_by, submitter_certification_text,
                      locked_at_utc, created_at_utc, updated_at_utc,
                      (SELECT COUNT(*) FROM pay_entry_corrections c WHERE c.entry_id=pay_entries.id) AS correction_count
               FROM pay_entries
               WHERE period_id=?
               ORDER BY entry_date, user_email""",
            (period_id,),
        )
    else:
        rows = await db.fetchall(
            """SELECT id, user_email, display_name, entry_date, local_number, address, hourly_rate,
                      lost_wage_input_type, lost_wage_amount, lost_wage_hourly_rate,
                      compensation_stub_id, hours, mileage_miles, mileage_rate, mileage_amount,
                      rentals_amount, meals_amount, hotel_amount, miscellaneous_amount,
                      president_diff_hours, president_diff_rate, president_diff_amount,
                      wage_scale_id, notes, review_status, review_note, reviewed_by, reviewed_at_utc,
                      submitter_certified_at_utc, submitter_certified_by, submitter_certification_text,
                      locked_at_utc, created_at_utc, updated_at_utc,
                      (SELECT COUNT(*) FROM pay_entry_corrections c WHERE c.entry_id=pay_entries.id) AS correction_count
               FROM pay_entries
               WHERE period_id=? AND user_email=?
               ORDER BY entry_date, user_email""",
            (period_id, actor.email),
        )
    entries: list[dict[str, object]] = []
    for row in rows:
        entries.append(
            {
                "id": row[0],
                "user_email": row[1],
                "display_name": row[2],
                "entry_date": row[3],
                "local_number": row[4],
                "address": row[5],
                "hourly_rate": row[6],
                "lost_wage_input_type": row[7],
                "lost_wage_amount": row[8],
                "lost_wage_hourly_rate": row[9],
                "compensation_stub_id": row[10],
                "hours": row[11],
                "mileage_miles": row[12],
                "mileage_rate": row[13],
                "mileage_amount": row[14],
                "rentals_amount": row[15],
                "meals_amount": row[16],
                "hotel_amount": row[17],
                "miscellaneous_amount": row[18],
                "president_diff_hours": row[19],
                "president_diff_rate": row[20],
                "president_diff_amount": row[21],
                "wage_scale_id": row[22],
                "notes": row[23],
                "review_status": row[24] or "pending",
                "review_note": row[25],
                "reviewed_by": row[26],
                "reviewed_at_utc": row[27],
                "submitter_certified_at_utc": row[28],
                "submitter_certified_by": row[29],
                "submitter_certification_text": row[30],
                "locked_at_utc": row[31],
                "created_at_utc": row[32],
                "updated_at_utc": row[33],
                "correction_count": int(row[34] or 0),
            }
        )
    return entries


async def upsert_entry(
    db: Db,
    *,
    period_id: str,
    actor: PayActor,
    data: dict[str, object],
    pay_cfg: Any,
    require_submitter_certification: bool = False,
) -> dict[str, object]:
    period = await get_pay_period(db, period_id)
    if not period:
        raise ValueError("pay period not found")
    if str(period["status"]) != "open":
        raise ValueError("pay period is locked")

    entry_date = str(data.get("entry_date") or "").strip()
    if not entry_date:
        raise ValueError("entry_date is required")
    period_start = date.fromisoformat(str(period["period_start"]))
    period_end = date.fromisoformat(str(period["period_end"]))
    parsed_entry_date = date.fromisoformat(entry_date)
    if parsed_entry_date < period_start or parsed_entry_date > period_end:
        raise ValueError("entry_date is outside the pay period")

    target_email = (normalize_email(data.get("user_email")) or actor.email) if actor.can_edit_all else actor.email
    if not target_email:
        raise ValueError("user_email is required")

    hours = _quantity(data.get("hours"))
    requested_president_diff_hours = _quantity(data.get("president_diff_hours"))
    profile = await pay_profile_by_email(db, email=target_email, active_only=True)
    compensation_stub_id: str | None = None
    weekly_basis_hours = 40.0
    normalized_wage_type = "expense_only"
    lost_wage_amount = Decimal("0.00")
    lost_wage_hourly_rate = Decimal("0.00")
    hourly_rate = Decimal("0.00")
    president_diff_hours = Decimal("0.0000")
    diff = DifferentialResult(
        wage_scale_id=None,
        diff_rate=Decimal("0.00"),
        diff_amount=Decimal("0.00"),
        lost_wage_hourly_rate=Decimal("0.00"),
        presidential_hourly_rate=Decimal("0.00"),
    )

    if not profile:
        if hours > 0 or requested_president_diff_hours > 0:
            raise ValueError("pay profile required before submitting lost-wage hours")
    else:
        pay_basis = str(profile["pay_basis"] or "expense_only")
        weekly_basis_hours = float(profile["weekly_basis_hours"] or 40.0)
        profile_hourly = _money(profile.get("calculated_hourly_rate"))
        if pay_basis == "expense_only":
            if hours > 0 or requested_president_diff_hours > 0:
                raise ValueError("pay profile is expense-only; lost-wage hours are not allowed")
        elif pay_basis == "commission" and hours > 0:
            required_payroll_month = required_commission_payroll_month(parsed_entry_date)
            compensation_stub = await compensation_stub_for_payroll_month(
                db,
                user_email=target_email,
                payroll_month=required_payroll_month,
            )
            if not compensation_stub:
                raise ValueError(
                    f"commission payroll proof for {required_payroll_month} is required before submitting lost-wage hours"
                )
            compensation_stub_id = str(compensation_stub["id"])
            weekly_basis_hours = float(compensation_stub.get("weekly_basis_hours") or weekly_basis_hours)
            profile_hourly = _money(compensation_stub.get("calculated_hourly_rate"))
            if profile_hourly <= 0:
                raise ValueError("commission payroll proof must calculate a positive wage rate")
        elif (hours > 0 or requested_president_diff_hours > 0) and profile_hourly <= 0:
            raise ValueError("pay profile must have a positive wage rate before submitting lost-wage hours")
        if pay_basis != "expense_only":
            normalized_wage_type = pay_basis
            hourly_rate = profile_hourly
            lost_wage_hourly_rate = profile_hourly
            if pay_basis == "commission":
                lost_wage_amount = profile_hourly
            else:
                lost_wage_amount = _money(profile.get("base_wage_amount"))
            if pay_basis == "president":
                can_enter_president_diff = actor.can_edit_all or actor.can_lock or actor.is_president or actor.email == target_email
                effective_requested_diff_hours = requested_president_diff_hours if can_enter_president_diff else Decimal("0.00")
                existing_week_hours = await president_week_scheduled_hours(
                    db,
                    period_start=period_start,
                    entry_date=parsed_entry_date,
                    period_id=period_id,
                    user_email=target_email,
                    exclude_entry_date=entry_date,
                )
                president_diff_hours = validate_president_week_scheduled_hours(
                    existing_week_hours=existing_week_hours,
                    union_hours=hours,
                    requested_diff_hours=effective_requested_diff_hours,
                )
                diff = await calculate_president_differential(
                    db,
                    entry_date=entry_date,
                    weekly_basis_hours=weekly_basis_hours,
                    president_diff_hours=president_diff_hours,
                    target_scale=pay_cfg.president_target_scale,
                    target_multiplier=pay_cfg.president_target_multiplier,
                    lost_wage_input_type=profile.get("base_wage_input_type"),
                    lost_wage_amount=profile.get("base_wage_amount"),
                )
                lost_wage_hourly_rate = diff.presidential_hourly_rate
                hourly_rate = diff.presidential_hourly_rate

    row = await db.fetchone(
        "SELECT id, locked_at_utc FROM pay_entries WHERE period_id=? AND user_email=? AND entry_date=?",
        (period_id, target_email, entry_date),
    )
    if row and row[1]:
        raise ValueError("entry is locked")
    entry_id = str(row[0]) if row else f"pay-entry-{uuid4().hex}"
    now = utcnow()
    submitted_display_name = str(data.get("display_name") or "").strip() if actor.can_edit_all else ""
    submitted_address = str(data.get("address") or "").strip()
    profile_default_address = str((profile or {}).get("default_address") or "").strip()
    entry_address = submitted_address or profile_default_address

    certify_value = data.get("submitter_certified")
    submitter_certified = certify_value is True or str(certify_value or "").strip().lower() in {"1", "true", "yes", "on"}
    certification_text = str(
        data.get("submitter_certification_text")
        or "I certify that this daily lost-wage and expense entry is accurate and was submitted by me."
    ).strip()
    if require_submitter_certification and not submitter_certified and not actor.can_lock:
        raise ValueError("submitter certification is required before saving a pay entry")
    certified_at = now if submitter_certified else None
    certified_by = actor.email if submitter_certified else None

    values = {
        "display_name": str(
            submitted_display_name
            or (profile or {}).get("principal_display_name")
            or actor.display_name
            or target_email
        ).strip(),
        "local_number": str(data.get("local_number") or "").strip(),
        "address": entry_address,
        "hourly_rate": float(hourly_rate),
        "lost_wage_input_type": normalized_wage_type,
        "lost_wage_amount": float(lost_wage_amount),
        "lost_wage_hourly_rate": float(lost_wage_hourly_rate),
        "compensation_stub_id": compensation_stub_id,
        "hours": float(hours),
        "mileage_miles": float(_quantity(data.get("mileage_miles"))),
        "mileage_rate": float(_mileage_rate(data.get("mileage_rate"))),
        "mileage_amount": float(_money(data.get("mileage_amount"))),
        "rentals_amount": float(_money(data.get("rentals_amount"))),
        "meals_amount": float(_money(data.get("meals_amount"))),
        "hotel_amount": float(_money(data.get("hotel_amount"))),
        "miscellaneous_amount": float(_money(data.get("miscellaneous_amount"))),
        "president_diff_hours": float(president_diff_hours),
        "president_diff_rate": float(diff.diff_rate),
        "president_diff_amount": float(diff.diff_amount),
        "wage_scale_id": diff.wage_scale_id,
        "notes": str(data.get("notes") or "").strip(),
        "submitter_certified_at_utc": certified_at,
        "submitter_certified_by": certified_by,
        "submitter_certification_text": certification_text if submitter_certified else None,
    }
    await db.exec(
        """
        INSERT INTO pay_entries(
          id, period_id, user_email, display_name, entry_date, local_number, address,
          hourly_rate, lost_wage_input_type, lost_wage_amount, lost_wage_hourly_rate,
          compensation_stub_id, hours, mileage_miles, mileage_rate, mileage_amount,
          rentals_amount, meals_amount, hotel_amount, miscellaneous_amount,
          president_diff_hours, president_diff_rate, president_diff_amount, wage_scale_id,
          notes, submitter_certified_at_utc, submitter_certified_by, submitter_certification_text,
          created_at_utc, updated_at_utc
        ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        ON CONFLICT(period_id, user_email, entry_date) DO UPDATE SET
          display_name=excluded.display_name,
          local_number=excluded.local_number,
          address=excluded.address,
          hourly_rate=excluded.hourly_rate,
          lost_wage_input_type=excluded.lost_wage_input_type,
          lost_wage_amount=excluded.lost_wage_amount,
          lost_wage_hourly_rate=excluded.lost_wage_hourly_rate,
          compensation_stub_id=excluded.compensation_stub_id,
          hours=excluded.hours,
          mileage_miles=excluded.mileage_miles,
          mileage_rate=excluded.mileage_rate,
          mileage_amount=excluded.mileage_amount,
          rentals_amount=excluded.rentals_amount,
          meals_amount=excluded.meals_amount,
          hotel_amount=excluded.hotel_amount,
          miscellaneous_amount=excluded.miscellaneous_amount,
          president_diff_hours=excluded.president_diff_hours,
          president_diff_rate=excluded.president_diff_rate,
          president_diff_amount=excluded.president_diff_amount,
          wage_scale_id=excluded.wage_scale_id,
          notes=excluded.notes,
          submitter_certified_at_utc=excluded.submitter_certified_at_utc,
          submitter_certified_by=excluded.submitter_certified_by,
          submitter_certification_text=excluded.submitter_certification_text,
          review_status='pending',
          updated_at_utc=excluded.updated_at_utc
        """,
        (
            entry_id,
            period_id,
            target_email,
            values["display_name"],
            entry_date,
            values["local_number"],
            values["address"],
            values["hourly_rate"],
            values["lost_wage_input_type"],
            values["lost_wage_amount"],
            values["lost_wage_hourly_rate"],
            values["compensation_stub_id"],
            values["hours"],
            values["mileage_miles"],
            values["mileage_rate"],
            values["mileage_amount"],
            values["rentals_amount"],
            values["meals_amount"],
            values["hotel_amount"],
            values["miscellaneous_amount"],
            values["president_diff_hours"],
            values["president_diff_rate"],
            values["president_diff_amount"],
            values["wage_scale_id"],
            values["notes"],
            values["submitter_certified_at_utc"],
            values["submitter_certified_by"],
            values["submitter_certification_text"],
            now,
            now,
        ),
    )
    if profile and submitted_address and not profile_default_address:
        await db.exec(
            "UPDATE pay_profiles SET default_address=?, updated_at_utc=?, updated_by=? WHERE principal_email=?",
            (submitted_address, now, actor.email, target_email),
        )

    await add_pay_event(
        db,
        period_id=period_id,
        entry_id=entry_id,
        event_type="entry_upserted",
        actor=actor.email,
        details={"entry_date": entry_date, "user_email": target_email},
    )
    entries = await list_entries(db, period_id=period_id, actor=PayActor(target_email, None, "guest", True, True, False))
    return next(row for row in entries if row["id"] == entry_id)



async def delete_pay_entry(
    db: Db,
    *,
    entry_id: str,
    actor: PayActor,
) -> dict[str, object]:
    row = await db.fetchone(
        """SELECT e.id, e.period_id, e.user_email, e.display_name, e.entry_date, e.locked_at_utc,
                  p.status, e.hours, e.mileage_miles, e.mileage_amount,
                  e.rentals_amount, e.meals_amount, e.hotel_amount, e.miscellaneous_amount,
                  e.president_diff_hours, e.president_diff_amount, e.notes,
                  (SELECT COUNT(*) FROM pay_attachments a WHERE a.entry_id=e.id) AS attachment_count,
                  (SELECT COUNT(*) FROM pay_entry_corrections c WHERE c.entry_id=e.id) AS correction_count
           FROM pay_entries e
           JOIN pay_periods p ON p.id=e.period_id
           WHERE e.id=?""",
        (entry_id,),
    )
    if not row:
        raise ValueError("entry not found")
    owner_email = normalize_email(row[2])
    if not (actor.can_edit_all or actor.can_lock or owner_email == actor.email):
        raise PermissionError("cannot delete another user's entry")
    if row[5] or str(row[6] or "") != "open":
        raise ValueError("entry is locked")
    snapshot = {
        "entry_id": row[0],
        "period_id": row[1],
        "user_email": row[2],
        "display_name": row[3],
        "entry_date": row[4],
        "hours": row[7],
        "mileage_miles": row[8],
        "mileage_amount": row[9],
        "rentals_amount": row[10],
        "meals_amount": row[11],
        "hotel_amount": row[12],
        "miscellaneous_amount": row[13],
        "president_diff_hours": row[14],
        "president_diff_amount": row[15],
        "notes": row[16],
        "attachment_count": int(row[17] or 0),
        "correction_count": int(row[18] or 0),
    }
    attachment_rows = await db.fetchall("SELECT id, local_path FROM pay_attachments WHERE entry_id=?", (entry_id,))
    await db.exec("UPDATE pay_events SET entry_id=NULL WHERE entry_id=?", (entry_id,))
    await db.exec("DELETE FROM pay_entry_corrections WHERE entry_id=?", (entry_id,))
    await db.exec("DELETE FROM pay_attachments WHERE entry_id=?", (entry_id,))
    await db.exec("DELETE FROM pay_entries WHERE id=?", (entry_id,))
    removed_files: list[str] = []
    for attachment_id, local_path in attachment_rows:
        try:
            path = Path(str(local_path or ""))
            if path.exists() and path.is_file():
                path.unlink()
                removed_files.append(str(path))
        except OSError:
            pass
    snapshot["removed_attachment_files"] = removed_files
    await add_pay_event(
        db,
        period_id=str(row[1]),
        event_type="entry_deleted",
        actor=actor.email,
        details=snapshot,
    )
    return {"ok": True, **snapshot}


_PAY_REVIEW_STATUSES = {"pending", "approved", "needs_fix", "rejected"}


def normalize_pay_review_status(value: object) -> str:
    status = str(value or "pending").strip().lower()
    if status not in _PAY_REVIEW_STATUSES:
        raise ValueError("review_status must be pending, approved, needs_fix, or rejected")
    return status


async def review_pay_entry(
    db: Db,
    *,
    entry_id: str,
    actor: PayActor,
    review_status: str,
    review_note: str | None = None,
) -> dict[str, object]:
    if not actor.can_lock:
        raise PermissionError("treasurer access required")
    normalized_status = normalize_pay_review_status(review_status)
    note = str(review_note or "").strip() or None
    if normalized_status in {"needs_fix", "rejected"} and not note:
        raise ValueError("review_note is required")
    row = await db.fetchone("SELECT period_id FROM pay_entries WHERE id=?", (entry_id,))
    if not row:
        raise ValueError("entry not found")
    now = utcnow()
    await db.exec(
        """UPDATE pay_entries
           SET review_status=?, review_note=?, reviewed_by=?, reviewed_at_utc=?, updated_at_utc=?
           WHERE id=?""",
        (normalized_status, note, actor.email, now, now, entry_id),
    )
    await add_pay_event(
        db,
        period_id=str(row[0]),
        entry_id=entry_id,
        event_type="entry_reviewed",
        actor=actor.email,
        details={"review_status": normalized_status, "review_note": note},
    )
    entries = await list_entries(db, period_id=str(row[0]), actor=PayActor(actor.email, actor.display_name, actor.role, True, True, True))
    return next(item for item in entries if item["id"] == entry_id)


async def create_pay_entry_correction(
    db: Db,
    *,
    period_id: str,
    actor: PayActor,
    data: dict[str, object],
    pay_cfg: Any,
) -> dict[str, object]:
    if not actor.can_lock:
        raise PermissionError("treasurer access required")
    target_email = normalize_email(data.get("user_email"))
    if not target_email:
        raise ValueError("user_email is required")
    entry_date = str(data.get("entry_date") or "").strip()
    if not entry_date:
        raise ValueError("entry_date is required")
    amount_fields = {
        "hours": _quantity(data.get("hours")),
        "mileage_miles": _quantity(data.get("mileage_miles")),
        "mileage_rate": _quantity(data.get("mileage_rate")),
        "mileage_amount": _money(data.get("mileage_amount")),
        "rentals_amount": _money(data.get("rentals_amount")),
        "meals_amount": _money(data.get("meals_amount")),
        "hotel_amount": _money(data.get("hotel_amount")),
        "miscellaneous_amount": _money(data.get("miscellaneous_amount")),
    }
    for key, value in amount_fields.items():
        if value < 0:
            raise ValueError("correction values cannot be negative")
    if not any(value > 0 for key, value in amount_fields.items() if key != "mileage_rate"):
        raise ValueError("correction must add at least one amount")
    existing = await db.fetchone(
        "SELECT id FROM pay_entries WHERE period_id=? AND user_email=? AND entry_date=?",
        (period_id, target_email, entry_date),
    )
    if not existing:
        await upsert_entry(
            db,
            period_id=period_id,
            actor=PayActor(actor.email, actor.display_name, actor.role, True, True, True),
            data={
                "user_email": target_email,
                "display_name": data.get("display_name") or target_email,
                "entry_date": entry_date,
                "local_number": data.get("local_number") or "3106",
                "address": data.get("address") or "",
                "notes": data.get("notes") or "Treasurer correction entry",
            },
            pay_cfg=pay_cfg,
        )
        existing = await db.fetchone(
            "SELECT id FROM pay_entries WHERE period_id=? AND user_email=? AND entry_date=?",
            (period_id, target_email, entry_date),
        )
    if not existing:
        raise RuntimeError("failed to create correction target entry")
    entry_id = str(existing[0])
    if amount_fields["mileage_miles"] > 0 and amount_fields["mileage_rate"] <= 0:
        row = await db.fetchone("SELECT mileage_rate FROM pay_entries WHERE id=?", (entry_id,))
        amount_fields["mileage_rate"] = _quantity(row[0] if row else 0)
    correction_id = f"pay-correction-{uuid4().hex}"
    now = utcnow()
    await db.exec(
        """INSERT INTO pay_entry_corrections(
             id, period_id, entry_id, target_user_email, display_name, entry_date,
             hours, mileage_miles, mileage_rate, mileage_amount, rentals_amount,
             meals_amount, hotel_amount, miscellaneous_amount, notes, created_by, created_at_utc
           ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            correction_id,
            period_id,
            entry_id,
            target_email,
            str(data.get("display_name") or "").strip() or None,
            entry_date,
            float(amount_fields["hours"]),
            float(amount_fields["mileage_miles"]),
            float(amount_fields["mileage_rate"]),
            float(amount_fields["mileage_amount"]),
            float(amount_fields["rentals_amount"]),
            float(amount_fields["meals_amount"]),
            float(amount_fields["hotel_amount"]),
            float(amount_fields["miscellaneous_amount"]),
            str(data.get("notes") or "").strip() or None,
            actor.email,
            now,
        ),
    )
    await db.exec(
        """UPDATE pay_entries
           SET hours=hours+?,
               mileage_miles=mileage_miles+?,
               mileage_rate=CASE WHEN ? > 0 THEN ? ELSE mileage_rate END,
               mileage_amount=mileage_amount+?,
               rentals_amount=rentals_amount+?,
               meals_amount=meals_amount+?,
               hotel_amount=hotel_amount+?,
               miscellaneous_amount=miscellaneous_amount+?,
               review_status='pending',
               updated_at_utc=?
           WHERE id=?""",
        (
            float(amount_fields["hours"]),
            float(amount_fields["mileage_miles"]),
            float(amount_fields["mileage_rate"]),
            float(amount_fields["mileage_rate"]),
            float(amount_fields["mileage_amount"]),
            float(amount_fields["rentals_amount"]),
            float(amount_fields["meals_amount"]),
            float(amount_fields["hotel_amount"]),
            float(amount_fields["miscellaneous_amount"]),
            now,
            entry_id,
        ),
    )
    await add_pay_event(
        db,
        period_id=period_id,
        entry_id=entry_id,
        event_type="entry_correction_added",
        actor=actor.email,
        details={"correction_id": correction_id, "user_email": target_email, "entry_date": entry_date},
    )
    row = await db.fetchone("SELECT * FROM pay_entry_corrections WHERE id=?", (correction_id,))
    entries = await list_entries(db, period_id=period_id, actor=PayActor(actor.email, actor.display_name, actor.role, True, True, True))
    return {"id": correction_id, "entry": next(item for item in entries if item["id"] == entry_id)}


def decode_content_base64(value: str) -> bytes:
    raw = str(value or "").strip()
    if raw.startswith("data:") and "," in raw:
        raw = raw.split(",", 1)[1]
    return base64.b64decode(raw, validate=True)


def detect_content_type(filename: str, declared: str | None, content: bytes) -> str:
    declared_type = str(declared or "").split(";", 1)[0].strip().lower()
    if declared_type in _RECEIPT_CONTENT_TYPES:
        return declared_type
    lowered = filename.lower()
    if content.startswith(b"%PDF") or lowered.endswith(".pdf"):
        return "application/pdf"
    if content.startswith(b"\xff\xd8\xff") or lowered.endswith((".jpg", ".jpeg")):
        return "image/jpeg"
    if content.startswith(b"\x89PNG\r\n\x1a\n") or lowered.endswith(".png"):
        return "image/png"
    return declared_type or "application/octet-stream"


def validate_receipt_file(*, filename: str, content_type: str, content: bytes, max_file_bytes: int) -> None:
    if not content:
        raise ValueError("file is empty")
    if len(content) > max_file_bytes:
        raise ValueError("file exceeds maximum size")
    if content_type not in _RECEIPT_CONTENT_TYPES:
        raise ValueError("only PDF, JPEG, and PNG files are allowed")
    if content_type == "application/pdf" and not content.startswith(b"%PDF"):
        raise ValueError("PDF content is invalid")
    if content_type == "image/jpeg" and not content.startswith(b"\xff\xd8\xff"):
        raise ValueError("JPEG content is invalid")
    if content_type == "image/png" and not content.startswith(b"\x89PNG\r\n\x1a\n"):
        raise ValueError("PNG content is invalid")


def scan_with_clamav(*, content: bytes, host: str, port: int, timeout_seconds: int) -> str:
    if not host:
        raise RuntimeError("ClamAV host is not configured")
    with socket.create_connection((host, int(port)), timeout=float(timeout_seconds)) as sock:
        sock.settimeout(float(timeout_seconds))
        sock.sendall(b"zINSTREAM\0")
        for start in range(0, len(content), 1024 * 1024):
            chunk = content[start : start + 1024 * 1024]
            sock.sendall(len(chunk).to_bytes(4, "big"))
            sock.sendall(chunk)
        sock.sendall((0).to_bytes(4, "big"))
        response = sock.recv(4096).decode("utf-8", errors="replace").strip()
    if "FOUND" in response:
        raise ValueError(f"virus scan failed: {response}")
    if "OK" not in response:
        raise RuntimeError(f"virus scan did not complete cleanly: {response}")
    return response


async def store_attachment(
    db: Db,
    *,
    cfg: Any,
    period_id: str,
    entry_id: str,
    actor: PayActor,
    attachment_type: str,
    filename: str,
    content_type: str | None,
    content: bytes,
    scan: bool = True,
    mileage_miles: object | None = None,
    mileage_rate: object | None = None,
    mileage_amount: object | None = None,
) -> dict[str, object]:
    entry = await db.fetchone(
        "SELECT user_email, locked_at_utc FROM pay_entries WHERE id=? AND period_id=?",
        (entry_id, period_id),
    )
    if not entry:
        raise ValueError("entry not found")
    if entry[1]:
        raise ValueError("entry is locked")
    if not actor.can_edit_all and normalize_email(entry[0]) != actor.email:
        raise PermissionError("cannot attach files to another user's entry")

    safe_name = safe_filename(filename, fallback="receipt")
    detected_type = detect_content_type(safe_name, content_type, content)
    validate_receipt_file(
        filename=safe_name,
        content_type=detected_type,
        content=content,
        max_file_bytes=int(cfg.pay_portal.receipt_max_file_bytes),
    )

    scan_result = "scan-skipped"
    if scan:
        scan_result = scan_with_clamav(
            content=content,
            host=cfg.pay_portal.clamav_host,
            port=cfg.pay_portal.clamav_port,
            timeout_seconds=cfg.pay_portal.clamav_timeout_seconds,
        )

    attachment_id = f"pay-att-{uuid4().hex}"
    ext = _RECEIPT_CONTENT_TYPES[detected_type]
    stored_filename = f"{attachment_id}{ext}"
    period_dir = Path(cfg.data_root) / "pay" / period_id / "attachments"
    period_dir.mkdir(parents=True, exist_ok=True)
    local_path = period_dir / stored_filename
    local_path.write_bytes(content)
    sha = hashlib.sha256(content).hexdigest()
    await db.exec(
        """INSERT INTO pay_attachments(
             id, period_id, entry_id, uploaded_by, attachment_type, original_filename,
             stored_filename, local_path, content_type, size_bytes, sha256, scan_status,
             scan_result, mileage_miles, mileage_rate, mileage_amount, created_at_utc
           ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            attachment_id,
            period_id,
            entry_id,
            actor.email,
            attachment_type,
            safe_name,
            stored_filename,
            str(local_path),
            detected_type,
            len(content),
            sha,
            "clean",
            scan_result,
            float(_quantity(mileage_miles)) if mileage_miles is not None else None,
            float(_quantity(mileage_rate)) if mileage_rate is not None else None,
            float(_money(mileage_amount)) if mileage_amount is not None else None,
            utcnow(),
        ),
    )
    await add_pay_event(
        db,
        period_id=period_id,
        entry_id=entry_id,
        event_type="attachment_uploaded",
        actor=actor.email,
        details={"filename": safe_name, "attachment_type": attachment_type, "sha256": sha},
    )
    return {
        "id": attachment_id,
        "attachment_type": attachment_type,
        "filename": safe_name,
        "content_type": detected_type,
        "size_bytes": len(content),
        "sha256": sha,
        "scan_status": "clean",
        "mileage_miles": float(_quantity(mileage_miles)) if mileage_miles is not None else None,
        "mileage_rate": float(_quantity(mileage_rate)) if mileage_rate is not None else None,
        "mileage_amount": float(_money(mileage_amount)) if mileage_amount is not None else None,
    }


async def store_compensation_stub(
    db: Db,
    *,
    cfg: Any,
    actor: PayActor,
    user_email: str | None,
    base_wage_input_type: object,
    base_wage_amount: object,
    weekly_basis_hours: object,
    commission_month_1_amount: object,
    commission_month_2_amount: object,
    commission_month_3_amount: object,
    filename: str,
    content_type: str | None,
    content: bytes,
    payroll_month: object | None = None,
    notes: str | None = None,
    scan: bool = True,
) -> dict[str, object]:
    target_email = normalize_email(user_email) or actor.email
    if not target_email:
        raise ValueError("user_email is required")
    if not actor.can_edit_all and target_email != actor.email:
        raise PermissionError("cannot upload lost wage proof for another member")

    safe_name = safe_filename(filename, fallback="pay-stub")
    detected_type = detect_content_type(safe_name, content_type, content)
    validate_receipt_file(
        filename=safe_name,
        content_type=detected_type,
        content=content,
        max_file_bytes=int(cfg.pay_portal.receipt_max_file_bytes),
    )
    scan_result = "scan-skipped"
    if scan:
        scan_result = scan_with_clamav(
            content=content,
            host=cfg.pay_portal.clamav_host,
            port=cfg.pay_portal.clamav_port,
            timeout_seconds=cfg.pay_portal.clamav_timeout_seconds,
        )

    basis = _quantity(weekly_basis_hours)
    if basis <= 0:
        basis = Decimal("40")
    normalized_payroll_month = normalize_payroll_month(payroll_month, default_for=date.today())
    calc = calculate_commission_compensation(
        base_wage_input_type=base_wage_input_type,
        base_wage_amount=base_wage_amount,
        weekly_basis_hours=basis,
        commission_month_1_amount=commission_month_1_amount,
        commission_month_2_amount=commission_month_2_amount,
        commission_month_3_amount=commission_month_3_amount,
    )
    stub_id = f"pay-comp-{uuid4().hex}"
    ext = _RECEIPT_CONTENT_TYPES[detected_type]
    stored_filename = f"{stub_id}{ext}"
    stub_dir = Path(cfg.data_root) / "pay" / "compensation-stubs" / safe_filename(target_email, fallback="member")
    stub_dir.mkdir(parents=True, exist_ok=True)
    local_path = stub_dir / stored_filename
    local_path.write_bytes(content)
    sha = hashlib.sha256(content).hexdigest()
    now = utcnow()
    await db.exec(
        """INSERT INTO pay_compensation_stubs(
             id, user_email, uploaded_by, base_wage_input_type, base_wage_amount,
             weekly_basis_hours, commission_month_1_amount, commission_month_2_amount,
             commission_month_3_amount, commission_average_monthly, commission_hourly_rate,
             calculated_hourly_rate, payroll_month, original_filename, stored_filename, local_path,
             content_type, size_bytes, sha256, scan_status, scan_result, notes, created_at_utc
           ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            stub_id,
            target_email,
            actor.email,
            calc.base_wage_input_type,
            float(calc.base_wage_amount),
            float(basis),
            float(calc.commission_month_1_amount),
            float(calc.commission_month_2_amount),
            float(calc.commission_month_3_amount),
            float(calc.commission_average_monthly),
            float(calc.commission_hourly_rate),
            float(calc.calculated_hourly_rate),
            normalized_payroll_month,
            safe_name,
            stored_filename,
            str(local_path),
            detected_type,
            len(content),
            sha,
            "clean",
            scan_result,
            str(notes or "").strip(),
            now,
        ),
    )
    await add_pay_event(
        db,
        period_id=None,
        event_type="compensation_stub_uploaded",
        actor=actor.email,
        details={
            "user_email": target_email,
            "payroll_month": normalized_payroll_month,
            "filename": safe_name,
            "sha256": sha,
            "calculated_hourly_rate": float(calc.calculated_hourly_rate),
        },
    )
    saved = await latest_compensation_stub(db, user_email=target_email)
    return saved or {
        "id": stub_id,
        "user_email": target_email,
        "calculated_hourly_rate": float(calc.calculated_hourly_rate),
    }


async def list_attachments(db: Db, *, period_id: str, actor: PayActor) -> list[dict[str, object]]:
    select_sql = """SELECT a.id, a.entry_id, a.attachment_type, a.original_filename, a.content_type,
                         a.size_bytes, a.sha256, a.scan_status, a.sharepoint_url, a.created_at_utc,
                         e.user_email, e.entry_date, a.mileage_miles, a.mileage_rate, a.mileage_amount,
                         e.locked_at_utc, p.status,
                         (SELECT COUNT(*) FROM pay_attachments ma
                          WHERE ma.entry_id=a.entry_id AND ma.attachment_type='mileage_pdf'
                            AND ma.removed_at_utc IS NULL) AS active_mileage_count,
                         e.display_name, e.local_number, e.notes,
                         e.mileage_miles, e.mileage_rate, e.mileage_amount
                  FROM pay_attachments a
                  JOIN pay_entries e ON e.id = a.entry_id
                  JOIN pay_periods p ON p.id = a.period_id
                  WHERE a.period_id=? AND a.removed_at_utc IS NULL
             AND COALESCE(e.review_status, 'pending') IN ('pending', 'approved')"""
    if actor.can_view_all:
        rows = await db.fetchall(f"{select_sql} ORDER BY e.entry_date, a.created_at_utc", (period_id,))
    else:
        rows = await db.fetchall(
            f"{select_sql} AND e.user_email=? ORDER BY e.entry_date, a.created_at_utc",
            (period_id, actor.email),
        )
    attachments: list[dict[str, object]] = []
    for row in rows:
        is_mileage = str(row[2] or "") == "mileage_pdf"
        has_summary = row[14] is not None
        entry_locked = bool(row[15]) or str(row[16] or "") != "open"
        active_mileage_count = int(row[17] or 0)
        fallback_miles = row[21] if row[21] is not None else row[12]
        fallback_rate = row[22] if row[22] is not None else row[13]
        fallback_amount = row[23] if row[23] is not None else row[14]
        legacy_totals_available = is_mileage and fallback_amount is not None and _money(fallback_amount) > 0
        owner_email = normalize_email(row[10])
        actor_can_remove = actor.can_edit_all or actor.can_lock or owner_email == actor.email
        can_remove = is_mileage and not entry_locked and actor_can_remove and (
            has_summary or active_mileage_count == 1 or legacy_totals_available
        )
        remove_reason = ""
        if is_mileage and not actor_can_remove:
            remove_reason = "Only the submitter or treasurer can remove this report."
        elif is_mileage and entry_locked:
            remove_reason = "Entry is locked."
        elif is_mileage and not can_remove:
            remove_reason = "This older report cannot be safely removed automatically."
        attachments.append(
            {
                "id": row[0],
                "entry_id": row[1],
                "attachment_type": row[2],
                "filename": row[3],
                "content_type": row[4],
                "size_bytes": row[5],
                "sha256": row[6],
                "scan_status": row[7],
                "sharepoint_url": row[8],
                "created_at_utc": row[9],
                "user_email": row[10],
                "entry_date": row[11],
                "mileage_miles": row[12] if row[12] is not None else fallback_miles,
                "mileage_rate": row[13] if row[13] is not None else fallback_rate,
                "mileage_amount": row[14] if row[14] is not None else fallback_amount,
                "display_name": row[18],
                "local_number": row[19],
                "description": row[20],
                "legacy_summary": is_mileage and not has_summary and legacy_totals_available,
                "active_mileage_count": active_mileage_count,
                "can_remove": can_remove,
                "remove_reason": remove_reason,
            }
        )
    return attachments


async def attachment_for_actor(
    db: Db,
    *,
    attachment_id: str,
    actor: PayActor,
    include_removed: bool = False,
) -> dict[str, object]:
    removed_clause = "" if include_removed else "AND a.removed_at_utc IS NULL"
    row = await db.fetchone(
        f"""SELECT a.id, a.period_id, a.entry_id, a.attachment_type, a.original_filename,
                   a.local_path, a.content_type, a.size_bytes, a.sha256, a.scan_status,
                   a.mileage_miles, a.mileage_rate, a.mileage_amount, a.removed_at_utc,
                   e.user_email, e.entry_date, e.locked_at_utc, e.mileage_miles, e.mileage_rate,
                   e.mileage_amount, p.status,
                   (SELECT COUNT(*) FROM pay_attachments ma
                    WHERE ma.entry_id=a.entry_id AND ma.attachment_type='mileage_pdf'
                      AND ma.removed_at_utc IS NULL) AS active_mileage_count
            FROM pay_attachments a
            JOIN pay_entries e ON e.id = a.entry_id
            JOIN pay_periods p ON p.id = a.period_id
            WHERE a.id=? {removed_clause}""",
        (attachment_id,),
    )
    if not row:
        raise ValueError("attachment not found")
    if not (actor.can_view_all or actor.can_edit_all or actor.can_lock) and normalize_email(row[14]) != actor.email:
        raise PermissionError("cannot access another user's attachment")
    return {
        "id": row[0],
        "period_id": row[1],
        "entry_id": row[2],
        "attachment_type": row[3],
        "filename": row[4],
        "local_path": row[5],
        "content_type": row[6],
        "size_bytes": row[7],
        "sha256": row[8],
        "scan_status": row[9],
        "mileage_miles": row[10],
        "mileage_rate": row[11],
        "mileage_amount": row[12],
        "removed_at_utc": row[13],
        "user_email": row[14],
        "entry_date": row[15],
        "locked_at_utc": row[16],
        "entry_mileage_miles": row[17],
        "entry_mileage_rate": row[18],
        "entry_mileage_amount": row[19],
        "period_status": row[20],
        "active_mileage_count": int(row[21] or 0),
    }


async def remove_mileage_attachment(
    db: Db,
    *,
    attachment_id: str,
    actor: PayActor,
    reason: str | None = None,
) -> dict[str, object]:
    attachment = await attachment_for_actor(db, attachment_id=attachment_id, actor=actor)
    owner_email = normalize_email(attachment.get("user_email"))
    if not (actor.can_edit_all or actor.can_lock or owner_email == actor.email):
        raise PermissionError("cannot remove another user's mileage report")
    if str(attachment["attachment_type"] or "") != "mileage_pdf":
        raise ValueError("only mileage reports can be removed here")
    if attachment.get("locked_at_utc") or str(attachment.get("period_status") or "") != "open":
        raise ValueError("entry is locked")

    miles_value = attachment.get("mileage_miles")
    rate_value = attachment.get("mileage_rate")
    amount_value = attachment.get("mileage_amount")
    subtract_entry_totals = True
    if amount_value is None:
        active_mileage_count = int(attachment.get("active_mileage_count") or 0)
        if active_mileage_count != 1:
            entry_amount = _money(attachment.get("entry_mileage_amount"))
            if entry_amount <= 0:
                raise ValueError("this older mileage report cannot be safely removed automatically")
            miles_value = Decimal("0.00")
            rate_value = attachment.get("entry_mileage_rate")
            amount_value = Decimal("0.00")
            subtract_entry_totals = False
        else:
            miles_value = attachment.get("entry_mileage_miles")
            rate_value = attachment.get("entry_mileage_rate")
            amount_value = attachment.get("entry_mileage_amount")
            await db.exec(
                """UPDATE pay_attachments
                   SET mileage_miles=?, mileage_rate=?, mileage_amount=?
                   WHERE id=?""",
                (float(_quantity(miles_value)), float(_quantity(rate_value)), float(_money(amount_value)), attachment_id),
            )

    remove_miles = _quantity(miles_value)
    remove_amount = _money(amount_value)
    current_miles = _quantity(attachment.get("entry_mileage_miles"))
    current_amount = _money(attachment.get("entry_mileage_amount"))
    new_miles = max(current_miles - remove_miles, Decimal("0.00")) if subtract_entry_totals else current_miles
    new_amount = max(current_amount - remove_amount, Decimal("0.00")) if subtract_entry_totals else current_amount
    new_rate = Decimal("0.00") if new_miles == 0 else _quantity(attachment.get("entry_mileage_rate"))
    now = utcnow()
    await db.exec(
        """UPDATE pay_attachments
           SET removed_at_utc=?, removed_by=?, removed_reason=?
           WHERE id=? AND removed_at_utc IS NULL""",
        (now, actor.email, str(reason or "submitter removed mileage report").strip(), attachment_id),
    )
    await db.exec(
        """UPDATE pay_entries
           SET mileage_miles=?, mileage_rate=?, mileage_amount=?, updated_at_utc=?
           WHERE id=? AND period_id=?""",
        (float(new_miles), float(new_rate), float(new_amount), now, attachment["entry_id"], attachment["period_id"]),
    )
    await add_pay_event(
        db,
        period_id=str(attachment["period_id"]),
        entry_id=str(attachment["entry_id"]),
        event_type="mileage_report_removed",
        actor=actor.email,
        details={
            "attachment_id": attachment_id,
            "filename": attachment.get("filename"),
            "mileage_miles": float(remove_miles),
            "mileage_amount": float(remove_amount),
            "subtracted_entry_totals": subtract_entry_totals,
        },
    )
    return {
        **attachment,
        "removed_at_utc": now,
        "mileage_miles": float(remove_miles),
        "mileage_rate": float(_quantity(rate_value)),
        "mileage_amount": float(remove_amount),
        "subtracted_entry_totals": subtract_entry_totals,
    }


def _cell_set_text(cell: Any, text: object) -> None:
    cell.text = str(text or "")


def _cell_set_paragraphs(cell: Any, lines: list[str]) -> None:
    cell.text = ""
    if not lines:
        return
    if cell.paragraphs:
        cell.paragraphs[0].text = lines[0]
    else:
        cell.add_paragraph(lines[0])
    for line in lines[1:]:
        cell.add_paragraph(line)


def _cell_text(cell: Any) -> str:
    return re.sub(r"\s+", " ", str(getattr(cell, "text", "") or "")).strip()


def _set_paragraph_if_prefix(doc: Document, prefix: str, value: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            paragraph.text = value
            return


def _set_voucher_signature_line(
    doc: Document,
    *,
    submitter_name: str,
    approved_by_text: str,
) -> None:
    submitter_text = str(submitter_name or "Submitted electronically").strip() or "Submitted electronically"
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip().startswith("Signature"):
            continue
        signature_replacements = [submitter_text, approved_by_text]
        replacement_index = 0
        if "\t" in paragraph.text and len(paragraph.runs) > 1:
            for run in paragraph.runs:
                if replacement_index >= len(signature_replacements):
                    break
                if run.text.strip() != "Signature":
                    continue
                replacement = signature_replacements[replacement_index]
                leading = run.text[: len(run.text) - len(run.text.lstrip())]
                trailing = run.text[len(run.text.rstrip()) :]
                run.text = f"{leading}{replacement}{trailing}"
                if "{{" in replacement:
                    run.font.size = Pt(4)
                replacement_index += 1
            if replacement_index == len(signature_replacements):
                return
        paragraph.clear()
        paragraph.add_run(submitter_text)
        paragraph.add_run("\t")
        approved_run = paragraph.add_run(approved_by_text)
        if "{{" in approved_by_text:
            approved_run.font.size = Pt(4)
        paragraph.add_run("\tPaid by")
        return


def _entry_amounts(row: dict[str, object]) -> dict[str, Decimal]:
    hourly_rate = _money(row.get("lost_wage_hourly_rate") or row.get("hourly_rate"))
    hours_amount = _money(row.get("hours")) * hourly_rate
    return {
        "hours": hours_amount.quantize(_CURRENCY, rounding=ROUND_HALF_UP),
        "mileage": _money(row.get("mileage_amount")),
        "rentals": _money(row.get("rentals_amount")),
        "meals": _money(row.get("meals_amount")),
        "hotel": _money(row.get("hotel_amount")),
        "misc": _money(row.get("miscellaneous_amount")),
        "president_diff": _money(row.get("president_diff_amount")),
    }


def _voucher_quantity(value: object) -> Decimal:
    return _quantity(value).quantize(_CURRENCY, rounding=ROUND_HALF_UP)


def _voucher_quantity_text(value: object) -> str:
    amount = _voucher_quantity(value)
    return "" if amount == 0 else f"{amount:.2f}"


def _entry_weekly_quantities(row: dict[str, object]) -> dict[str, Decimal]:
    return {
        "hours": _voucher_quantity(row.get("hours")),
        "mileage": _voucher_quantity(row.get("mileage_miles")),
        "rentals": _money(row.get("rentals_amount")),
        "meals": _money(row.get("meals_amount")),
        "hotel": _money(row.get("hotel_amount")),
        "misc": _money(row.get("miscellaneous_amount")),
        "president_diff": _voucher_quantity(row.get("president_diff_hours")),
    }


def pay_entry_daily_narrative_lines(entries: list[dict[str, object]]) -> list[str]:
    notes_by_day: dict[date, list[str]] = {}
    seen_by_day: dict[date, set[str]] = {}
    for entry in entries:
        try:
            entry_day = date.fromisoformat(str(entry.get("entry_date")))
        except Exception:
            continue
        note = re.sub(r"\s+", " ", str(entry.get("notes") or "").strip())
        if not note:
            continue
        seen = seen_by_day.setdefault(entry_day, set())
        dedupe_key = note.lower()
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        notes_by_day.setdefault(entry_day, []).append(note)

    lines: list[str] = []
    for entry_day in sorted(notes_by_day):
        day_label = entry_day.strftime("%a %m/%d/%Y")
        lines.append(f"{day_label}: {'; '.join(notes_by_day[entry_day])}")
    return lines


def _split_narrative_line(line: str) -> tuple[str, str]:
    if ": " not in line:
        return "", line
    label, narrative = line.split(": ", 1)
    return label.strip(), narrative.strip()


def _remove_front_receipt_explanation_rows(doc: Document) -> None:
    for table in doc.tables[:2]:
        for row in list(table.rows):
            row_text = " ".join(_cell_text(cell) for cell in row.cells).lower()
            if "attach necessary receipts" not in row_text:
                continue
            try:
                table._tbl.remove(row._tr)  # noqa: SLF001
            except Exception:
                for cell in row.cells:
                    _cell_set_text(cell, "")


def _find_totals_row(doc: Document) -> Any | None:
    for table in doc.tables:
        for row in table.rows:
            cells = list(row.cells)
            if cells and _cell_text(cells[0]).upper() == "TOTALS":
                return row
    return None


def _write_daily_narrative(
    doc: Document,
    *,
    period_start: str,
    period_end: str,
    narrative_lines: list[str],
) -> None:
    if not narrative_lines:
        return
    doc.add_page_break()
    doc.add_paragraph("Daily Narrative by Date")
    doc.add_paragraph(f"Pay period {period_start} to {period_end}")
    table = doc.add_table(rows=1, cols=2)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    header = table.rows[0].cells
    _cell_set_text(header[0], "Date")
    _cell_set_text(header[1], "Narrative")
    for line in narrative_lines:
        label, narrative = _split_narrative_line(line)
        cells = table.add_row().cells
        _cell_set_text(cells[0], label)
        _cell_set_text(cells[1], narrative)


def fill_pay_voucher_docx(
    *,
    template_path: str,
    output_path: str,
    period_start: str,
    period_end: str,
    entries: list[dict[str, object]],
    include_signature_placeholders: bool = False,
    filer_signer_index: int = 1,
    paid_by_signer_index: int = 2,
) -> None:
    doc = Document(template_path)
    first = entries[0] if entries else {}
    local_number = str(first.get("local_number") or "3106").strip()
    display_name = str(first.get("display_name") or first.get("user_email") or "").strip()
    address = str(first.get("address") or "").strip()
    hourly_rate = _currency_text(first.get("lost_wage_hourly_rate") or first.get("hourly_rate"))

    _set_paragraph_if_prefix(doc, "Local #", f"Local # {local_number}    Date: {period_end}")
    _set_paragraph_if_prefix(doc, "Name", f"Name {display_name}")
    _set_paragraph_if_prefix(doc, "Address", f"Address {address}")
    _set_paragraph_if_prefix(doc, "Hourly Rate", f"Hourly Rate {hourly_rate}")

    period_start_date = date.fromisoformat(period_start)
    table_pairs = [(0, 0), (1, 7)]
    row_names = {
        "hours": 1,
        "mileage": 2,
        "rentals": 3,
        "meals": 4,
        "hotel": 5,
        "misc": 6,
        "president_diff": 7,
    }
    totals_by_category = {key: Decimal("0.00") for key in row_names}
    narrative_lines = pay_entry_daily_narrative_lines(entries)

    for table_index, day_offset in table_pairs:
        if table_index >= len(doc.tables):
            continue
        table = doc.tables[table_index]
        for col in range(1, min(8, len(table.rows[0].cells))):
            current = period_start_date + timedelta(days=day_offset + col - 1)
            _cell_set_text(table.rows[0].cells[col], current.strftime("%a\n%m/%d"))
        day_totals = [Decimal("0.00") for _ in range(7)]
        for entry in entries:
            try:
                entry_day = date.fromisoformat(str(entry.get("entry_date")))
            except Exception:
                continue
            index = (entry_day - period_start_date).days - day_offset
            if index < 0 or index > 6:
                continue
            col = index + 1
            amounts = _entry_amounts(entry)
            weekly_values = _entry_weekly_quantities(entry)
            for key, row_no in row_names.items():
                totals_by_category[key] += amounts[key]
                value = weekly_values[key]
                current_value = _money(table.rows[row_no].cells[col].text)
                _cell_set_text(table.rows[row_no].cells[col], _voucher_quantity_text(current_value + value))
            day_totals[index] += weekly_values["hours"] + weekly_values["president_diff"]
        for key, row_no in row_names.items():
            row_total = Decimal("0.00")
            for col in range(1, 8):
                row_total += _money(table.rows[row_no].cells[col].text)
            _cell_set_text(table.rows[row_no].cells[8], _voucher_quantity_text(row_total))
        for index, day_total in enumerate(day_totals):
            _cell_set_text(table.rows[8].cells[index + 1], _voucher_quantity_text(day_total))
        _cell_set_text(table.rows[8].cells[8], _voucher_quantity_text(sum(day_totals, Decimal("0.00"))))

    _remove_front_receipt_explanation_rows(doc)

    totals_row = _find_totals_row(doc)
    if totals_row is not None:
        values = [
            totals_by_category["hours"],
            totals_by_category["mileage"],
            totals_by_category["rentals"],
            totals_by_category["meals"],
            totals_by_category["hotel"],
            totals_by_category["misc"],
            sum(totals_by_category.values(), Decimal("0.00")),
        ]
        for idx, value in enumerate(values, start=1):
            if idx < len(totals_row.cells):
                _cell_set_text(totals_row.cells[idx], _currency_total_text(value))

    if include_signature_placeholders:
        paid_by_index = max(1, int(paid_by_signer_index or 1))
        approved_by_text = (
            f"{{{{Sig_es_:signer{paid_by_index}:signer{paid_by_index}_signature}}}}    "
            f"{{{{Dte_es_:signer{paid_by_index}:signer{paid_by_index}_date}}}}"
        )
    else:
        approved_by_text = "Signature"
    _set_voucher_signature_line(doc, submitter_name=display_name, approved_by_text=approved_by_text)

    _write_daily_narrative(
        doc,
        period_start=period_start,
        period_end=period_end,
        narrative_lines=narrative_lines,
    )

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def merge_pdfs(input_paths: list[str], output_path: str) -> None:
    paths = [str(path) for path in input_paths if path and Path(path).exists()]
    if not paths:
        raise RuntimeError("no PDF files available to merge")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    try:
        from pypdf import PdfReader, PdfWriter  # type: ignore

        writer = PdfWriter()
        for path in paths:
            reader = PdfReader(path)
            for page in reader.pages:
                writer.add_page(page)
        with open(output_path, "wb") as out:
            writer.write(out)
        return
    except ImportError:
        pass
    result = subprocess.run(
        ["pdfunite", *paths, output_path],
        check=False,
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(f"pdfunite failed: {result.stderr[:400]}")


def image_to_pdf(image_path: str, output_path: str) -> None:
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    with Image.open(image_path) as img:
        converted = img.convert("RGB")
        converted.save(output_path, "PDF", resolution=100.0)


def _extract_first_pdf(zip_bytes: bytes) -> bytes | None:
    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            names = sorted(zf.namelist())
            for name in names:
                if name.lower().endswith(".pdf"):
                    data = zf.read(name)
                    if data:
                        return data
    except Exception:
        return None
    return None


async def _attachments_for_packet(db: Db, period_id: str) -> list[dict[str, object]]:
    rows = await db.fetchall(
        """SELECT a.id, a.entry_id, a.attachment_type, a.original_filename, a.local_path,
                  a.content_type, a.sha256, e.user_email, e.entry_date
           FROM pay_attachments a
           JOIN pay_entries e ON e.id = a.entry_id
           WHERE a.period_id=?
             AND a.removed_at_utc IS NULL
             AND COALESCE(e.review_status, 'pending') IN ('pending', 'approved')
           ORDER BY e.user_email, e.entry_date, a.created_at_utc""",
        (period_id,),
    )
    return [
        {
            "id": row[0],
            "entry_id": row[1],
            "attachment_type": row[2],
            "original_filename": row[3],
            "local_path": row[4],
            "content_type": row[5],
            "sha256": row[6],
            "user_email": row[7],
            "entry_date": row[8],
        }
        for row in rows
    ]


async def _compensation_stubs_for_packet(db: Db, period_id: str) -> list[dict[str, object]]:
    rows = await db.fetchall(
        """SELECT DISTINCT s.id, s.user_email, s.original_filename, s.local_path,
                  s.content_type, s.sha256
           FROM pay_entries e
           JOIN pay_compensation_stubs s ON s.id = e.compensation_stub_id
           WHERE e.period_id=?
             AND COALESCE(e.review_status, 'pending') IN ('pending', 'approved')
           ORDER BY s.user_email, s.created_at_utc""",
        (period_id,),
    )
    return [
        {
            "id": row[0],
            "user_email": row[1],
            "original_filename": row[2],
            "local_path": row[3],
            "content_type": row[4],
            "sha256": row[5],
        }
        for row in rows
    ]


def _pay_sharepoint_upload_target(*, cfg: Any, folder_path: str) -> tuple[str, str]:
    path_parts = [part.strip() for part in str(folder_path or "").replace("\\", "/").split("/") if part.strip()]
    inferred_library = ""
    inferred_folder = ""
    for index, part in enumerate(path_parts):
        lowered = part.lower()
        if lowered == "documents" or lowered.endswith(" - documents"):
            inferred_library = part
            inferred_folder = "/".join(path_parts[index + 1 :])
            break
    folder = inferred_folder or "/".join(path_parts)
    configured_library = str(getattr(cfg.pay_portal, "sharepoint_library", "") or "").strip()
    default_library = str(getattr(cfg.graph, "document_library", "") or "").strip()
    return inferred_library or configured_library or default_library, folder


async def _upload_if_configured(
    *,
    cfg: Any,
    graph: Any,
    folder_path: str,
    filename: str,
    local_path: str,
) -> tuple[str | None, str | None]:
    library, upload_folder_path = _pay_sharepoint_upload_target(cfg=cfg, folder_path=folder_path)
    site_hostname, site_path = _pay_graph_site(cfg=cfg)
    if not (graph and site_hostname and site_path and library and upload_folder_path):
        return None, None
    libraries = [library]
    default_library = str(getattr(cfg.graph, "document_library", "") or "").strip()
    if default_library and default_library not in libraries:
        libraries.append(default_library)
    last_error: Exception | None = None
    for candidate_library in libraries:
        try:
            uploaded = graph.upload_local_file_to_folder_path(
                site_hostname=site_hostname,
                site_path=site_path,
                library=candidate_library,
                folder_path=upload_folder_path,
                filename=filename,
                local_path=local_path,
            )
            return uploaded.web_url, uploaded.path
        except RuntimeError as exc:
            last_error = exc
            if "Could not find document library drive named" not in str(exc):
                raise
    if last_error:
        raise last_error
    return None, None


def pay_packet_signer_order(*, grouped_entry_emails: list[str], president_signer_email: str) -> tuple[list[str], int]:
    _ = grouped_entry_emails
    president = normalize_email(president_signer_email)
    if not president:
        raise ValueError("president signer email is required")
    return [president], 1


def _send_president_signature_email(
    *,
    cfg: Any,
    mailer: Any | None,
    president_email: str,
    period_start: str,
    period_end: str,
    signing_link: str | None,
    sharepoint_unsigned_url: str | None,
    folder_path: str,
    packet_id: str,
) -> dict[str, object] | None:
    if mailer is None or not bool(getattr(cfg.email, "enabled", False)):
        return None
    link = str(signing_link or "").strip()
    if not link:
        raise RuntimeError("DocuSeal did not return a president signing link")
    subject = f"Pay packet ready for signature: {period_start} to {period_end}"
    text_body = (
        f"The Local 3106 pay packet for {period_start} to {period_end} is ready for your signature.\n\n"
        f"Sign here: {link}\n"
        f"SharePoint packet folder: {folder_path}\n"
    )
    if sharepoint_unsigned_url:
        text_body += f"Unsigned packet: {sharepoint_unsigned_url}\n"
    html_body = (
        f"<p>The Local 3106 pay packet for <strong>{html.escape(period_start)} to {html.escape(period_end)}</strong> "
        "is ready for your signature.</p>"
        f'<p><a href="{html.escape(link, quote=True)}">Open DocuSeal signing link</a></p>'
        f"<p>SharePoint packet folder: {html.escape(folder_path)}</p>"
    )
    if sharepoint_unsigned_url:
        html_body += f'<p>Unsigned packet: <a href="{html.escape(sharepoint_unsigned_url, quote=True)}">Open packet</a></p>'
    sent = mailer.send_mail(
        to_recipients=[president_email],
        subject=subject,
        text_body=text_body,
        html_body=html_body,
        custom_headers={"X-CWA3106-Pay-Packet-ID": packet_id},
    )
    return {
        "sent": True,
        "graph_message_id": getattr(sent, "graph_message_id", ""),
        "internet_message_id": getattr(sent, "internet_message_id", None),
    }


async def lock_period_and_send_packet(
    *,
    db: Db,
    cfg: Any,
    graph: Any,
    docuseal: Any,
    period_id: str,
    actor: PayActor,
    president_signer_email: str | None,
    docx_to_pdf_func: Any,
    mailer: Any | None = None,
) -> dict[str, object]:
    _pay_lock_log("started", period_id=period_id, actor=actor.email)
    if not actor.can_lock:
        _pay_lock_log("permission_denied", period_id=period_id, actor=actor.email)
        raise PermissionError("treasurer access required")
    _pay_lock_log("load_period", period_id=period_id)
    period = await get_pay_period(db, period_id)
    if not period:
        _pay_lock_log("period_missing", period_id=period_id)
        raise ValueError("pay period not found")
    _pay_lock_log(
        "period_loaded",
        period_id=period_id,
        status=str(period["status"]),
        revision=int(period["revision"]),
    )
    if str(period["status"]) not in {"open", "locked"}:
        _pay_lock_log("period_not_open", period_id=period_id, status=str(period["status"]))
        raise ValueError("pay period is already sent or completed")
    _pay_lock_log("list_entries", period_id=period_id)
    entries = await list_entries(
        db,
        period_id=period_id,
        actor=PayActor(actor.email, actor.display_name, actor.role, True, True, True),
    )
    all_entries = entries
    includable_statuses = {"pending", "approved"}
    excluded_entries = [
        entry for entry in all_entries if str(entry.get("review_status") or "pending") not in includable_statuses
    ]
    entries = [entry for entry in all_entries if str(entry.get("review_status") or "pending") in includable_statuses]
    _pay_lock_log(
        "entries_loaded",
        period_id=period_id,
        total_entries=len(all_entries),
        includable_entries=len(entries),
        excluded_entries=len(excluded_entries),
    )
    if not entries:
        _pay_lock_log("no_includable_entries", period_id=period_id)
        raise ValueError("cannot lock a pay period with no pending or approved entries")
    _pay_lock_log("resolve_president_signer", period_id=period_id, explicit=bool(president_signer_email))
    president_signer = await president_email(db, explicit=president_signer_email, pay_cfg=cfg.pay_portal)
    if not president_signer:
        _pay_lock_log("president_signer_missing", period_id=period_id)
        raise ValueError("president signer email is required")
    _pay_lock_log("president_signer_resolved", period_id=period_id, president_signer=president_signer)

    packet_id = f"pay-packet-{uuid4().hex}"
    _pay_lock_log("packet_id_created", period_id=period_id, packet_id=packet_id)
    packet_dir = Path(cfg.data_root) / "pay" / period_id / "packet" / packet_id
    voucher_dir = packet_dir / "vouchers"
    support_dir = packet_dir / "support"
    voucher_dir.mkdir(parents=True, exist_ok=True)
    support_dir.mkdir(parents=True, exist_ok=True)
    period_start = str(period["period_start"])
    period_end = str(period["period_end"])

    grouped: dict[str, list[dict[str, object]]] = {}
    for entry in entries:
        grouped.setdefault(str(entry["user_email"]), []).append(entry)
    ordered_groups = sorted(grouped.items())
    signer_order, president_signer_index = pay_packet_signer_order(
        grouped_entry_emails=[email for email, _ in ordered_groups],
        president_signer_email=president_signer,
    )
    _pay_lock_log(
        "entries_grouped",
        period_id=period_id,
        packet_id=packet_id,
        person_count=len(ordered_groups),
        signer_count=len(signer_order),
        president_signer_index=president_signer_index,
    )

    voucher_paths: list[str] = []
    voucher_pdf_paths: list[str] = []
    anchor_pdf_paths: list[str] = []
    packet_pdf_paths: list[str] = []
    alignment_pdf_paths: list[str] = []
    support_pdf_paths_by_user: dict[str, list[str]] = {}
    _pay_lock_log("collect_support_documents", period_id=period_id, packet_id=packet_id)
    packet_attachments = await _attachments_for_packet(db, period_id)
    _pay_lock_log(
        "support_documents_loaded",
        period_id=period_id,
        packet_id=packet_id,
        attachment_count=len(packet_attachments),
    )
    for attachment in packet_attachments:
        source = Path(str(attachment["local_path"]))
        if not source.exists():
            continue
        user_email = normalize_email(str(attachment["user_email"]))
        if str(attachment["content_type"]) == "application/pdf":
            support_pdf_paths_by_user.setdefault(user_email, []).append(str(source))
            continue
        target = support_dir / f"{attachment['id']}.pdf"
        image_to_pdf(str(source), str(target))
        support_pdf_paths_by_user.setdefault(user_email, []).append(str(target))

    for index, (email, rows) in enumerate(ordered_groups, start=1):
        _pay_lock_log(
            "generate_person_voucher",
            period_id=period_id,
            packet_id=packet_id,
            user_email=email,
            ordinal=index,
            entry_count=len(rows),
        )
        label = safe_filename(rows[0].get("display_name") or email, fallback=f"voucher-{index}")
        docx_path = voucher_dir / f"{index:02d}-{label}.docx"
        anchor_docx_path = voucher_dir / f"{index:02d}-{label}.anchor.docx"
        filer_signer_index = 0
        fill_pay_voucher_docx(
            template_path=cfg.pay_portal.voucher_template_path,
            output_path=str(docx_path),
            period_start=period_start,
            period_end=period_end,
            entries=rows,
            include_signature_placeholders=False,
        )
        fill_pay_voucher_docx(
            template_path=cfg.pay_portal.voucher_template_path,
            output_path=str(anchor_docx_path),
            period_start=period_start,
            period_end=period_end,
            entries=rows,
            include_signature_placeholders=True,
            filer_signer_index=filer_signer_index,
            paid_by_signer_index=president_signer_index,
        )
        voucher_pdf = docx_to_pdf_func(
            str(docx_path),
            str(voucher_dir),
            cfg.libreoffice_timeout_seconds,
            engine=cfg.docx_pdf_engine,
            graph_uploader=graph,
            graph_site_hostname=_pay_graph_site(cfg=cfg)[0],
            graph_site_path=_pay_graph_site(cfg=cfg)[1],
            graph_library=cfg.graph.document_library,
            graph_temp_folder_path=_pay_docx_pdf_graph_temp_folder(cfg=cfg),
        )
        anchor_pdf = docx_to_pdf_func(
            str(anchor_docx_path),
            str(voucher_dir),
            cfg.libreoffice_timeout_seconds,
            engine=cfg.docx_pdf_engine,
            graph_uploader=graph,
            graph_site_hostname=_pay_graph_site(cfg=cfg)[0],
            graph_site_path=_pay_graph_site(cfg=cfg)[1],
            graph_library=cfg.graph.document_library,
            graph_temp_folder_path=_pay_docx_pdf_graph_temp_folder(cfg=cfg),
        )
        voucher_paths.append(str(docx_path))
        voucher_pdf_paths.append(voucher_pdf)
        anchor_pdf_paths.append(anchor_pdf)
        support_pdf_paths = support_pdf_paths_by_user.get(normalize_email(email), [])
        packet_pdf_paths.extend([voucher_pdf, *support_pdf_paths])
        alignment_pdf_paths.extend([anchor_pdf, *support_pdf_paths])
        _pay_lock_log(
            "person_voucher_generated",
            period_id=period_id,
            packet_id=packet_id,
            user_email=email,
            support_pdf_count=len(support_pdf_paths),
        )

    _pay_lock_log(
        "merge_packet",
        period_id=period_id,
        packet_id=packet_id,
        packet_pdf_count=len(packet_pdf_paths),
        alignment_pdf_count=len(alignment_pdf_paths),
    )
    unsigned_packet_path = str(packet_dir / f"{period_start}_to_{period_end}_packet.pdf")
    alignment_packet_path = str(packet_dir / f"{period_start}_to_{period_end}_alignment.pdf")
    merge_pdfs(packet_pdf_paths, unsigned_packet_path)
    merge_pdfs(alignment_pdf_paths, alignment_packet_path)
    packet_bytes = Path(unsigned_packet_path).read_bytes()
    alignment_bytes = Path(alignment_packet_path).read_bytes()
    sha = hashlib.sha256(packet_bytes).hexdigest()
    _pay_lock_log(
        "packet_merged",
        period_id=period_id,
        packet_id=packet_id,
        packet_bytes=len(packet_bytes),
        alignment_bytes=len(alignment_bytes),
    )

    folder_path = pay_period_folder_path(
        root_folder=cfg.pay_portal.sharepoint_root_folder,
        period_start=period_start,
        period_end=period_end,
    )
    _pay_lock_log(
        "upload_unsigned_packet",
        period_id=period_id,
        packet_id=packet_id,
        folder_path=folder_path,
    )
    sharepoint_unsigned_url, _ = await _upload_if_configured(
        cfg=cfg,
        graph=graph,
        folder_path=folder_path,
        filename=Path(unsigned_packet_path).name,
        local_path=unsigned_packet_path,
    )
    _pay_lock_log(
        "unsigned_packet_uploaded",
        period_id=period_id,
        packet_id=packet_id,
        uploaded=bool(sharepoint_unsigned_url),
    )
    _pay_lock_log(
        "upload_generated_vouchers",
        period_id=period_id,
        packet_id=packet_id,
        voucher_count=len(voucher_paths),
    )
    for path in voucher_paths:
        await _upload_if_configured(
            cfg=cfg,
            graph=graph,
            folder_path="/".join((folder_path, "Generated")),
            filename=Path(path).name,
            local_path=path,
        )
    _pay_lock_log("generated_vouchers_uploaded", period_id=period_id, packet_id=packet_id)
    _pay_lock_log(
        "upload_receipts_and_mileage",
        period_id=period_id,
        packet_id=packet_id,
        attachment_count=len(packet_attachments),
    )
    for attachment in packet_attachments:
        web_url, sp_path = await _upload_if_configured(
            cfg=cfg,
            graph=graph,
            folder_path="/".join((folder_path, "Receipts and Mileage")),
            filename=safe_filename(attachment["original_filename"], fallback=str(attachment["id"])),
            local_path=str(attachment["local_path"]),
        )
        if web_url or sp_path:
            await db.exec(
                "UPDATE pay_attachments SET sharepoint_url=?, sharepoint_path=? WHERE id=?",
                (web_url, sp_path, attachment["id"]),
            )

    _pay_lock_log("receipts_and_mileage_uploaded", period_id=period_id, packet_id=packet_id)
    compensation_stubs = await _compensation_stubs_for_packet(db, period_id)
    _pay_lock_log(
        "upload_lost_wage_proof",
        period_id=period_id,
        packet_id=packet_id,
        stub_count=len(compensation_stubs),
    )
    for stub in compensation_stubs:
        web_url, sp_path = await _upload_if_configured(
            cfg=cfg,
            graph=graph,
            folder_path="/".join((folder_path, "Lost Wage Proof")),
            filename=safe_filename(stub["original_filename"], fallback=str(stub["id"])),
            local_path=str(stub["local_path"]),
        )
        if web_url or sp_path:
            await db.exec(
                "UPDATE pay_compensation_stubs SET sharepoint_url=?, sharepoint_path=? WHERE id=?",
                (web_url, sp_path, stub["id"]),
            )

    _pay_lock_log("lost_wage_proof_uploaded", period_id=period_id, packet_id=packet_id)
    _pay_lock_log(
        "docuseal_create_submission",
        period_id=period_id,
        packet_id=packet_id,
        signer_count=len(signer_order),
    )
    submission = docuseal.create_submission(
        pdf_bytes=packet_bytes,
        alignment_pdf_bytes=alignment_bytes,
        signers=signer_order,
        title=f"Local 3106 Pay Packet {period_start} to {period_end}",
        metadata={
            "pay_period_id": period_id,
            "pay_packet_id": packet_id,
            "form_key": _PAY_FORM_KEY,
            "president_signer_email": president_signer,
            "signer_order": signer_order,
        },
        template_id=resolve_docuseal_template_id(cfg, template_key=_PAY_FORM_KEY, doc_type=_PAY_FORM_KEY),
        form_key=_PAY_FORM_KEY,
    )
    _pay_lock_log(
        "docuseal_submission_created",
        period_id=period_id,
        packet_id=packet_id,
        docuseal_submission_id=submission.submission_id,
    )
    signer_links_by_email: dict[str, str] = {}
    _pay_lock_log("docuseal_resolve_signing_links", period_id=period_id, packet_id=packet_id)
    try:
        signer_links_by_email = docuseal.extract_signing_links_by_email(submission.raw)
    except Exception:
        signer_links_by_email = {}
    if not signer_links_by_email:
        try:
            signer_links_by_email = docuseal.fetch_signing_links_by_email(submission_id=submission.submission_id)
        except Exception:
            signer_links_by_email = {}
    first_signer = signer_order[0]
    signing_link = signer_links_by_email.get(first_signer.lower()) or submission.signing_link
    _pay_lock_log(
        "docuseal_signing_link_resolved",
        period_id=period_id,
        packet_id=packet_id,
        first_signer=first_signer,
        has_signing_link=bool(signing_link),
    )

    now = utcnow()
    _pay_lock_log("record_packet", period_id=period_id, packet_id=packet_id)
    await db.exec(
        """INSERT INTO pay_packets(
             id, period_id, revision, status, voucher_paths_json, voucher_pdf_paths_json,
             unsigned_packet_path, unsigned_packet_sha256, docuseal_submission_id,
             docuseal_signing_link, sharepoint_unsigned_url, created_at_utc, updated_at_utc
           ) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            packet_id,
            period_id,
            int(period["revision"]),
            "awaiting_signature",
            json.dumps(voucher_paths),
            json.dumps(voucher_pdf_paths),
            unsigned_packet_path,
            sha,
            submission.submission_id,
            signing_link,
            sharepoint_unsigned_url,
            now,
            now,
        ),
    )
    await db.exec(
        """UPDATE pay_periods
           SET status='awaiting_signature', locked_by=?, locked_at_utc=?, president_email=?,
               sharepoint_folder_path=?, updated_at_utc=?
           WHERE id=?""",
        (actor.email, now, president_signer, folder_path, now, period_id),
    )
    await db.exec(
        """UPDATE pay_entries
              SET locked_at_utc=?
            WHERE period_id=? AND locked_at_utc IS NULL
              AND COALESCE(review_status, 'pending') IN ('pending', 'approved')""",
        (now, period_id),
    )
    _pay_lock_log("packet_recorded", period_id=period_id, packet_id=packet_id)
    _pay_lock_log(
        "email_president",
        period_id=period_id,
        packet_id=packet_id,
        president_signer=president_signer,
        email_enabled=bool(getattr(cfg.email, "enabled", False)),
    )
    president_email_delivery = _send_president_signature_email(
        cfg=cfg,
        mailer=mailer,
        president_email=president_signer,
        period_start=period_start,
        period_end=period_end,
        signing_link=signing_link,
        sharepoint_unsigned_url=sharepoint_unsigned_url,
        folder_path=folder_path,
        packet_id=packet_id,
    )
    _pay_lock_log(
        "president_email_finished",
        period_id=period_id,
        packet_id=packet_id,
        sent=bool(president_email_delivery),
    )
    await add_pay_event(
        db,
        period_id=period_id,
        packet_id=packet_id,
        event_type="period_locked_packet_sent",
        actor=actor.email,
        details={
            "docuseal_submission_id": submission.submission_id,
            "signing_link": signing_link,
            "president_signer_email": president_signer,
            "signer_order": signer_order,
            "excluded_entries": excluded_entries,
            "president_email_delivery": president_email_delivery,
        },
    )
    _pay_lock_log(
        "complete",
        period_id=period_id,
        packet_id=packet_id,
        docuseal_submission_id=submission.submission_id,
        president_email_sent=bool(president_email_delivery),
    )
    return {
        "packet_id": packet_id,
        "period_id": period_id,
        "status": "awaiting_signature",
        "docuseal_submission_id": submission.submission_id,
        "signing_link": signing_link,
        "signer_order": signer_order,
        "sharepoint_unsigned_url": sharepoint_unsigned_url,
        "excluded_entries": excluded_entries,
        "president_email_sent": bool(president_email_delivery),
        "president_email_delivery": president_email_delivery,
    }


async def create_revision(db: Db, *, period_id: str, actor: PayActor) -> dict[str, object]:
    if not actor.can_lock:
        raise PermissionError("treasurer access required")
    period = await get_pay_period(db, period_id)
    if not period:
        raise ValueError("pay period not found")
    start = str(period["period_start"])
    end = str(period["period_end"])
    next_revision = int(period["revision"]) + 1
    new_id = period_id_for(date.fromisoformat(start), date.fromisoformat(end), next_revision)
    now = utcnow()
    await db.exec(
        """INSERT INTO pay_periods(id, period_start, period_end, status, revision, created_at_utc, updated_at_utc)
           VALUES(?,?,?,?,?,?,?)""",
        (new_id, start, end, "open", next_revision, now, now),
    )
    await add_pay_event(
        db,
        period_id=new_id,
        event_type="revision_created",
        actor=actor.email,
        details={"source_period_id": period_id, "revision": next_revision},
    )
    return await get_pay_period(db, new_id) or {"id": new_id}


async def handle_pay_docuseal_completion(
    *,
    db: Db,
    cfg: Any,
    graph: Any,
    mailer: Any,
    docuseal: Any,
    submission_id: str,
    payload: dict[str, object],
) -> dict[str, object] | None:
    _ = payload
    row = await db.fetchone(
        """SELECT p.id, p.period_id, p.status, pp.period_start, pp.period_end, pp.sharepoint_folder_path
           FROM pay_packets p
           JOIN pay_periods pp ON pp.id = p.period_id
           WHERE p.docuseal_submission_id=?""",
        (submission_id,),
    )
    if not row:
        return None
    packet_id, period_id, status, period_start, period_end, folder_path = row
    if str(status) == "completed":
        return {"ok": True, "deduped": True, "pay_packet_id": packet_id}

    packet_dir = Path(cfg.data_root) / "pay" / str(period_id) / "packet" / str(packet_id)
    packet_dir.mkdir(parents=True, exist_ok=True)
    signed_path: str | None = None
    audit_path: str | None = None
    signed_bytes: bytes | None = None
    artifacts = docuseal.download_completed_artifacts(submission_id=submission_id)
    zip_bytes = artifacts.get("completed_zip_bytes")
    if isinstance(zip_bytes, (bytes, bytearray)) and zip_bytes:
        audit_path = str(packet_dir / "docuseal_completed.zip")
        Path(audit_path).write_bytes(bytes(zip_bytes))
        signed_bytes = _extract_first_pdf(bytes(zip_bytes))
    if not signed_bytes:
        direct_pdf_bytes = artifacts.get("signed_pdf_bytes")
        if isinstance(direct_pdf_bytes, (bytes, bytearray)) and direct_pdf_bytes:
            signed_bytes = bytes(direct_pdf_bytes)
    if signed_bytes:
        signed_path = str(packet_dir / f"{period_start}_to_{period_end}_signed.pdf")
        Path(signed_path).write_bytes(signed_bytes)
    else:
        await add_pay_event(
            db,
            period_id=period_id,
            packet_id=packet_id,
            event_type="docuseal_completion_missing_signed_artifact",
            actor="docuseal",
            details={
                "docuseal_submission_id": submission_id,
                "completed_zip_present": isinstance(zip_bytes, (bytes, bytearray)) and bool(zip_bytes),
                "signed_pdf_present": False,
            },
        )
        raise RuntimeError("DocuSeal completion did not include a signed pay packet PDF")

    target_folder = str(folder_path or "").strip() or pay_period_folder_path(
        root_folder=cfg.pay_portal.sharepoint_root_folder,
        period_start=str(period_start),
        period_end=str(period_end),
    )
    signed_url: str | None = None
    audit_url: str | None = None
    if signed_path:
        signed_url, _ = await _upload_if_configured(
            cfg=cfg,
            graph=graph,
            folder_path=target_folder,
            filename=Path(signed_path).name,
            local_path=signed_path,
        )
    if audit_path:
        audit_url, _ = await _upload_if_configured(
            cfg=cfg,
            graph=graph,
            folder_path="/".join((target_folder, "Audit")),
            filename=Path(audit_path).name,
            local_path=audit_path,
        )

    now = utcnow()
    await db.exec(
        """UPDATE pay_packets
           SET status='completed', signed_packet_path=?, audit_zip_path=?,
               sharepoint_signed_url=?, sharepoint_audit_url=?, completed_at_utc=?, updated_at_utc=?
           WHERE id=?""",
        (signed_path, audit_path, signed_url, audit_url, now, now, packet_id),
    )
    await db.exec(
        """UPDATE pay_periods
           SET status='completed', completed_at_utc=?, sharepoint_folder_path=?, updated_at_utc=?
           WHERE id=?""",
        (now, target_folder, now, period_id),
    )
    await add_pay_event(
        db,
        period_id=period_id,
        packet_id=packet_id,
        event_type="docuseal_completion_processed",
        actor="docuseal",
        details={"docuseal_submission_id": submission_id, "signed_url": signed_url, "audit_url": audit_url},
    )

    if mailer is not None and cfg.email.enabled:
        recipients = await treasurer_recipients(
            db,
            fallback=cfg.email.internal_recipients,
            pay_cfg=cfg.pay_portal,
        )
        attachments = None
        if signed_bytes and len(signed_bytes) <= cfg.email.max_attachment_bytes:
            attachments = [
                MailAttachment(
                    filename=Path(signed_path or "pay_packet_signed.pdf").name,
                    content_type="application/pdf",
                    content_bytes=signed_bytes,
                )
            ]
        if recipients:
            mailer.send_mail(
                to_recipients=recipients,
                subject=f"Pay packet signed: {period_start} to {period_end}",
                text_body=(
                    f"The pay packet for {period_start} to {period_end} has been signed.\n\n"
                    f"SharePoint folder: {target_folder}\n"
                    f"Signed packet: {signed_url or signed_path or 'unavailable'}\n"
                ),
                attachments=attachments,
            )

    return {"ok": True, "handled": True, "pay_packet_id": packet_id, "period_id": period_id}


def _parse_decimal_input(text: object) -> tuple[Decimal, int]:
    cleaned = str(text or "").replace("$", "").replace(",", "").strip()
    if not re.fullmatch(r"[+-]?(?:\d+(?:\.\d+)?|\.\d+)", cleaned):
        raise ValueError("value must be a number")
    value = Decimal(cleaned)
    places = len(cleaned.split(".", 1)[1]) if "." in cleaned else 0
    return value, places


def _format_decimal(value: Decimal, places: int) -> str:
    quant = Decimal(1).scaleb(-places)
    return f"{value.quantize(quant, rounding=ROUND_HALF_UP):f}"


def parse_rate_user_input(rate_text: object) -> tuple[Decimal, str]:
    value, places = _parse_decimal_input(rate_text)
    if value <= 0:
        raise ValueError("rate must be greater than 0")
    if abs(value) >= Decimal("10"):
        rate = value / Decimal("100")
        return rate, _format_decimal(rate, places + 2)
    return value, _format_decimal(value, places or 2)


def mileage_rate_from_settings(settings: dict[str, object], year: int) -> tuple[Decimal, str]:
    rates = settings.get("irs_rates")
    raw = None
    if isinstance(rates, dict):
        raw = rates.get(str(year))
    if raw is None:
        raw = "0.67"
    return parse_rate_user_input(raw)


def _google_leg_cache_key(*, api_key: str, origin: str, destination: str) -> tuple[str, str]:
    key_hash = hashlib.sha256(str(api_key or "").encode("utf-8")).hexdigest()[:12]
    route_key = json.dumps([str(origin or "").strip(), str(destination or "").strip()], sort_keys=True)
    return key_hash, route_key


def _google_leg(*, api_key: str, origin: str, destination: str) -> dict[str, object]:
    cache_key = _google_leg_cache_key(api_key=api_key, origin=origin, destination=destination)
    with _GOOGLE_LEG_CACHE_LOCK:
        cached = _GOOGLE_LEG_CACHE.get(cache_key)
    if cached is not None:
        return dict(cached)

    directions_url = "https://maps.googleapis.com/maps/api/directions/json"
    resp = requests.get(
        directions_url,
        params={"origin": origin, "destination": destination, "mode": "driving", "key": api_key},
        timeout=30,
    )
    data = resp.json()
    if data.get("status") != "OK" or not data.get("routes"):
        raise RuntimeError(f"Directions failed between '{origin}' and '{destination}': {data.get('status')}")
    route = data["routes"][0]
    leg = route["legs"][0]
    distance_meters = int((leg.get("distance") or {}).get("value", 0) or 0)
    distance_miles = Decimal(distance_meters) / _METERS_PER_MILE if distance_meters else Decimal("0")
    steps = []
    for step in leg.get("steps", []):
        text = html.unescape(re.sub("<.*?>", "", str(step.get("html_instructions") or ""))).strip()
        if text:
            steps.append(text)
    map_bytes = None
    polyline = (route.get("overview_polyline") or {}).get("points")
    if polyline:
        static_url = "https://maps.googleapis.com/maps/api/staticmap"
        static_params = [
            ("size", "600x400"),
            ("path", f"enc:{polyline}"),
            ("markers", f"color:blue|label:S|{leg.get('start_address', origin)}"),
            ("markers", f"color:red|label:E|{leg.get('end_address', destination)}"),
            ("key", api_key),
        ]
        try:
            map_resp = requests.get(static_url, params=static_params, timeout=30)
            if map_resp.status_code == 200 and map_resp.content:
                map_bytes = map_resp.content
        except requests.RequestException:
            map_bytes = None
    result = {
        "origin": leg.get("start_address", origin),
        "destination": leg.get("end_address", destination),
        "distance_text": (leg.get("distance") or {}).get("text") or "",
        "distance_miles": distance_miles,
        "turn_by_turn": steps,
        "map_bytes": map_bytes,
    }
    with _GOOGLE_LEG_CACHE_LOCK:
        _GOOGLE_LEG_CACHE[cache_key] = dict(result)
    return result


def validate_mileage_locations(*, google_maps_api_key: str, locations: list[str]) -> dict[str, object]:
    if not google_maps_api_key:
        raise RuntimeError("Google Maps API key is not configured")
    cleaned = [str(item or "").strip() for item in locations if str(item or "").strip()]
    if len(cleaned) < 2:
        raise ValueError("at least two locations are required")
    legs: list[dict[str, object]] = []
    resolved = list(cleaned)
    total_miles = Decimal("0")
    for idx in range(len(cleaned) - 1):
        leg = _google_leg(api_key=google_maps_api_key, origin=cleaned[idx], destination=cleaned[idx + 1])
        origin = str(leg.get("origin") or cleaned[idx]).strip()
        destination = str(leg.get("destination") or cleaned[idx + 1]).strip()
        resolved[idx] = origin
        resolved[idx + 1] = destination
        miles = _quantity(leg.get("distance_miles"))
        total_miles += miles
        legs.append(
            {
                "origin_input": cleaned[idx],
                "destination_input": cleaned[idx + 1],
                "origin": origin,
                "destination": destination,
                "distance_text": leg.get("distance_text") or "",
                "distance_miles": float(miles.quantize(_MILES, rounding=ROUND_HALF_UP)),
            }
        )
    return {
        "locations": resolved,
        "legs": legs,
        "total_miles": float(total_miles.quantize(_MILES, rounding=ROUND_HALF_UP)),
    }


def _prefetch_google_legs(
    *,
    api_key: str,
    routes: list[list[str]],
    progress_callback: Any = None,
    max_workers: int = _GOOGLE_LEG_PREFETCH_WORKERS,
) -> None:
    if not api_key:
        return
    pairs: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for locations in routes:
        for idx in range(len(locations) - 1):
            pair = (locations[idx], locations[idx + 1])
            if pair in seen:
                continue
            seen.add(pair)
            pairs.append(pair)
    missing: list[tuple[str, str]] = []
    with _GOOGLE_LEG_CACHE_LOCK:
        for origin, destination in pairs:
            key = _google_leg_cache_key(api_key=api_key, origin=origin, destination=destination)
            if key not in _GOOGLE_LEG_CACHE:
                missing.append((origin, destination))
    if not missing:
        _emit_progress(
            progress_callback,
            stage="mileage",
            current=len(pairs),
            total=len(pairs),
            message="Mileage route cache ready",
        )
        return

    completed = 0
    worker_count = max(1, min(int(max_workers or 1), len(missing)))
    with ThreadPoolExecutor(max_workers=worker_count) as executor:
        futures = {
            executor.submit(_google_leg, api_key=api_key, origin=origin, destination=destination): (origin, destination)
            for origin, destination in missing
        }
        for future in as_completed(futures):
            future.result()
            completed += 1
            _emit_progress(
                progress_callback,
                stage="mileage",
                current=completed,
                total=len(missing),
                message=f"Fetched mileage route {completed}/{len(missing)}",
            )


def build_mileage_pdf(
    *,
    name: str,
    local_number: str,
    date_str: str,
    description: str,
    rate: Decimal,
    rate_display: str,
    locations: list[str],
    google_maps_api_key: str,
) -> tuple[bytes, Decimal, Decimal]:
    try:
        from reportlab.lib.units import inch
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Image as ReportLabImage
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:  # pragma: no cover - runtime packaging guard
        raise RuntimeError("reportlab is required for mileage PDF generation") from exc

    if not google_maps_api_key:
        raise RuntimeError("Google Maps API key is not configured")
    if len(locations) < 2:
        raise ValueError("at least two locations are required")

    legs = []
    total_miles = Decimal("0")
    for idx in range(len(locations) - 1):
        leg = _google_leg(
            api_key=google_maps_api_key,
            origin=locations[idx],
            destination=locations[idx + 1],
        )
        total_miles += leg["distance_miles"]  # type: ignore[operator]
        legs.append(leg)

    reimbursement = (total_miles * rate).quantize(_CURRENCY, rounding=ROUND_HALF_UP)
    total_miles_display = total_miles.quantize(_MILES, rounding=ROUND_HALF_UP)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Mileage Report", styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"Name: {name}", styles["Normal"]),
    ]
    if local_number:
        story.append(Paragraph(f"Local: {local_number}", styles["Normal"]))
    story.extend(
        [
            Paragraph(f"Date: {date_str}", styles["Normal"]),
            Paragraph(f"Description: {description}", styles["Normal"]),
            Spacer(1, 12),
            Paragraph(f"IRS Standard Mileage Rate: ${rate_display} per mile", styles["Normal"]),
            Paragraph(f"Total Distance: {_decimal_text(total_miles_display)} miles", styles["Normal"]),
            Paragraph(f"Total Reimbursement: ${_decimal_text(reimbursement)}", styles["Normal"]),
            Spacer(1, 18),
        ]
    )
    for index, leg in enumerate(legs, start=1):
        story.extend(
            [
                Paragraph(f"Route {index}:", styles["Heading3"]),
                Paragraph(f"From: {leg['origin']}", styles["Normal"]),
                Paragraph(f"To: {leg['destination']}", styles["Normal"]),
                Paragraph(
                    f"Distance: {leg.get('distance_text') or _decimal_text(leg['distance_miles']) + ' mi'}",
                    styles["Normal"],
                ),
                Spacer(1, 8),
            ]
        )
    story.extend(
        [
            Spacer(1, 12),
            Paragraph("Distances computed by Google Directions API on the server side.", styles["Normal"]),
        ]
    )
    story.append(PageBreak())
    for index, leg in enumerate(legs, start=1):
        story.extend(
            [
                Paragraph(f"Route {index} Detail", styles["Heading2"]),
                Spacer(1, 6),
                Paragraph(f"From: {leg['origin']}", styles["Normal"]),
                Paragraph(f"To: {leg['destination']}", styles["Normal"]),
                Paragraph(
                    f"Distance: {leg.get('distance_text') or _decimal_text(leg['distance_miles']) + ' mi'}",
                    styles["Normal"],
                ),
                Spacer(1, 12),
            ]
        )
        if leg.get("map_bytes"):
            map_image = ReportLabImage(io.BytesIO(leg["map_bytes"]))  # type: ignore[arg-type]
            map_image.drawWidth = 6 * inch
            map_image.drawHeight = 4 * inch
            story.extend([map_image, Spacer(1, 12)])
        story.append(Paragraph("Turn-by-Turn Directions:", styles["Heading3"]))
        for step in leg.get("turn_by_turn", []):
            story.append(Paragraph(str(step), styles["Normal"]))
        story.append(PageBreak())
    doc.build(story)
    return buf.getvalue(), reimbursement, total_miles_display


async def create_mileage_attachment(
    *,
    db: Db,
    cfg: Any,
    period_id: str,
    entry_id: str,
    actor: PayActor,
    name: str,
    local_number: str,
    date_str: str,
    description: str,
    locations: list[str],
    rate_text: str | None,
    graph: Any | None = None,
) -> dict[str, object]:
    settings = await pay_settings(db, pay_cfg=cfg.pay_portal)
    name = str(name or "").strip() or str(actor.display_name or actor.email or "Pay User").strip()
    year = datetime.strptime(date_str, "%Y-%m-%d").year
    rate, rate_display = parse_rate_user_input(rate_text) if rate_text else mileage_rate_from_settings(settings, year)
    pdf_bytes, reimbursement, total_miles = build_mileage_pdf(
        name=name,
        local_number=local_number,
        date_str=date_str,
        description=description,
        rate=rate,
        rate_display=rate_display,
        locations=locations,
        google_maps_api_key=cfg.pay_portal.google_maps_api_key,
    )
    yyyymmdd = date_str.replace("-", "")
    filename = f"{yyyymmdd} {name}.pdf"
    attachment = await store_attachment(
        db,
        cfg=cfg,
        period_id=period_id,
        entry_id=entry_id,
        actor=actor,
        attachment_type="mileage_pdf",
        filename=filename,
        content_type="application/pdf",
        content=pdf_bytes,
        scan=False,
        mileage_miles=total_miles,
        mileage_rate=rate,
        mileage_amount=reimbursement,
    )
    legacy_sharepoint: dict[str, str | None] = {}
    legacy_root = str(getattr(cfg.pay_portal, "mileage_legacy_sharepoint_folder", "") or "").strip()
    if graph is not None and legacy_root:
        stored = await db.fetchone("SELECT local_path FROM pay_attachments WHERE id=?", (attachment["id"],))
        local_path = str(stored[0]) if stored else ""
        if local_path:
            legacy_folder_path = mileage_legacy_report_folder_path(root_folder=legacy_root, report_date=date_str)
            try:
                web_url, sp_path = await _upload_if_configured(
                    cfg=cfg,
                    graph=graph,
                    folder_path=legacy_folder_path,
                    filename=filename,
                    local_path=local_path,
                )
                if web_url or sp_path:
                    legacy_sharepoint = {
                        "legacy_sharepoint_url": web_url,
                        "legacy_sharepoint_path": sp_path,
                    }
                    await add_pay_event(
                        db,
                        period_id=period_id,
                        entry_id=entry_id,
                        event_type="mileage_legacy_sharepoint_uploaded",
                        actor=actor.email,
                        details={"filename": filename, "sharepoint_path": sp_path},
                    )
            except Exception as exc:  # noqa: BLE001 - mileage generation should survive legacy copy failures.
                await add_pay_event(
                    db,
                    period_id=period_id,
                    entry_id=entry_id,
                    event_type="mileage_legacy_sharepoint_upload_failed",
                    actor=actor.email,
                    details={"filename": filename, "error": str(exc)[:500]},
                )
    row = await db.fetchone(
        "SELECT mileage_miles, mileage_amount FROM pay_entries WHERE id=? AND period_id=?",
        (entry_id, period_id),
    )
    current_mileage_miles = _quantity(row[0] if row else 0)
    current_mileage_amount = _money(row[1] if row else 0)
    await db.exec(
        """UPDATE pay_entries
           SET mileage_miles=?, mileage_rate=?, mileage_amount=?, updated_at_utc=?
           WHERE id=? AND period_id=?""",
        (
            float(current_mileage_miles + total_miles),
            float(rate),
            float(current_mileage_amount + reimbursement),
            utcnow(),
            entry_id,
            period_id,
        ),
    )
    return {
        **attachment,
        **legacy_sharepoint,
        "mileage_miles": float(total_miles),
        "mileage_rate": float(rate),
        "mileage_amount": float(reimbursement),
        "reimbursement": float(reimbursement),
    }
