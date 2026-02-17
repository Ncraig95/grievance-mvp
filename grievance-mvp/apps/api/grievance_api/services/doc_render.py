from __future__ import annotations

import html
import re
import tempfile
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docxtpl import DocxTemplate


_JINJA_PLACEHOLDER_RE = re.compile(r"{{\s*([^{}]+?)\s*}}")
_LEFTOVER_PLACEHOLDER_RE = re.compile(r"{{.*?}}", flags=re.DOTALL)
_XML_TAG_RE = re.compile(r"<[^>]+>")


def _escape_signature_placeholders(xml_text: str) -> str:
    def _replace(match: re.Match[str]) -> str:
        expr = match.group(1).strip()
        # DocuSeal signature/date tags use ":" and are not valid Jinja variables.
        if ":" in expr:
            return "{% raw %}" + match.group(0) + "{% endraw %}"
        return match.group(0)

    return _JINJA_PLACEHOLDER_RE.sub(_replace, xml_text)


def _prepare_template_docx(template_path: str) -> str:
    source = Path(template_path)
    if not source.exists():
        return template_path

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with ZipFile(source, "r") as zin, ZipFile(tmp_path, "w") as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    patched = _escape_signature_placeholders(data.decode("utf-8", errors="ignore"))
                    data = patched.encode("utf-8")
                zout.writestr(info, data)
        return str(tmp_path)
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise


def _replace_leftover_placeholders(xml_text: str, context: dict) -> str:
    def _replace(match: re.Match[str]) -> str:
        raw = match.group(0)
        inner = _XML_TAG_RE.sub("", raw[2:-2]).strip()
        if not inner or ":" in inner:
            return raw
        if inner not in context:
            return raw
        value = context.get(inner)
        if value is None:
            return ""
        return html.escape(str(value), quote=False)

    return _LEFTOVER_PLACEHOLDER_RE.sub(_replace, xml_text)


def _postprocess_rendered_docx(out_path: str, context: dict) -> None:
    source = Path(out_path)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with ZipFile(source, "r") as zin, ZipFile(tmp_path, "w") as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    patched = _replace_leftover_placeholders(data.decode("utf-8", errors="ignore"), context)
                    data = patched.encode("utf-8")
                zout.writestr(info, data)
        tmp_path.replace(source)
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise


def render_docx(template_path: str, context: dict, out_path: str) -> None:
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    patched_template_path: str | None = None
    try:
        patched_template_path = _prepare_template_docx(template_path)
        tpl = DocxTemplate(patched_template_path)
        tpl.render(context)
        tpl.save(out_path)
        _postprocess_rendered_docx(out_path, context)
        return
    except Exception:
        # Fallback for missing/invalid template files during bootstrap/testing.
        doc = Document()
        doc.add_heading("Grievance Document", level=1)
        for key, value in context.items():
            doc.add_paragraph(f"{key}: {value}")
        doc.save(out_path)
    finally:
        if patched_template_path and patched_template_path != template_path:
            Path(patched_template_path).unlink(missing_ok=True)
