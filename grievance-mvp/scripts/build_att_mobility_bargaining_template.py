from __future__ import annotations

import argparse
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE_DOC = Path("Docx Files Template/AT&TMobility2023BargainingSuggestion_Form.docx")
DEFAULT_OUTPUT = SOURCE_DOC


def _set_cell_border(cell, **kwargs) -> None:  # noqa: ANN001
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def _set_cell_margins(cell, *, top: int = 0, start: int = 0, bottom: int = 0, end: int = 0) -> None:  # noqa: ANN001
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_row_height(row, *, height_twips: int, rule: WD_ROW_HEIGHT_RULE = WD_ROW_HEIGHT_RULE.AT_LEAST) -> None:  # noqa: ANN001
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "exact" if rule == WD_ROW_HEIGHT_RULE.EXACTLY else "atLeast")


def _set_default_font(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        if style_name in styles:
            styles[style_name].font.name = "Arial"


def _append_field_run(paragraph, instruction: str) -> None:  # noqa: ANN001
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), "18")
    r_pr.append(sz_cs)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld_simple.append(run)
    paragraph._p.append(fld_simple)


def _extract_logo_to_temp(source_doc: Path) -> Path | None:
    if not source_doc.exists():
        return None
    with zipfile.ZipFile(source_doc) as archive:
        names = [name for name in archive.namelist() if name.startswith("word/media/")]
        if not names:
            return None
        suffix = Path(names[0]).suffix or ".png"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(archive.read(names[0]))
            return Path(tmp.name)


def _add_heading(document: Document) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("ONLY ONE CONTRACT ARTICLE PER FORM")
    run.bold = True
    run.underline = True
    run.font.size = Pt(13)
    heading.paragraph_format.space_after = Pt(8)


def _add_top_block(document: Document, logo_path: Path | None) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4.9)

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    if logo_path and logo_path.exists():
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(str(logo_path), width=Inches(0.9))

    nested = right.add_table(rows=1, cols=1)
    nested.alignment = WD_TABLE_ALIGNMENT.RIGHT
    nested.autofit = False
    nested.columns[0].width = Inches(4.7)
    box_cell = nested.cell(0, 0)
    box_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_border(
        box_cell,
        left={"val": "single", "sz": 18, "color": "000000"},
        top={"val": "single", "sz": 18, "color": "000000"},
        right={"val": "single", "sz": 18, "color": "000000"},
        bottom={"val": "single", "sz": 18, "color": "000000"},
    )

    p1 = box_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("FOR LOCAL USE ONLY:")
    r1.bold = True
    r1.font.size = Pt(8.5)

    p2 = box_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run("ARTICLE AFFECTED: ")
    r2.bold = True
    r2.font.size = Pt(10)
    p2.add_run("{{Txt_es_:signer1:article_affected}}")
    p2.add_run(" __________________________")


def _add_title(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("BARGAINING DEMAND AT&T MOBILITY")
    run.bold = True
    run.font.size = Pt(12)


def _add_labeled_line(
    document: Document,
    label: str,
    placeholder: str,
    *,
    align: str = "left",
    space_after_pt: float = 3,
) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == "right" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after_pt)
    label_run = p.add_run(f"{label} ")
    label_run.bold = True
    label_run.font.size = Pt(10.5)
    value_run = p.add_run(f"{{{{ {placeholder} }}}}")
    value_run.underline = True


def _add_dynamic_section(document: Document, title: str, rows_key: str) -> None:
    title_p = document.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(0)
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(10.5)

    table = document.add_table(rows=3, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.45)

    start_row = table.rows[0]
    data_row = table.rows[1]
    end_row = table.rows[2]
    _set_row_height(start_row, height_twips=40)
    _set_row_height(data_row, height_twips=180)
    _set_row_height(end_row, height_twips=40)

    start_cell = start_row.cells[0]
    start_cell.text = f"{{%tr for line in {rows_key} %}}"
    _set_cell_border(
        start_cell,
        left={"val": "nil"},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "nil"},
    )
    _set_cell_margins(start_cell)

    data_cell = data_row.cells[0]
    data_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_border(
        data_cell,
        left={"val": "nil"},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "single", "sz": 8, "color": "000000"},
    )
    _set_cell_margins(data_cell, start=90, end=36)
    body = data_cell.paragraphs[0]
    body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.line_spacing = 1.0
    body.add_run("{{ line.text }}")

    end_cell = end_row.cells[0]
    end_cell.text = "{%tr endfor %}"
    _set_cell_border(
        end_cell,
        left={"val": "nil"},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "nil"},
    )
    _set_cell_margins(end_cell)


def _add_contact_block(document: Document) -> None:
    phones = document.add_paragraph()
    phones.paragraph_format.space_before = Pt(6)
    work = phones.add_run("TELEPHONE: (WORK) ")
    work.bold = True
    phones.add_run("{{ work_phone }}")
    phones.add_run("    (HOME) ").bold = True
    phones.add_run("{{ home_phone }}")

    email = document.add_paragraph()
    email_run = email.add_run("NON-WORK EMAIL ADDRESS: ")
    email_run.bold = True
    email.add_run("{{ non_work_email }}")


def _add_signature_line(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    label_run = p.add_run("SIGNATURE OF LOCAL PRESIDENT: ")
    label_run.bold = True
    p.add_run("{{Sig_es_:signer1:signature}}")
    p.add_run(" ________________________________")


def _add_footer_page_number(document: Document) -> None:
    footer = document.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    prefix = p.runs[0] if p.runs else p.add_run()
    prefix.text = "Page "
    prefix.font.name = "Arial"
    prefix.font.size = Pt(9)
    _append_field_run(p, "PAGE")


def build_template(output_path: Path, source_doc: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.start_type = WD_SECTION.NEW_PAGE
    _set_default_font(document)

    logo_path = _extract_logo_to_temp(source_doc)
    try:
        _add_heading(document)
        _add_top_block(document, logo_path)
        _add_title(document)
        _add_labeled_line(document, "LOCAL#:", "local_number", align="right")
        _add_labeled_line(document, "DEMAND FROM LOCAL:", "demand_from_local", space_after_pt=6)
        _add_labeled_line(
            document,
            "TITLE OF MOBILITY MEMBER SUBMITTING DEMAND:",
            "submitting_member_title",
            space_after_pt=6,
        )
        _add_labeled_line(
            document,
            "NAME OF MOBILITY MEMBER SUBMITTING DEMAND:",
            "submitting_member_name",
            space_after_pt=6,
        )
        _add_dynamic_section(document, "DEMAND:", "demand_rows")
        _add_dynamic_section(document, "REASON:", "reason_rows")
        _add_dynamic_section(document, "SPECIFIC EXAMPLES:", "specific_examples_rows")
        _add_contact_block(document)
        _add_signature_line(document)
        _add_footer_page_number(document)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        tmp_output = output_path.with_name(f".{output_path.name}.tmp")
        try:
            document.save(str(tmp_output))
            tmp_output.replace(output_path)
        finally:
            tmp_output.unlink(missing_ok=True)
    finally:
        if logo_path:
            logo_path.unlink(missing_ok=True)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-doc", default=str(SOURCE_DOC))
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT))
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    build_template(Path(args.output), Path(args.source_doc))


if __name__ == "__main__":
    main()
